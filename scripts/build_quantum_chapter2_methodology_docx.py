from __future__ import annotations

from pathlib import Path
from textwrap import fill
import os

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT_DIR = Path("artifacts")
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR = OUT_DIR / "chapter2_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / "arXiv_양자분야_제2장_분석대상_및_방법론_고품질_재작성_20260704.docx"

NAVY = "12314F"
BLUE = "4C8ED9"
LIGHT_BLUE = "E8F2FA"
GREEN = "55A884"
LIGHT_GREEN = "EAF5EF"
ORANGE = "E99B4C"
LIGHT_ORANGE = "FFF1E3"
PURPLE = "816CC6"
LIGHT_PURPLE = "F0ECFA"
RED = "D9635C"
LIGHT_RED = "FCEAE8"
YELLOW = "F5C85B"
LIGHT_YELLOW = "FFF8DE"
GRAY = "6C7884"
LIGHT_GRAY = "F3F5F7"
DARK = "203040"
WHITE = "FFFFFF"
BORDER = "C9D6E2"

# ---------------------------------------------------------------------------
# Font setup for graphics
# ---------------------------------------------------------------------------
font_candidates = [
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc",
]
font_path = next((p for p in font_candidates if Path(p).exists()), None)
if font_path:
    font_prop = font_manager.FontProperties(fname=font_path)
    plt.rcParams["font.family"] = font_prop.get_name()
else:
    font_prop = None
    plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False


def rgb(hexstr: str) -> RGBColor:
    return RGBColor.from_string(hexstr)


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "sz", "space", "color"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_column_widths(table, widths_cm) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def set_run_font(run, size=10, bold=False, color=DARK, name="맑은 고딕") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = rgb(color)


def set_paragraph_format(paragraph, before=0, after=4, line=1.25, keep=False) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if keep:
        pf.keep_with_next = True
    pPr = paragraph._p.get_or_add_pPr()
    widow = OxmlElement("w:widowControl")
    pPr.append(widow)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_hyperlink(paragraph, text: str, url: str, color="0563C1", underline=True) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    rFonts.set(qn("w:ascii"), "Arial")
    rPr.append(rFonts)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(rPr)
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        set_run_font(r, 16, True, NAVY)
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "14")
        bottom.set(qn("w:space"), "3")
        bottom.set(qn("w:color"), BLUE)
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        set_run_font(r, 12.5, True, NAVY)
    return p


def add_bullet(doc, text: str, level=0, color=DARK, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.55 + level * 0.45)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    set_paragraph_format(p, after=3, line=1.25)
    marker = "•" if level == 0 else "–"
    r = p.add_run(f"{marker} ")
    set_run_font(r, 10.2, True, BLUE)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, 10.2, True, color)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, 10.2, False, color)
    else:
        r = p.add_run(text)
        set_run_font(r, 10.2, False, color)


def add_body(doc, text: str, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    set_paragraph_format(p, after=5, line=1.35)
    r = p.add_run(text)
    set_run_font(r, 10.2, bold, color)


def add_callout(doc, title: str, bullets: list[str], color=LIGHT_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, color)
    set_cell_border(cell, top={"val":"single","sz":8,"color":accent}, bottom={"val":"single","sz":8,"color":accent}, left={"val":"single","sz":8,"color":accent}, right={"val":"single","sz":8,"color":accent})
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, 11, True, NAVY)
    set_paragraph_format(p, after=4)
    for text in bullets:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        set_paragraph_format(p, after=2, line=1.2)
        r = p.add_run("• ")
        set_run_font(r, 9.8, True, accent)
        r = p.add_run(text)
        set_run_font(r, 9.8, False, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(p, before=6, after=3, line=1.0, keep=True)
    r = p.add_run(text)
    set_run_font(r, 9.5, True, NAVY)


def add_figure_caption(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, before=2, after=7, line=1.0)
    r = p.add_run(text)
    set_run_font(r, 9.2, True, GRAY)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], header_fill=NAVY, font_size=8.8) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_column_widths(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_shading(c, header_fill)
        set_cell_border(c, top={"val":"single","sz":6,"color":BORDER}, bottom={"val":"single","sz":6,"color":BORDER}, left={"val":"single","sz":6,"color":BORDER}, right={"val":"single","sz":6,"color":BORDER})
        set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p, after=0, line=1.05)
        r = p.add_run(h)
        set_run_font(r, font_size, True, WHITE)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            set_cell_shading(c, WHITE if r_idx % 2 == 0 else LIGHT_GRAY)
            set_cell_border(c, top={"val":"single","sz":4,"color":BORDER}, bottom={"val":"single","sz":4,"color":BORDER}, left={"val":"single","sz":4,"color":BORDER}, right={"val":"single","sz":4,"color":BORDER})
            set_cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, after=0, line=1.08)
            rr = p.add_run(str(value))
            set_run_font(rr, font_size, i == 0, DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Figure generation helpers
# ---------------------------------------------------------------------------
def save_period_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.8, 3.4), dpi=200)
    ax.axis("off")
    # full collection box
    ax.add_patch(FancyBboxPatch((0.05, 0.55), 0.90, 0.28, boxstyle="round,pad=0.015,rounding_size=0.02", facecolor=LIGHT_BLUE, edgecolor=BLUE, linewidth=2))
    ax.text(0.07, 0.74, "수집 원자료", fontsize=14, fontweight="bold", color="#12314F", va="center")
    ax.text(0.25, 0.74, "2025.06.06 ~ 2026.06.06", fontsize=14, fontweight="bold", color="#12314F", va="center")
    ax.text(0.68, 0.74, "2,427건", fontsize=17, fontweight="bold", color="#12314F", va="center")
    ax.text(0.86, 0.74, "13개월처럼 보이나\n마지막 월은 부분월", fontsize=10.5, fontweight="bold", color="#6C7884", ha="center", va="center")
    # analysis box
    ax.add_patch(FancyBboxPatch((0.05, 0.12), 0.72, 0.28, boxstyle="round,pad=0.015,rounding_size=0.02", facecolor=LIGHT_GREEN, edgecolor=GREEN, linewidth=2))
    ax.text(0.07, 0.31, "수치 분석 DB", fontsize=14, fontweight="bold", color="#12314F", va="center")
    ax.text(0.25, 0.31, "2025.06 ~ 2026.05", fontsize=14, fontweight="bold", color="#12314F", va="center")
    ax.text(0.58, 0.31, "2,400건", fontsize=17, fontweight="bold", color="#12314F", va="center")
    # excluded box
    ax.add_patch(FancyBboxPatch((0.80, 0.12), 0.15, 0.28, boxstyle="round,pad=0.015,rounding_size=0.02", facecolor=LIGHT_RED, edgecolor=RED, linewidth=2))
    ax.text(0.875, 0.31, "2026.06\n27건 제외", fontsize=12, fontweight="bold", color="#A83D37", ha="center", va="center")
    ax.add_patch(FancyArrowPatch((0.50, 0.54), (0.50, 0.42), arrowstyle="-|>", mutation_scale=18, linewidth=2, color=GRAY))
    ax.text(0.48, 0.49, "완전월 정렬", fontsize=10.5, fontweight="bold", color=GRAY, ha="right", va="center")
    ax.text(0.50, 0.02, "월평균 200.0편 · 월별 성장률·최근 모멘텀·분야별 추세의 비교 가능성 확보", fontsize=11, fontweight="bold", color="#203040", ha="center")
    fig.tight_layout(pad=0.2)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def save_search_logic(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.8, 4.4), dpi=200)
    ax.axis("off")
    boxes = [
        (0.04, 0.47, 0.25, 0.36, LIGHT_BLUE, BLUE, "양자 기술 축", "QUANTUM\nQUANTOMETER\nQUANTIZATION"),
        (0.38, 0.47, 0.36, 0.36, LIGHT_ORANGE, ORANGE, "응용·기능 축", "센서·측정 / 시간·주파수\n전자기장·관성 / 컴퓨팅·통신"),
        (0.81, 0.47, 0.15, 0.36, LIGHT_GREEN, GREEN, "수집 결과", "관련성 높은\n코어 논문군"),
    ]
    for x, y, w, h, fc, ec, title, body in boxes:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.025", facecolor=fc, edgecolor=ec, linewidth=2.2))
        ax.text(x+w/2, y+h-0.09, title, fontsize=13, fontweight="bold", color="#12314F", ha="center", va="center")
        ax.text(x+w/2, y+0.13, body, fontsize=11.2, fontweight="bold", color="#203040", ha="center", va="center", linespacing=1.35)
    ax.text(0.335, 0.65, "AND", fontsize=16, fontweight="bold", color="#12314F", ha="center")
    ax.add_patch(FancyArrowPatch((0.29, 0.65), (0.37, 0.65), arrowstyle="-|>", mutation_scale=18, linewidth=2, color=GRAY))
    ax.add_patch(FancyArrowPatch((0.74, 0.65), (0.80, 0.65), arrowstyle="-|>", mutation_scale=18, linewidth=2, color=GRAY))
    # bottom interpretation boxes
    labels = [
        (0.04, 0.10, 0.28, LIGHT_GREEN, GREEN, "강점", "제목 기반 고정밀 선별\n핵심 연구 주제의 직접성 확보"),
        (0.36, 0.10, 0.28, LIGHT_YELLOW, YELLOW, "해석 기준", "전체 양자 연구의 완전 모집단이 아닌\n전략적 코어 분석 모집단으로 해석"),
        (0.68, 0.10, 0.28, LIGHT_RED, RED, "보완 과제", "QUANTIZATION 노이즈 점검\n초록·전문 기반 확장 검색 병행"),
    ]
    for x, y, w, fc, ec, title, body in labels:
        ax.add_patch(FancyBboxPatch((x, y), w, 0.24, boxstyle="round,pad=0.015,rounding_size=0.02", facecolor=fc, edgecolor=ec, linewidth=1.8))
        ax.text(x+0.03, y+0.17, title, fontsize=11.5, fontweight="bold", color="#12314F", va="center")
        ax.text(x+0.03, y+0.07, body, fontsize=9.7, fontweight="bold", color="#203040", va="center", linespacing=1.3)
    fig.tight_layout(pad=0.15)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def save_pipeline(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.8, 5.3), dpi=200)
    ax.axis("off")
    steps = [
        ("1", "검색식 설계", "양자축 × 응용축"),
        ("2", "arXiv 수집", "메타데이터·제출일"),
        ("3", "PDF 확보", "파일·페이지·경로"),
        ("4", "전문 추출", "본문·참고문헌·단어수"),
        ("5", "정제·표준화", "저자·기관·국가·카테고리"),
        ("6", "분석 DB", "12개월·2,400건"),
    ]
    colors = [(LIGHT_BLUE, BLUE), (LIGHT_BLUE, BLUE), (LIGHT_ORANGE, ORANGE), (LIGHT_ORANGE, ORANGE), (LIGHT_PURPLE, PURPLE), (LIGHT_GREEN, GREEN)]
    x0, y, w, h, gap = 0.025, 0.55, 0.135, 0.30, 0.026
    for i, ((num, title, sub), (fc, ec)) in enumerate(zip(steps, colors)):
        x = x0 + i*(w+gap)
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.012,rounding_size=0.025", facecolor=fc, edgecolor=ec, linewidth=2))
        ax.text(x+0.027, y+h-0.055, num, fontsize=14, fontweight="bold", color=WHITE, ha="center", va="center", bbox=dict(boxstyle="circle,pad=0.25", fc=ec, ec=ec))
        ax.text(x+w/2, y+0.17, title, fontsize=11.5, fontweight="bold", color="#12314F", ha="center")
        ax.text(x+w/2, y+0.075, sub, fontsize=9.4, fontweight="bold", color="#5E6974", ha="center", linespacing=1.3)
        if i < len(steps)-1:
            ax.add_patch(FancyArrowPatch((x+w+0.004, y+h/2), (x+w+gap-0.004, y+h/2), arrowstyle="-|>", mutation_scale=14, linewidth=1.6, color=GRAY))
    # output axes
    out_y = 0.10
    out_boxes = [
        (0.04, 0.22, LIGHT_BLUE, BLUE, "연구 규모·성장", "월별 활동·모멘텀"),
        (0.28, 0.22, LIGHT_ORANGE, ORANGE, "분야·카테고리", "코어·융합 확장"),
        (0.52, 0.22, LIGHT_PURPLE, PURPLE, "연구 주체", "저자·기관·이중 국가기준"),
        (0.76, 0.20, LIGHT_GREEN, GREEN, "협업 구조", "공저·다기관·국제협력"),
    ]
    for x, ww, fc, ec, title, sub in out_boxes:
        ax.add_patch(FancyBboxPatch((x, out_y), ww, 0.22, boxstyle="round,pad=0.012,rounding_size=0.02", facecolor=fc, edgecolor=ec, linewidth=1.7))
        ax.text(x+ww/2, out_y+0.14, title, fontsize=11.2, fontweight="bold", color="#12314F", ha="center")
        ax.text(x+ww/2, out_y+0.055, sub, fontsize=9.4, fontweight="bold", color="#5E6974", ha="center")
    ax.add_patch(FancyArrowPatch((0.50, 0.52), (0.50, 0.34), arrowstyle="-|>", mutation_scale=18, linewidth=2, color=GRAY))
    fig.tight_layout(pad=0.1)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def save_integrated_framework(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.8, 5.2), dpi=200)
    ax.axis("off")
    # center
    ax.add_patch(FancyBboxPatch((0.31, 0.31), 0.38, 0.38, boxstyle="round,pad=0.025,rounding_size=0.035", facecolor="#FFFFFF", edgecolor=NAVY, linewidth=2.6))
    ax.text(0.50, 0.59, "정량 분석 코어", fontsize=16, fontweight="bold", color="#12314F", ha="center")
    ax.text(0.50, 0.47, "규모·성장  |  분야·카테고리", fontsize=11.5, fontweight="bold", color="#203040", ha="center")
    ax.text(0.50, 0.39, "연구 주체  |  협업 구조", fontsize=11.5, fontweight="bold", color="#203040", ha="center")
    outer = [
        (0.04, 0.58, 0.22, 0.25, LIGHT_BLUE, BLUE, "시장", "투자·수요·산업 채택"),
        (0.74, 0.58, 0.22, 0.25, LIGHT_ORANGE, ORANGE, "기술", "오류정정·확장성·벤치마크"),
        (0.04, 0.14, 0.22, 0.25, LIGHT_GREEN, GREEN, "정책", "국가전략·예산·조달"),
        (0.74, 0.14, 0.22, 0.25, LIGHT_PURPLE, PURPLE, "표준화", "PQC·QKD·시험·인증"),
    ]
    for x, y, w, h, fc, ec, title, sub in outer:
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.025", facecolor=fc, edgecolor=ec, linewidth=2))
        ax.text(x+w/2, y+0.16, title, fontsize=14, fontweight="bold", color="#12314F", ha="center")
        ax.text(x+w/2, y+0.065, sub, fontsize=10, fontweight="bold", color="#5E6974", ha="center")
    arrows = [((0.26,0.70),(0.31,0.61)),((0.74,0.70),(0.69,0.61)),((0.26,0.27),(0.31,0.39)),((0.74,0.27),(0.69,0.39))]
    for a,b in arrows:
        ax.add_patch(FancyArrowPatch(a,b,arrowstyle="<->",mutation_scale=15,linewidth=1.7,color=GRAY))
    ax.text(0.50, 0.20, "교차검증 원칙: 정량 신호 ↔ 외부 환경 근거 ↔ 한계·대안", fontsize=11.5, fontweight="bold", color="#12314F", ha="center")
    ax.text(0.50, 0.08, "단일 수치의 서술을 넘어 연구량의 원인·기술적 의미·정책 활용성을 함께 해석함", fontsize=10.5, fontweight="bold", color="#5E6974", ha="center")
    fig.tight_layout(pad=0.15)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def save_evidence_hierarchy(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11.8, 4.6), dpi=200)
    ax.axis("off")
    levels = [
        (0.12, 0.68, 0.76, 0.20, NAVY, "1차 공식 근거", "정부·국제표준기관·공공 로드맵", "정책·표준·예산·전환 일정의 기준"),
        (0.18, 0.47, 0.64, 0.17, BLUE, "검증 기술 근거", "동료평가 논문·DARPA 등 독립 벤치마킹", "기술 진전과 성숙도의 검증"),
        (0.24, 0.28, 0.52, 0.15, GREEN, "산업·시장 근거", "시장 모니터·산업협회·생태계 보고서", "채택·투자·수요 변화의 해석"),
        (0.30, 0.10, 0.40, 0.13, ORANGE, "기업 로드맵", "IBM·Google·Quantinuum 등", "목표와 달성 성과를 구분하여 활용"),
    ]
    for x,y,w,h,c,title,source,use in levels:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.016,rounding_size=0.025",facecolor=c,edgecolor=c,linewidth=1.8,alpha=0.95))
        ax.text(x+0.025,y+h*0.62,title,fontsize=12,fontweight="bold",color=WHITE,va="center")
        ax.text(x+w*0.42,y+h*0.62,source,fontsize=10,fontweight="bold",color=WHITE,ha="center",va="center")
        ax.text(x+w-0.025,y+h*0.62,use,fontsize=9.5,fontweight="bold",color=WHITE,ha="right",va="center")
    ax.text(0.50,0.94,"근거자료의 위계와 해석 통제",fontsize=15,fontweight="bold",color="#12314F",ha="center")
    fig.tight_layout(pad=0.15)
    fig.savefig(path, bbox_inches="tight", facecolor="white")
    plt.close(fig)


period_png = ASSET_DIR / "fig_2_1_period.png"
search_png = ASSET_DIR / "fig_2_2_search.png"
pipeline_png = ASSET_DIR / "fig_2_3_pipeline.png"
framework_png = ASSET_DIR / "fig_2_4_framework.png"
hierarchy_png = ASSET_DIR / "fig_2_5_evidence.png"
save_period_figure(period_png)
save_search_logic(search_png)
save_pipeline(pipeline_png)
save_integrated_framework(framework_png)
save_evidence_hierarchy(hierarchy_png)

# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.55)
sec.bottom_margin = Cm(1.45)
sec.left_margin = Cm(1.65)
sec.right_margin = Cm(1.65)
sec.header_distance = Cm(0.65)
sec.footer_distance = Cm(0.65)

styles = doc.styles
styles["Normal"].font.name = "맑은 고딕"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
styles["Normal"].font.size = Pt(10.2)

# Header/footer
header = sec.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("arXiv 양자 분야 논문 분석 | 제2장 분석 대상 및 방법론")
set_run_font(r, 8.5, True, NAVY)
footer = sec.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("분석 기준: 2025.06~2026.05  |  ")
set_run_font(r, 8.5, False, GRAY)
add_page_number(p)

# Cover
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cover.autofit = False
cell = cover.cell(0, 0)
set_cell_shading(cell, NAVY)
set_cell_margins(cell, 700, 350, 700, 350)
set_cell_border(cell, top={"val":"single","sz":0,"color":NAVY}, bottom={"val":"single","sz":0,"color":NAVY}, left={"val":"single","sz":0,"color":NAVY}, right={"val":"single","sz":0,"color":NAVY})
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHAPTER 2")
set_run_font(r, 15, True, "A7C8E8")
p = cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("분석 대상 및 방법론")
set_run_font(r, 27, True, WHITE)
p = cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("arXiv 기반 양자 분야 논문 수집·분석 개별 보고서")
set_run_font(r, 12, False, "D8E8F4")
p = cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
r = p.add_run("정량 분석 × 시장·기술·정책·표준화 정성 분석의 통합 방법론")
set_run_font(r, 11, True, "F7DCA5")

doc.add_paragraph()
add_callout(doc, "본 장의 작성 기준", [
    "수집 원자료는 2025년 6월 6일부터 2026년 6월 6일까지의 2,427건으로 관리함.",
    "수치 분석은 2026년 6월 부분월 27건을 제외한 2025년 6월~2026년 5월의 12개월·2,400건을 기준으로 수행함.",
    "제3장 연구 규모·성장, 제4장 분야·카테고리, 제5장 연구 주체, 제6장 협업 구조의 공통 모집단과 지표 정의를 고정함.",
    "정량 결과는 시장·기술·정책·표준화 자료와 교차해석하여 단순 통계가 아닌 전략적 의미를 도출하도록 설계함.",
], LIGHT_BLUE, BLUE)

doc.add_page_break()

# 2.1
add_heading(doc, "2.1 본 장의 목적 및 전체 보고서 내 위치", 1)
add_body(doc, "본 장은 arXiv 양자 분야 논문 분석의 대상 범위, 검색식, 데이터 수집·정제 절차, 분석 단위, 지표 산식, 정성 근거의 반영 원칙을 정의하는 기준 장으로 작성함.")
add_bullet(doc, "제3장 연구 규모·성장 분석에서 월별 활동 수준과 최근 모멘텀을 동일 기준으로 비교할 수 있도록 완전월 분석 원칙을 설정함.")
add_bullet(doc, "제4장 분야·카테고리 분포 분석에서 국문명을 메인으로 사용하고 arXiv 영문 카테고리 코드·명칭을 괄호로 병기하도록 표준화함.")
add_bullet(doc, "제5장 연구 주체 분석에서 저자·제1저자·기관을 구분하고, 국가 분석은 제1저자 국가 기준과 전체 기관 국가 기준의 두 축으로 수행하도록 정의함.")
add_bullet(doc, "제6장 협업 구조 분석에서 공저·다기관·다국가 협력의 최소 판정 기준과 중복 집계 방식을 고정함.")
add_bullet(doc, "제7장 종합 인사이트에서 정량 신호와 최신 시장·기술·정책·표준화 근거를 결합하여 실행 가능한 시사점으로 전환하도록 설계함.")

add_table_caption(doc, "표 2-1. 본 장의 정의 항목과 후속 장 연계")
add_table(doc,
    ["정의 항목", "본 장의 역할", "후속 장 연계"],
    [
        ["분석 대상", "원자료와 수치 분석 DB를 구분하여 모집단을 확정함", "제3장~제6장 수치의 기준을 통일함"],
        ["검색식", "양자 기술 축과 응용·기능 축의 결합 범위를 제시함", "제4장의 분야 포괄성과 노이즈를 해석함"],
        ["데이터 구축", "메타데이터·PDF·전문·참고문헌·기관·국가 정보를 통합함", "제5장·제6장의 연구 주체와 협업 분석을 지원함"],
        ["지표 체계", "활동·집중·주체·협업·품질 지표를 정의함", "제3장~제7장의 비교·통합을 가능하게 함"],
        ["정성 근거", "시장·기술·정책·표준 자료의 위계와 활용 원칙을 설정함", "제3장~제7장의 정량 결과를 외부 환경과 연결함"],
    ], [3.2, 6.3, 7.7], font_size=8.9)

add_callout(doc, "핵심 방법론 원칙", [
    "동일한 원천 논문 테이블을 분석 목적별로 재구성하되 모집단은 변경하지 않도록 관리함.",
    "수치의 크기만 제시하지 않고 발생 배경, 기술적 의미, 정책·산업 활용성과 함께 해석함.",
    "완전한 정보가 아닌 기관·국가 추론값은 커버리지와 한계를 병기하여 과도한 순위 해석을 방지함.",
], LIGHT_GREEN, GREEN)

# 2.2
add_heading(doc, "2.2 분석 대상 데이터와 완전월 기준", 1)
add_body(doc, "분석 원자료는 코넬대학교가 제공하는 arXiv의 양자 관련 논문을 자동 수집한 데이터로 구성함. 원자료는 2026년 6월 6일까지 포함하나, 2026년 6월은 6일치만 포함된 부분월이므로 월별 비교와 성장률 산출에서는 제외함.")
add_table_caption(doc, "표 2-2. 수집 원자료와 수치 분석 DB의 구분")
add_table(doc,
    ["항목", "수집 원자료", "수치 분석 DB", "적용 이유"],
    [
        ["분석 기간", "2025.06.06~2026.06.06", "2025.06~2026.05", "동일한 월 길이를 가진 완전 12개월을 비교함"],
        ["논문 수", "2,427건", "2,400건", "2026년 6월 부분월 27건을 제외함"],
        ["월 수", "13개월처럼 보이나 마지막 월은 부분월", "12개월", "평균·성장률·최근 모멘텀의 왜곡을 방지함"],
        ["월평균", "186.7편", "200.0편", "부분월 포함으로 인한 인위적 평균 하락을 보정함"],
        ["활용 목적", "원자료 보존·품질 점검", "정량 분석·비교·통합 해석", "원자료와 분석용 DB의 역할을 분리함"],
    ], [2.8, 4.0, 4.0, 6.2], font_size=8.6)

doc.add_picture(str(period_png), width=Cm(17.1))
add_figure_caption(doc, "그림 2-1. 수집 원자료와 12개월 수치 분석 DB의 범위 조정")
add_body(doc, "부분월 제외는 단순 편집 조치가 아니라 시계열 비교의 내적 타당성을 확보하는 핵심 통제 조건으로 적용함. 2026년 6월 27건을 포함할 경우 최근 연구량이 급감한 것처럼 보이는 착시가 발생하므로, 제3장의 월별 성장률과 최근 모멘텀뿐 아니라 제4장의 월별 카테고리 변화율에도 동일한 제외 기준을 적용함.")

# 2.3
add_heading(doc, "2.3 검색식 구조와 전략적 모집단", 1)
add_body(doc, "검색식은 제목(title) 필드에서 양자 기술 키워드와 응용·기능 키워드가 동시에 등장하는 논문을 선별하도록 설계함. 제목 기반 검색은 관련성이 명확한 코어 논문을 우선 확보하는 장점이 있으나, 초록·본문에서만 관련성이 확인되는 주변 논문은 누락될 수 있으므로 완전 모집단이 아닌 전략적 분석 모집단으로 해석함.")
add_callout(doc, "적용 검색식", [
    '((((ti:QUANTUM OR ti:QUANTOMETER) OR ti:QUANTIZATION) AND (((((((((((((((ti:SENSOR OR ti:SENSING) OR ti:SENSE) OR ti:DETECT) OR ti:MEASUREMENT) OR ti:INERTIA) OR ti:TIME) OR ti:FREQUENCY) OR ti:"MAGNETIC FIELD") OR ti:"ELECTRIC FIELD") OR ti:"LIGHT BASED") OR ti:COMPUTER) OR ti:COMPUTING) OR ti:COMPUTATION) OR ti:COMPUTATIONAL) OR ti:COMMUNICATION))) AND (submittedDate:[202506060000 TO 202606062359])'
], LIGHT_GRAY, NAVY)

add_table_caption(doc, "표 2-3. 검색 축별 키워드와 해석 범위")
add_table(doc,
    ["검색 축", "대표 키워드", "선별 의미", "해석·보완 기준"],
    [
        ["양자 기술", "QUANTUM, QUANTOMETER, QUANTIZATION", "양자기술 또는 양자화 개념이 제목에 명시된 논문을 선별함", "QUANTIZATION의 신호처리·모델압축 노이즈를 후속 분류에서 점검함"],
        ["센싱·계측", "SENSOR, SENSING, SENSE, DETECT, MEASUREMENT", "양자센서·검출·정밀측정 관련 논문을 포착함", "양자물리(quant-ph), 광학(physics.optics), 응집물질 계열의 교차를 확인함"],
        ["시간·주파수·전자기장", "TIME, FREQUENCY, MAGNETIC FIELD, ELECTRIC FIELD, INERTIA", "원자시계·주파수표준·자기장·전기장·관성 측정을 포함함", "센싱 세부 기술군의 실질적 확장 축으로 해석함"],
        ["컴퓨팅·통신", "COMPUTER, COMPUTING, COMPUTATION, COMPUTATIONAL, COMMUNICATION", "양자컴퓨팅·시뮬레이션·통신 관련 논문을 포함함", "협의의 센싱 분석이 아닌 광의의 양자기술 생태계 분석으로 해석함"],
    ], [2.8, 4.2, 5.2, 5.0], font_size=8.4)

doc.add_picture(str(search_png), width=Cm(17.1))
add_figure_caption(doc, "그림 2-2. 양자 기술 축과 응용·기능 축의 결합 검색 논리")
add_bullet(doc, "정밀도 우선 방식으로 분석 코어를 구축하되, 후속 토픽모델링에서는 초록·PDF 전문을 활용하여 검색식 밖의 세부 기술어를 재탐색함.")
add_bullet(doc, "검색식에 포함된 컴퓨팅·통신 키워드로 인해 양자물리(quant-ph) 외 암호·보안(cs.CR), 기계학습(cs.LG), 광학(physics.optics), 응집물질 계열이 포함될 수 있음을 전제로 함.")
add_bullet(doc, "카테고리 명칭은 보고서에서 국문을 메인으로 표기하고 영문 코드 또는 영문명을 괄호로 병기함.")

# 2.4
add_heading(doc, "2.4 데이터 수집·정제·분석 DB 구축", 1)
add_body(doc, "데이터 구축은 검색식 설계, arXiv 메타데이터 수집, PDF 확보, 전문 텍스트 추출, 저자·기관·국가·카테고리 정비, 12개월 분석 DB 생성의 순서로 수행함. 논문 단위의 원천 테이블을 유지하면서 각 장의 목적에 맞게 월·카테고리·연구주체·협업 단위로 재구성함.")
doc.add_picture(str(pipeline_png), width=Cm(17.1))
add_figure_caption(doc, "그림 2-3. arXiv 양자 분야 논문 데이터 구축 및 분석 흐름")

add_table_caption(doc, "표 2-4. 데이터 구축 단계별 산출 정보와 후속 활용")
add_table(doc,
    ["단계", "주요 수행 내용", "주요 산출 필드", "후속 활용"],
    [
        ["1. 검색식 설계", "양자축과 응용·기능축을 AND 조건으로 결합함", "검색식·기간 조건", "분석 범위와 노이즈 수준을 결정함"],
        ["2. arXiv 수집", "제출일 기준 메타데이터를 자동 수집함", "arXiv ID, 제목, 초록, 저자, 카테고리, DOI, URL", "제3장 연구량과 제4장 분야 분석의 기초자료로 활용함"],
        ["3. PDF 확보", "논문별 PDF 파일과 경로·페이지 정보를 관리함", "PDF 경로, 파일 크기, 페이지 수", "본문 분석 가능성과 품질상태를 확인함"],
        ["4. 전문 추출", "본문·단어수·참고문헌 정보를 추출함", "pdf_full_text, word_count, reference_count", "토픽모델링·성능지표·응용분야 추출 기반으로 활용함"],
        ["5. 연구주체 정비", "저자·기관명·국가 문자열을 정제하고 표준화 후보를 생성함", "author_count, first_author, institutions, country fields", "제5장 연구주체와 제6장 협업 구조에 활용함"],
        ["6. 분석 DB", "2026년 6월을 제외하고 분석 파생변수를 생성함", "month, category set, collaboration flags", "제3장~제7장의 공통 모집단으로 적용함"],
    ], [2.4, 5.0, 5.2, 5.0], font_size=8.2)

# 2.5
add_heading(doc, "2.5 분석 DB의 구성과 데이터 커버리지", 1)
add_body(doc, "수치 분석 DB는 2,400건을 기준으로 구성하며, 메타데이터와 PDF 전문을 결합한 구조로 관리함. 제목·초록·카테고리·PDF 경로의 커버리지는 분석에 충분한 수준이며, 기관·국가 정보는 추출·추론 방식의 특성을 고려하여 커버리지와 표준화 상태를 함께 제시함.")
add_table_caption(doc, "표 2-5. 분석 DB의 주요 품질·커버리지 지표")
add_table(doc,
    ["구분", "확보 수준", "해석", "적용 기준"],
    [
        ["분석 논문", "2,400건", "2025.06~2026.05 완전 12개월 기준으로 확정함", "모든 정량 분석의 분모로 사용함"],
        ["PDF 전문 텍스트", "2,397건·99.88%", "본문 기반 텍스트마이닝 확장성이 매우 높게 확보됨", "실패 3건은 별도 재처리 대상으로 관리함"],
        ["참고문헌 수", "약 96.29%", "연구 기반·지식 흐름 분석이 가능한 수준으로 보임", "추출 실패·과소추출 여부를 샘플 검증함"],
        ["기관명", "약 94.50%", "연구기관 분석 기반을 확보하였으나 표기 변형이 혼재함", "기관 표준화 전 순위는 탐색 신호로 해석함"],
        ["DOI", "약 16.13%", "프리프린트 데이터의 특성상 낮은 비율이 자연스러움", "품질·성숙도의 절대 지표로 사용하지 않음"],
        ["제1저자 국가", "1,951건·81.3%", "연구 주도 축을 파악하는 국가 기준으로 활용함", "제1저자 소속기관의 국가를 기준으로 집계함"],
        ["전체 기관 국가", "1,998건·83.3%", "모든 저자 소속기관의 참여 범위를 반영함", "논문별 동일 국가는 1회만 집계함"],
    ], [3.0, 3.6, 6.1, 5.2], font_size=8.3)
add_callout(doc, "커버리지 해석 원칙", [
    "PDF 확보율과 전문 추출률을 구분하여 파일 수집 성공과 텍스트 분석 성공을 별도로 관리함.",
    "DOI 미보유를 낮은 연구 품질로 해석하지 않고 arXiv의 조기 공개 특성을 반영함.",
    "기관명·국가 값은 명시적 소속정보, 문자열 추출, 도메인 추론 등 생성 경로를 구분하여 신뢰등급화할 필요가 있음.",
], LIGHT_YELLOW, ORANGE)

# 2.6
add_heading(doc, "2.6 분석 단위와 핵심 지표 체계", 1)
add_body(doc, "논문 단위를 기본키로 유지하되 분석 질문에 따라 월, 카테고리, 저자, 기관, 국가, 협업 단위로 재구성함. 지표는 각 장에서 독립적으로 설명 가능하면서 제7장의 종합 분석에서 결합할 수 있도록 설계함.")
add_table_caption(doc, "표 2-6. 장별 분석 단위·핵심 지표·해석 목적")
add_table(doc,
    ["분석 축", "분석 단위", "핵심 지표", "해석 목적"],
    [
        ["연구 규모·성장", "월·반기·최근 3개월", "논문 수, 월평균, 중앙값, 증감률, 이동평균, 변동계수", "연구 활동의 안정성·성장 모멘텀·피크 구간을 판단함"],
        ["분야·카테고리", "주 카테고리·전체 카테고리·교차쌍", "논문 수, 비중, 복수 카테고리율, HHI, 교차강도", "양자물리 코어와 융합 확장축을 식별함"],
        ["연구 주체", "저자·제1저자·기관·국가", "고유 수, 상위 비중, 롱테일, 국가별 건수·점유율", "연구 저변과 허브 주체, 국가별 주도·참여 구조를 파악함"],
        ["협업 구조", "공저 규모·기관쌍·국가쌍", "공동저자율, 평균 저자 수, 다기관율, 국제협력률", "연구 생태계의 팀 연구·개방성·협력 방식을 해석함"],
        ["정성 통합", "시장·기술·정책·표준 근거", "정량 신호와 외부 환경의 부합·차이·선행성", "연구 변화의 배경과 전략적 활용성을 검증함"],
    ], [3.1, 3.7, 6.3, 4.8], font_size=8.3)

add_heading(doc, "2.6.1 연구 주체의 국가 분석 이중 기준", 2)
add_body(doc, "국가별 연구역량은 한 가지 기준만으로 해석할 경우 주도성과 참여범위가 혼재될 수 있으므로 제1저자 국가 기준과 전체 기관 국가 기준을 병렬 적용함.")
add_table_caption(doc, "표 2-7. 국가 분석의 두 가지 기준")
add_table(doc,
    ["기준", "집계 방식", "주요 의미", "유의점"],
    [
        ["제1저자 국가 기준", "논문의 제1저자 소속기관 국가를 1건으로 집계함", "연구 의제의 주도·대표 축을 파악함", "제1저자가 실제 연구 책임자와 다를 수 있으며 미확인 소속을 제외함"],
        ["전체 기관 국가 기준", "모든 저자 소속기관의 국가를 추출하고 논문별 동일 국가는 1회 집계함", "국가별 참여 범위와 국제 연결성을 파악함", "다국가 논문은 여러 국가에 각각 1건씩 반영되어 합계가 논문 수를 초과함"],
    ], [3.6, 5.3, 5.0, 4.3], font_size=8.5)
add_bullet(doc, "저자 그래프에 이름을 표기할 경우 ‘저자명 (국가)’ 형식을 적용하며, 괄호의 국가는 소속기관에서 확인된 연구 수행 국가로 표기함.")
add_bullet(doc, "국가 정보가 확인되지 않은 저자는 ‘저자명 (미확인)’으로 표시하여 국적과 연구 수행 국가가 혼동되지 않도록 관리함.")

# 2.7 qualitative integration
add_heading(doc, "2.7 정량·정성 통합 분석 방법", 1)
add_body(doc, "본 보고서는 정량 지표의 크기와 순위를 제시하는 데 그치지 않고, 시장 수요, 기술 성숙도, 국가정책, 표준화·보안 전환이 연구활동에 제공하는 외부 자극을 함께 검토함. 정성 분석은 정량 결과를 정당화하기 위한 부연이 아니라 정량 신호의 원인·지속 가능성·전략적 의미를 검증하는 절차로 적용함.")
doc.add_picture(str(framework_png), width=Cm(17.1))
add_figure_caption(doc, "그림 2-4. 정량 분석과 시장·기술·정책·표준화 근거의 통합 프레임")

add_table_caption(doc, "표 2-8. 정량 분석 축별 정성 근거의 결합 방식")
add_table(doc,
    ["정량 분석 축", "연결할 정성 근거", "통합 해석 질문", "대표 참고자료"],
    [
        ["연구 규모·성장", "시장 투자, 기업 채택, 국가 예산·조달, 장기 기술로드맵", "월 200편 내외의 연구량이 일시적 관심인지 구조적 파이프라인인지 검증함", "McKinsey QTM 2026 [R6], NQI FY2025 [R2], UK Strategy [R5]"],
        ["분야·카테고리", "오류정정, PQC, QKD, 양자센싱·측정·시험표준", "교차 카테고리가 실제 기술·표준 이슈와 연결되는지 확인함", "NIST PQC [R1], Google QEC [R12], IEC/ISO JTC 3 [R38], ITU-T [R39]"],
        ["연구 주체", "국가전략, 연구인프라, 산업생태계, 인력정책", "상위 국가·기관의 집중이 정책·인프라와 부합하는지 검토함", "EU Strategy [R18], UK Strategy [R5], 한국 양자종합계획 [R23]"],
        ["협업 구조", "국제 공동프로그램, 표준화 컨소시엄, 공급망·실증 네트워크", "다기관·국제협력이 기술 융합·표준화 구조와 연결되는지 평가함", "Quantum Flagship SRIA [R4], GSMA PQC [R9], DARPA QBI [R14]"],
    ], [3.0, 5.1, 6.2, 4.6], font_size=8.0)

add_heading(doc, "2.7.1 정성 근거의 선택·검증 원칙", 2)
doc.add_picture(str(hierarchy_png), width=Cm(17.1))
add_figure_caption(doc, "그림 2-5. 정성 근거자료의 위계와 해석 통제")
add_bullet(doc, "정부·국제표준기관 자료는 정책방향, 예산, 표준화 일정, 전환 기준의 핵심 1차 근거로 활용함.")
add_bullet(doc, "동료평가 논문과 DARPA 등 독립 검증 프로그램은 오류정정·확장성·벤치마크의 기술 성숙도 판단에 활용함.")
add_bullet(doc, "시장·산업 전망은 기관별 가정이 다르므로 단일 수치를 확정값으로 사용하지 않고 복수 전망의 공통 방향과 범위를 비교함.")
add_bullet(doc, "기업 로드맵은 목표·투자·플랫폼 전략의 신호로 활용하되 이미 달성된 성과와 구분하고 외부 검증자료를 병행함.")
add_bullet(doc, "정량 신호와 외부 자료가 불일치할 경우 정량 결과를 폐기하지 않고 검색식 범위, 시차, 정책 집행 지연, 분류체계 차이를 대안 가설로 검토함.")

# 2.8 metrics
add_heading(doc, "2.8 주요 지표의 산식과 판정 기준", 1)
add_body(doc, "지표는 재현성과 장간 일관성을 확보하기 위해 분모, 중복 허용 여부, 결측 처리 기준을 명시함. 비율 지표는 원칙적으로 2,400건을 분모로 사용하되 기관·국가 등 정보 미확인 건이 존재하는 분석은 유효값 기준과 전체 모집단 기준을 함께 제시함.")
add_table_caption(doc, "표 2-9. 핵심 정량 지표의 정의·산식·해석")
add_table(doc,
    ["지표", "산식·집계 방식", "해석", "검증 포인트"],
    [
        ["월별 논문 수", "제출월별 논문 건수 Mₜ", "연구 활동의 절대 규모를 나타냄", "부분월·중복 버전을 제외함"],
        ["전후반기 증감률", "(후반기 논문 수−전반기 논문 수)÷전반기×100", "최근 6개월의 구조적 증가·감소 방향을 비교함", "각 구간을 동일한 6개월로 구성함"],
        ["복수 카테고리 비율", "카테고리 2개 이상 논문 수÷전체 논문 수×100", "융합·교차 연구의 확장성을 나타냄", "주 카테고리와 전체 카테고리를 구분함"],
        ["HHI 집중도", "Σ(분야 또는 국가 점유율×100)²", "상위 분야·국가의 집중 정도를 비교함", "집계기준이 다른 HHI는 직접 비교하지 않음"],
        ["공동저자율", "저자 2명 이상 논문 수÷전체 논문 수×100", "팀 기반 연구의 보편성을 나타냄", "저자 파싱 오류와 컨소시엄 저자를 점검함"],
        ["다기관 연구 비율", "고유 기관 2개 이상 논문 수÷전체 논문 수×100", "기관 간 협업의 확산 정도를 나타냄", "기관명 표준화 전후 값을 비교함"],
        ["국제협력률", "기관 국가 2개 이상 논문 수÷전체 논문 수×100", "국경을 넘는 공동연구의 개방성을 나타냄", "국가 미확인 논문의 영향을 병기함"],
        ["커버리지", "비결측 유효 건수÷전체 논문 수×100", "지표의 신뢰 가능한 적용 범위를 나타냄", "결측·추론·표준화 단계를 분리함"],
    ], [3.0, 6.2, 5.2, 4.2], font_size=7.9)

# 2.9 QC limitations
add_heading(doc, "2.9 품질관리와 방법론상 제한사항", 1)
add_body(doc, "분석 결과의 신뢰도를 높이기 위해 수집·정제·집계·시각화 단계별 점검 기준을 적용함. 검색식의 장점과 한계를 동시에 명시하고, 기관·국가·PDF 추출처럼 자동화 과정에서 오류 가능성이 있는 필드는 재처리·샘플 검증·신뢰등급화를 수행하도록 설계함.")
add_table_caption(doc, "표 2-10. 방법론상 위험요인과 통제·보완 방안")
add_table(doc,
    ["위험요인", "발생 가능 영향", "현재 통제", "후속 보완"],
    [
        ["제목 기반 검색", "초록·본문에만 관련성이 있는 논문을 누락할 수 있음", "관련성 높은 코어 모집단으로 해석함", "초록·전문 확장 검색과 분류모델을 병행함"],
        ["QUANTIZATION 노이즈", "양자물리와 무관한 모델 압축·신호처리 논문이 유입될 수 있음", "카테고리·초록·본문 키워드로 점검함", "규칙 기반 제외어와 지도학습 분류를 적용함"],
        ["arXiv 버전", "동일 논문의 개정 버전이 시계열 수치에 영향을 줄 수 있음", "arXiv ID 기준 대표 레코드를 사용함", "버전별 최초 제출일과 최종 수정일을 분리함"],
        ["기관명 변형", "동일 기관이 복수 명칭으로 분리되어 순위와 협업망이 왜곡될 수 있음", "기초 문자열 정제와 국가 추론을 적용함", "기관 표준사전·ROR·GRID 등 외부 식별자를 매핑함"],
        ["국가 추론 결측", "국가 순위·국제협력률이 과소추정될 수 있음", "제1저자 국가와 전체 기관 국가의 커버리지를 병기함", "소속문자열·도메인·외부기관 DB의 다중 검증을 수행함"],
        ["시장전망 편차", "기관별 시장규모·시점이 다르게 제시될 수 있음", "정책·공식·시장 자료의 위계를 구분함", "복수 전망의 가정·범위·기준연도를 병기함"],
        ["기업 로드맵 과대해석", "목표를 달성된 기술성과로 오인할 수 있음", "동료평가·독립 벤치마킹과 구분함", "달성 시점·검증지표·외부평가를 추적함"],
    ], [3.1, 5.2, 5.0, 5.3], font_size=7.9)

add_callout(doc, "중간 점검 체크리스트", [
    "논문 수와 월별 합계가 2,400건으로 일치하는지 검증함.",
    "표·그래프의 분모, 기간, 중복 허용 여부가 본 장 정의와 일치하는지 확인함.",
    "카테고리는 국문명을 메인으로, 영문 코드·명칭을 괄호로 병기함.",
    "저자 그래프의 라벨은 ‘저자명 (국가)’ 형식으로 표시하고 국가 미확인 여부를 숨기지 않음.",
    "표 제목은 표 위쪽, 그림·이미지 제목은 아래쪽에 배치함.",
    "그래프 라벨과 수치는 굵은 글꼴을 사용하고 플롯 영역 밖으로 벗어나지 않도록 여백을 검수함.",
], LIGHT_GREEN, GREEN)

# 2.10 chapter links
add_heading(doc, "2.10 후속 장 연결 구조", 1)
add_body(doc, "본 장에서 정의한 모집단·지표·정성 근거의 활용 원칙은 이후 모든 분석 장에 공통 적용함. 각 장은 독립 분석 결과를 제시하되 제7장의 통합 시사점에서 연구 규모, 분야 구조, 연구 주체, 협업 구조를 상호 연결하도록 구성함.")
add_table_caption(doc, "표 2-11. 후속 장의 핵심 질문과 본 장의 적용 기준")
add_table(doc,
    ["후속 장", "핵심 질문", "본 장에서 고정한 기준", "정성 통합 방향"],
    [
        ["제3장 연구 규모·성장", "양자 연구량이 안정적으로 증가하고 있는가", "12개월·2,400건, 월·반기·최근 3개월 지표", "시장 투자·국가예산·기업 로드맵과 연구량을 연결함"],
        ["제4장 분야·카테고리", "코어 분야와 융합 확장축은 무엇인가", "주·전체 카테고리, 복수분류, 국문+영문 병기", "오류정정·PQC·QKD·센싱·표준화 이슈와 교차함"],
        ["제5장 연구 주체", "누가 연구를 주도하고 참여하는가", "저자·제1저자·기관, 제1저자 국가·전체 기관 국가", "국가전략·연구인프라·산업생태계와 비교함"],
        ["제6장 협업 구조", "연구는 어떤 규모와 네트워크로 수행되는가", "공저·다기관·다국가 기준과 중복 집계 방식", "국제프로그램·표준 컨소시엄·공급망 협력과 연결함"],
        ["제7장 종합 시사점", "정량 신호를 어떤 R&D·산업·정책 전략으로 전환할 것인가", "정량·정성 교차검증 및 한계·대안 가설", "우선 기술군·협력대상·모니터링 체계를 제시함"],
    ], [3.3, 5.1, 5.7, 4.6], font_size=8.0)

add_heading(doc, "2.11 본 장 소결", 1)
add_bullet(doc, "수집 원자료 2,427건 중 2026년 6월 부분월 27건을 제외하고 12개월·2,400건을 수치 분석의 공통 기준으로 확정함.")
add_bullet(doc, "제목 기반 검색식은 높은 관련성을 갖는 양자 코어 논문군을 확보하는 데 강점이 있으나 전체 양자 연구의 완전 모집단으로 과대해석하지 않도록 통제함.")
add_bullet(doc, "메타데이터와 PDF 전문을 통합하여 월·카테고리·저자·기관·국가·협업의 다층 분석이 가능한 DB를 구축함.")
add_bullet(doc, "국가 분석은 제1저자 국가 기준과 전체 기관 국가 기준을 병렬 적용하여 연구 주도성과 참여 범위를 구분함.")
add_bullet(doc, "정량 분석은 최신 시장·기술·정책·표준화 자료와 교차검증하여 연구활동의 배경, 지속 가능성, 전략적 의미를 함께 해석함.")
add_bullet(doc, "다음 제3장에서는 본 장의 기준 DB를 활용하여 월별 연구량, 전후반기 증감, 피크·저점, 최근 모멘텀을 분석함.")

# selected references
add_heading(doc, "2.12 본 장에서 활용하는 주요 정성 근거자료", 1)
add_body(doc, "아래 자료는 방법론의 정성 통합 프레임을 설계하고 후속 장의 정량 결과를 교차해석하기 위한 대표 근거자료로 관리함. 전체 참고자료는 최종 통합 보고서의 부록에서 시장·기술·정책·표준화 유형별로 정리함.")
refs = [
    ("[R1]", "NIST CSRC, Post-Quantum Cryptography FIPS Approved", "PQC 표준과 전환 기준", "https://csrc.nist.gov/news/2024/postquantum-cryptography-fips-approved"),
    ("[R2]", "National Quantum Initiative, Annual Report FY2025", "미국 국가양자 연구·예산·기관활동", "https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf"),
    ("[R4]", "Quantum Flagship, Strategic Research and Industry Agenda 2030", "EU 기술·산업 로드맵", "https://qt.eu/about-quantum-flagship/strategic-research-and-industry-agenda-2030"),
    ("[R5]", "UK Government, National Quantum Strategy", "영국 연구·산업화·인력 전략", "https://www.gov.uk/government/publications/national-quantum-strategy"),
    ("[R6]", "McKinsey, Quantum Technology Monitor 2026", "시장·기업 채택·투자 신호", "https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/mckinsey-quantum-technology-monitor-2026-a-commercial-tipping-point"),
    ("[R9]", "GSMA, Post-Quantum Cryptography Documents", "이동통신망 PQC 전환", "https://www.gsma.com/solutions-and-impact/technologies/security/post-quantum-cryptography-documents/"),
    ("[R12]", "Nature, Quantum Error Correction Below the Surface Code Threshold", "오류정정 기술 성숙도 검증", "https://www.nature.com/articles/s41586-024-08449-y"),
    ("[R14]", "DARPA, Quantum Benchmarking Initiative", "유용한 양자컴퓨터의 독립 벤치마킹", "https://www.darpa.mil/research/programs/quantum-benchmarking-initiative"),
    ("[R18]", "European Commission, Quantum Europe Strategy", "EU 최신 양자 전략", "https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy"),
    ("[R38]", "IEC/ISO JTC 3, Strategic Business Plan", "양자 전 분야 국제표준화 구조", "https://assets.iec.ch/further_informations/49854/SBP_JTC%203_2026.05.pdf"),
    ("[R39]", "ITU-T Y.Sup98, Technical Considerations Towards Quantum Networks", "양자 네트워크 구조·표준화", "https://www.itu.int/rec/dologin_pub.asp?id=T-REC-Y.Sup98-202511-I%21%21PDF-E&lang=s&type=items"),
]
for rid, title, use, url in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.35)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    set_paragraph_format(p, after=3, line=1.18)
    r = p.add_run(f"{rid} ")
    set_run_font(r, 8.9, True, NAVY)
    add_hyperlink(p, title, url)
    r = p.add_run(f" — {use}에 활용함.")
    set_run_font(r, 8.9, False, DARK)

# Document metadata and core properties
cp = doc.core_properties
cp.title = "arXiv 양자 분야 제2장 분석 대상 및 방법론"
cp.subject = "2025년 6월~2026년 5월 12개월 기준, 정량·정성 통합 방법론"
cp.author = "OpenAI"
cp.keywords = "arXiv, 양자, 논문분석, 방법론, 시장, 기술, 정책, 표준화"

# Set update fields on open
settings = doc.settings._element
update_fields = OxmlElement("w:updateFields")
update_fields.set(qn("w:val"), "true")
settings.append(update_fields)

# Final save
doc.save(OUTPUT)
print(f"Created {OUTPUT}")

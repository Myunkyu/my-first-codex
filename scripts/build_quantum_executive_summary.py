from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path('artifacts')
ASSET_DIR = OUT_DIR / 'executive_summary_assets'
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / 'arXiv_양자분야_제1장_Executive_Summary_고품질_재작성_20260704.docx'

# -----------------------------------------------------------------------------
# Consolidated evidence from Chapters 2-7
# -----------------------------------------------------------------------------
TOTAL = 2400
PERIOD = '2025년 6월~2026년 5월'
MONTHLY = np.array([182,187,183,219,229,171,227,186,165,211,220,220])
MONTH_LABELS = ['25.06','25.07','25.08','25.09','25.10','25.11','25.12','26.01','26.02','26.03','26.04','26.05']
MONTHLY_AVG = 200.0
RECENT3_AVG = 217.0
HALF_GROWTH = 5.0
CV = 11.1

QUANT_PH_N = 1680
QUANT_PH_SHARE = 70.00
ALL_CATEGORY_N = 117
MULTI_CATEGORY_N = 1138
MULTI_CATEGORY_SHARE = 47.42
CATEGORY_HHI = 0.493

UNIQUE_AUTHORS = 10514
UNIQUE_FIRST_AUTHORS = 2248
VALID_INST = 1995
COUNTRIES = 35
FIRST_COUNTRY_US = 453
INST_COUNTRY_US = 524
KOREA_FIRST = 44
KOREA_INST = 48

COAUTHOR_N = 2113
COAUTHOR_SHARE = 88.0
AVG_AUTHORS = 5.41
MEDIAN_AUTHORS = 4
MULTI_INST_N = 710
MULTI_INST_SHARE = 29.6
INTL_N = 241
INTL_SHARE = 10.0

TECH = [
    ('양자컴퓨팅·SW/AI',4.8,5.0,4.6,'선도·상용화 가속'),
    ('양자통신·보안',3.8,4.8,4.2,'선도·상용화 가속'),
    ('양자센싱·계측',4.1,4.4,4.0,'선도·상용화 가속'),
    ('양자소자·소재',4.3,4.1,3.8,'핵심기반·공급망'),
    ('양자이론·시뮬레이션',4.4,2.8,3.2,'장기 원천기반'),
    ('양자화학·산업응용',3.3,4.0,3.5,'선택적 응용발굴'),
]
COUNTRY = [
    ('미국',453,524,133),('중국',277,296,66),('독일',161,199,79),
    ('영국',113,142,73),('일본',112,124,20),('한국',44,48,6),
]

NAVY='17365D'; BLUE='2F75B5'; LIGHT_BLUE='D9EAF7'; GREEN='548235'; LIGHT_GREEN='E2F0D9'; ORANGE='ED7D31'; LIGHT_ORANGE='FCE4D6'; PURPLE='7030A0'; LIGHT_PURPLE='E4DFEC'; GRAY='6B7280'; LIGHT_GRAY='F2F4F7'; DARK='243447'; WHITE='FFFFFF'; BORDER='CAD4DF'; YELLOW='FFD966'; RED='C00000'

def rgb(h): return RGBColor.from_string(h)

font_candidates = [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc',
    '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
]
font_path = next((p for p in font_candidates if Path(p).exists()), None)
if font_path:
    fp = font_manager.FontProperties(fname=font_path)
    plt.rcParams['font.family'] = fp.get_name()
else:
    plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.weight'] = 'bold'

# -----------------------------------------------------------------------------
# Charts - matplotlib default colors only
# -----------------------------------------------------------------------------
def finish(fig, path):
    fig.tight_layout(pad=1.0)
    fig.savefig(path, dpi=220, bbox_inches='tight')
    plt.close(fig)


def fig_dashboard(path):
    fig, ax = plt.subplots(figsize=(12, 6.2))
    ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    cards = [
        (.03,.56,.22,.32,colors[0],'연구활동',[f'{TOTAL:,}건',f'월평균 {MONTHLY_AVG:.1f}편',f'최근 3개월 {RECENT3_AVG:.1f}편']),
        (.27,.56,.22,.32,colors[1],'분야·융합',[f'quant-ph {QUANT_PH_SHARE:.2f}%',f'전체 카테고리 {ALL_CATEGORY_N}개',f'복수분류 {MULTI_CATEGORY_SHARE:.2f}%']),
        (.51,.56,.22,.32,colors[2],'연구 주체',[f'저자 {UNIQUE_AUTHORS:,}명',f'기관 {VALID_INST:,}개',f'{COUNTRIES}개국']),
        (.75,.56,.22,.32,colors[3],'협업 구조',[f'공동저자 {COAUTHOR_SHARE:.1f}%',f'다기관 {MULTI_INST_SHARE:.1f}%',f'국제협력 {INTL_SHARE:.1f}%']),
    ]
    for x,y,w,h,color,title,lines in cards:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor='white',edgecolor=color,linewidth=2.2))
        ax.text(x+w/2,y+h-.06,title,ha='center',fontsize=12.5,fontweight='bold',color=color)
        ax.text(x+w/2,y+.13,'\n'.join(lines),ha='center',va='center',fontsize=11,fontweight='bold',linespacing=1.5)
    ax.add_patch(FancyBboxPatch((.08,.12),.84,.27,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor='black',linewidth=2.2))
    ax.text(.50,.31,'종합 판단',ha='center',fontsize=11,fontweight='bold')
    ax.text(.50,.235,'안정적 고활동과 강한 과학 코어 위에서 융합·실증·표준·상용화 경쟁이 가속됨',ha='center',fontsize=14,fontweight='bold')
    ax.text(.50,.16,'한국은 전면 추격보다 기술 특화·공동 인프라·국제 허브·표준 선점의 결합전략이 적합함',ha='center',fontsize=10.5,fontweight='bold')
    for x in [.14,.38,.62,.86]:
        ax.add_patch(FancyArrowPatch((x,.55),(.50,.40),arrowstyle='-|>',mutation_scale=14,linewidth=1.2,alpha=.7))
    ax.text(.50,.96,'양자 분야 분석 핵심 대시보드',ha='center',fontsize=17,fontweight='bold')
    finish(fig,path)


def fig_growth(path):
    fig, ax = plt.subplots(figsize=(11.5,4.9))
    x = np.arange(12)
    ax.bar(x, MONTHLY, alpha=.65, label='월별 논문 수')
    ma = np.convolve(MONTHLY, np.ones(3)/3, mode='valid')
    ax.plot(np.arange(2,12), ma, marker='o', linewidth=2.2, label='3개월 이동평균')
    ax.axhline(MONTHLY_AVG, linestyle='--', linewidth=1.3, label=f'월평균 {MONTHLY_AVG:.0f}편')
    ax.set_xticks(x); ax.set_xticklabels(MONTH_LABELS,fontweight='bold')
    ax.set_ylabel('논문 수(편)',fontweight='bold')
    ax.set_ylim(0,260); ax.grid(axis='y',alpha=.25)
    ax.set_title('월별 연구활동과 최근 모멘텀',fontsize=16,fontweight='bold',pad=16)
    ax.legend(frameon=False,ncol=3,loc='lower left')
    for i,v in enumerate(MONTHLY):
        if i in [0,4,6,8,10,11]: ax.text(i,v+5,str(v),ha='center',fontsize=9,fontweight='bold')
    finish(fig,path)


def fig_portfolio(path):
    fig, ax = plt.subplots(figsize=(10.8,6.7))
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    for i,(name,research,readiness,growth,priority) in enumerate(TECH):
        size = 300 + growth*90
        ax.scatter(research,readiness,s=size,alpha=.72,edgecolors='black',linewidths=1,color=colors[i])
        dx,dy=(.05,.07)
        if name=='양자컴퓨팅·SW/AI': dx=-.8
        if name in ['양자통신·보안','양자소자·소재']: dy=-.18
        ax.text(research+dx,readiness+dy,f'{name}\n{priority}',fontsize=9.4,fontweight='bold')
    ax.axvline(4.0,linestyle='--',linewidth=1.1); ax.axhline(4.0,linestyle='--',linewidth=1.1)
    ax.set_xlim(2.8,5.25); ax.set_ylim(2.4,5.25)
    ax.set_xlabel('연구기반·학술역량(5점)',fontweight='bold')
    ax.set_ylabel('시장·정책·표준 준비도(5점)',fontweight='bold')
    ax.set_title('6대 기술축 전략 포트폴리오',fontsize=16,fontweight='bold',pad=17)
    ax.grid(alpha=.25)
    ax.text(4.65,5.05,'선도·상용화',ha='center',fontsize=10.5,fontweight='bold')
    ax.text(3.35,5.05,'표준·응용 선점',ha='center',fontsize=10.5,fontweight='bold')
    ax.text(4.65,2.55,'장기 기반투자',ha='center',fontsize=10.5,fontweight='bold')
    ax.text(3.35,2.55,'선택적 검증',ha='center',fontsize=10.5,fontweight='bold')
    ax.text(2.88,2.68,'※ 원 크기: 최근 성장 모멘텀\n※ 점수: 정량결과+외부자료 전문가 통합평가',fontsize=8.7,fontweight='bold')
    finish(fig,path)


def fig_country(path):
    fig, ax = plt.subplots(figsize=(9.2,6.5))
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    ax.plot([0,560],[0,560],linestyle='--',linewidth=1.2,label='주도=참여 기준선')
    for i,(name,lead,participation,hub) in enumerate(COUNTRY):
        ax.scatter(lead,participation,s=180+hub*5,alpha=.72,edgecolors='black',linewidths=1,color=colors[i])
        offsets={'미국':(-48,8),'중국':(8,5),'독일':(8,5),'영국':(8,5),'일본':(8,-18),'한국':(8,-18)}
        dx,dy=offsets[name]
        ax.annotate(f'{name}\n{lead}/{participation}',(lead,participation),xytext=(dx,dy),textcoords='offset points',fontsize=9.5,fontweight='bold')
    ax.set_xlim(0,570); ax.set_ylim(0,570); ax.grid(alpha=.25)
    ax.set_xlabel('제1저자 국가 논문 수: 연구 주도성',fontweight='bold')
    ax.set_ylabel('전체 기관 국가 연결 수: 참여 외연',fontweight='bold')
    ax.set_title('주요국 연구 주도–참여 포지션',fontsize=16,fontweight='bold',pad=16)
    ax.text(180,75,'한국: 자체 주도 기반은 존재하나\n규모와 네트워크 외연 확대 필요',ha='center',fontsize=9.5,fontweight='bold')
    ax.legend(frameon=False,loc='lower right')
    finish(fig,path)


def fig_strategy(path):
    fig, ax = plt.subplots(figsize=(11.5,6.0))
    ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    colors=plt.rcParams['axes.prop_cycle'].by_key()['color']
    center=(.50,.49)
    ax.add_patch(FancyBboxPatch((.36,.36),.28,.26,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor='black',linewidth=2.3))
    ax.text(.50,.54,'한국형 양자전략',ha='center',fontsize=15,fontweight='bold')
    ax.text(.50,.46,'특화 × 인프라 × 국제협력\n× 표준 × 데이터',ha='center',fontsize=10.5,fontweight='bold',linespacing=1.4)
    boxes=[
        (.04,.65,.25,.22,colors[0],'1. 기술 포트폴리오','컴퓨팅·SW/AI·통신·센싱 우선'),
        (.71,.65,.25,.22,colors[1],'2. 공동 인프라','극저온·광자·파운드리·시험평가'),
        (.04,.12,.25,.22,colors[2],'3. 국제 허브','미·독·영·일 기관·브리지 연구자'),
        (.71,.12,.25,.22,colors[3],'4. 표준·시장 전환','PQC·양자망·실증·조달·기업수요'),
        (.365,.73,.27,.17,colors[4],'5. 데이터 거버넌스','논문–특허–과제–표준–시장 연계'),
        (.365,.08,.27,.17,colors[5],'6. 인력·조직','융합인력·산업파견·PMO·KPI'),
    ]
    for x,y,w,h,color,title,sub in boxes:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor='white',edgecolor=color,linewidth=1.9))
        ax.text(x+w/2,y+h*.66,title,ha='center',fontsize=10.5,fontweight='bold',color=color)
        ax.text(x+w/2,y+h*.27,sub,ha='center',fontsize=8.7,fontweight='bold')
        ax.add_patch(FancyArrowPatch((x+w/2,y+h/2),center,arrowstyle='-|>',mutation_scale=12,linewidth=1.0,alpha=.6))
    ax.text(.50,.96,'6대 전략축 통합 실행모델',ha='center',fontsize=16,fontweight='bold')
    finish(fig,path)


def fig_roadmap(path):
    fig, ax = plt.subplots(figsize=(12,5.9))
    ax.axis('off'); ax.set_xlim(0,36); ax.set_ylim(0,5.9)
    colors=plt.rcParams['axes.prop_cycle'].by_key()['color']
    phases=[(0,6,'0~6개월\n정비·선별'),(6,18,'6~18개월\n검증·연계'),(18,36,'18~36개월\n확대·상용화')]
    for i,(s,e,title) in enumerate(phases):
        ax.add_patch(FancyBboxPatch((s+.2,5.05),e-s-.4,.62,boxstyle='round,pad=.02,rounding_size=.08',facecolor=colors[i],alpha=.65,edgecolor='black'))
        ax.text((s+e)/2,5.36,title,ha='center',va='center',fontsize=10.8,fontweight='bold')
    rows=[
        ('데이터·분류',[(0,6,'DB·기관·기술분류 정비'),(6,18,'논문–특허–표준 연계'),(18,36,'상시 조기경보 운영')]),
        ('기술·실증',[(0,6,'6대 기술축·후보 확정'),(6,18,'벤치마크·산업 PoC'),(18,36,'제품·조달·사업화 확대')]),
        ('인프라·협력',[(0,6,'시설·허브·파트너 맵'),(6,18,'공동시설·국제과제'),(18,36,'표준·공급망 연대')]),
        ('인력·조직',[(0,6,'핵심인력·PMO 구성'),(6,18,'공동박사·산업파견'),(18,36,'클러스터·기업 생태계')]),
        ('성과관리',[(0,6,'KPI 기준선 설정'),(6,18,'분기점검·중간평가'),(18,36,'투자 확대·전환·중단')]),
    ]
    for r,(label,items) in enumerate(rows):
        y=4.4-r*.82
        ax.text(-.2,y+.21,label,ha='right',va='center',fontsize=9.6,fontweight='bold')
        for j,(s,e,txt) in enumerate(items):
            ax.add_patch(FancyBboxPatch((s+.35,y),e-s-.7,.45,boxstyle='round,pad=.015,rounding_size=.05',facecolor='white',edgecolor=colors[j],linewidth=1.4))
            ax.text((s+e)/2,y+.225,txt,ha='center',va='center',fontsize=8.7,fontweight='bold')
    ax.text(18,5.85,'36개월 실행 로드맵',ha='center',fontsize=16,fontweight='bold')
    ax.text(18,.10,'각 단계에서 기술성·시장성·협업성·데이터 신뢰도를 재평가하여 투자 우선순위를 조정함',ha='center',fontsize=9.2,fontweight='bold')
    finish(fig,path)

FIGS = {
    'dashboard': ASSET_DIR/'fig_es_1_dashboard.png',
    'growth': ASSET_DIR/'fig_es_2_growth.png',
    'portfolio': ASSET_DIR/'fig_es_3_portfolio.png',
    'country': ASSET_DIR/'fig_es_4_country.png',
    'strategy': ASSET_DIR/'fig_es_5_strategy.png',
    'roadmap': ASSET_DIR/'fig_es_6_roadmap.png',
}
fig_dashboard(FIGS['dashboard'])
fig_growth(FIGS['growth'])
fig_portfolio(FIGS['portfolio'])
fig_country(FIGS['country'])
fig_strategy(FIGS['strategy'])
fig_roadmap(FIGS['roadmap'])

# -----------------------------------------------------------------------------
# Word helpers
# -----------------------------------------------------------------------------
def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_border(cell,color=BORDER,size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ['top','left','bottom','right']:
        el=borders.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)

def set_cell_margins(cell,top=80,start=90,bottom=80,end=90):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_run(run,size=10,bold=False,color=DARK):
    run.font.name='맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=rgb(color)

def pformat(p,before=0,after=4,line=1.25,keep=False):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line; pf.keep_with_next=keep

def add_page_number(p):
    run=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text='PAGE'; e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); run._r.extend([b,i,e])

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); pformat(p,before=12 if level==1 else 8,after=6,line=1.0,keep=True); r=p.add_run(text); set_run(r,15 if level==1 else 12,True,NAVY)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'3'); bottom.set(qn('w:color'),BLUE); pBdr.append(bottom); pPr.append(pBdr)
    return p

def add_bullet(doc,text,major=False):
    p=doc.add_paragraph(); pformat(p,after=3,line=1.28); p.paragraph_format.left_indent=Cm(.48); p.paragraph_format.first_line_indent=Cm(-.34); r=p.add_run('□ ' if major else '· '); set_run(r,10.2,True,BLUE if major else GRAY); r=p.add_run(text); set_run(r,10.2,major,DARK); return p

def add_caption(doc,text,figure=False):
    p=doc.add_paragraph(); pformat(p,before=3,after=6,line=1.0,keep=not figure); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); set_run(r,9.2,True,GRAY if figure else NAVY)

def add_table(doc,headers,rows,widths,font=8.0):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,NAVY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,after=0,line=1.05); r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cell=cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,WHITE if ri%2==0 else LIGHT_GRAY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pformat(p,after=0,line=1.08); r=p.add_run(str(v)); set_run(r,font,i==0,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0); return table

def add_callout(doc,title,lines,fill=LIGHT_BLUE,accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,accent,'8'); set_cell_margins(cell,130,150,130,150); p=cell.paragraphs[0]; pformat(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=cell.add_paragraph(); pformat(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.first_line_indent=Cm(-.25); r=p.add_run('• '); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_hyperlink(p,text,url):
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rPr.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)

# -----------------------------------------------------------------------------
# Build document
# -----------------------------------------------------------------------------
doc=Document(); sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35); sec.header_distance=Cm(.55); sec.footer_distance=Cm(.55)
style=doc.styles['Normal']; style.font.name='맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); style.font.size=Pt(10)
header=sec.header.paragraphs[0]; r=header.add_run('arXiv 양자 분야 논문 분석 | 제1장 Executive Summary'); set_run(r,8.4,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('분석 기준: 2025.06~2026.05, 12개월·2,400건  |  '); set_run(r,8.2,False,GRAY); add_page_number(footer)

cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; cell=cover.cell(0,0); set_cell_shading(cell,NAVY); set_cell_margins(cell,760,300,760,300); set_cell_border(cell,NAVY,'0'); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CHAPTER 1'); set_run(r,14,True,'9DC3E6'); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Executive Summary'); set_run(r,29,True,WHITE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('arXiv 양자 분야 연구동향 종합분석'); set_run(r,14,True,LIGHT_BLUE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,before=18); r=p.add_run(f'{PERIOD} | 12개월·{TOTAL:,}건'); set_run(r,11.5,True,YELLOW)

doc.add_paragraph(); kpis=[('2,400건','분석 논문'),('217편','최근 3개월 평균'),('47.42%','복수 카테고리'),('10,514명','고유 저자'),('10.0%','국제협력')]; t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(val,label) in enumerate(kpis):
    cell=t.cell(0,i); fill=[LIGHT_BLUE,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]; accent=[BLUE,GREEN,ORANGE,PURPLE,GRAY][i]; set_cell_shading(cell,fill); set_cell_border(cell,accent,'7'); set_cell_margins(cell,100,70,100,70); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,14.5,True,NAVY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_run(r,8.5,True,GRAY)

add_callout(doc,'최종 결론',[
    '양자 연구는 높은 활동성과 강한 양자물리 코어를 유지하면서 AI·광학·보안·소재·응용으로 확장되는 안정 성장형 융합 생태계로 판단함.',
    '연구 중심 경쟁은 오류정정·독립 벤치마크·PQC·양자네트워크 표준·산업 실증·공급망을 중심으로 검증·상용화 경쟁으로 이동하고 있음.',
    '한국은 전 기술영역의 양적 추격보다 컴퓨팅·SW/AI, 통신·보안, 센싱·계측의 선택적 선도와 소자·소재 인프라, 국제 허브·표준 연계를 결합하는 전략이 적합함.',
],LIGHT_BLUE,BLUE)
doc.add_page_break()

# ES 1
add_heading(doc,'1.1 분석 개요 및 해석 범위')
add_bullet(doc,f'분석대상은 arXiv에서 수집한 양자 관련 논문이며, 수치 비교는 데이터가 완전한 {PERIOD}의 12개월·{TOTAL:,}건을 기준으로 수행함.',True)
add_bullet(doc,'2026년 6월은 수집 종료일에 따른 부분월 27건으로 확인되어 연구량·성장률·분야·주체·협업의 비교분석에서 제외함.')
add_bullet(doc,'검색식은 제목의 QUANTUM·QUANTOMETER·QUANTIZATION과 센싱·검출·계측·컴퓨팅·통신 관련 용어를 결합하여 기술연관성이 높은 논문을 수집함.')
add_bullet(doc,'arXiv의 신속성을 활용해 최신 연구활동을 포착하되, 동료평가·인용·특허·시장성과를 직접 의미하지 않는다는 한계를 전제로 해석함.')
add_bullet(doc,'Executive Summary는 제2~7장의 정량결과와 최신 시장·기술·정책·표준화 자료를 결합하여 의사결정에 필요한 핵심 판단과 실행방향을 요약함.')

add_caption(doc,'표 1-1. 분석 범위와 핵심 품질관리 기준')
add_table(doc,['항목','적용 기준','품질관리·해석 원칙'],[
['분석기간',PERIOD,'2026년 6월 부분월을 제외하여 월별 비교 왜곡을 방지함'],
['분석대상',f'arXiv 논문 {TOTAL:,}건','제목 검색식 기반 선행신호로 활용하고 전문 정합성 검증을 병행함'],
['분야분석','주 카테고리와 전체 카테고리','학술분류를 국문 중심으로 제시하고 실무 기술축에 N:M 재매핑함'],
['주체분석','저자·제1저자·기관·국가','국가는 제1저자 국가와 전체 기관 국가를 병렬 적용함'],
['협업분석','공저·다기관·국제협력','기관명 표준화와 국가정보 커버리지 한계를 병기함'],
['정성분석','시장·정책·표준·기업 로드맵','분석기간 이후 자료는 향후 환경신호로 구분함'],
],[3.4,6.0,9.4],8.1)

# ES 2
add_heading(doc,'1.2 핵심 분석결과')
doc.add_picture(str(FIGS['dashboard']),width=Cm(18.1)); add_caption(doc,'그림 1-1. 양자 분야 분석 핵심 대시보드',True)
add_caption(doc,'표 1-2. 분야별 핵심 결과와 종합 해석')
add_table(doc,['분석 영역','핵심 수치','핵심 해석'],[
['연구 규모·성장',f'{TOTAL:,}건·월평균 {MONTHLY_AVG:.1f}편·최근 3개월 {RECENT3_AVG:.1f}편','급격한 일회성 팽창보다 높은 기반에서 조정과 재상승이 반복되는 안정 성장형 구조로 보임'],
['분야·카테고리',f'양자물리(quant-ph) {QUANT_PH_N:,}건·{QUANT_PH_SHARE:.2f}%, 복수 카테고리 {MULTI_CATEGORY_SHARE:.2f}%','원천과학 코어가 강하면서 AI·광학·보안·응집물질·응용물리로 융합이 확장됨'],
['연구 주체',f'저자 {UNIQUE_AUTHORS:,}명·제1저자 {UNIQUE_FIRST_AUTHORS:,}명·기관 {VALID_INST:,}개·{COUNTRIES}개국','개인·기관은 롱테일형으로 분산되나 연구량과 인프라는 주요국에 집중됨'],
['협업 구조',f'공동저자 {COAUTHOR_SHARE:.1f}%·다기관 {MULTI_INST_SHARE:.1f}%·국제협력 {INTL_SHARE:.1f}%','팀 연구는 보편적이나 기관·국가 경계를 넘는 협업은 선택적으로 형성됨'],
['전략 전환','오류정정·검증·표준·실증·공급망 정책 확대','논문 중심 경쟁에서 유용성·경제성·상호운용·배치가능성 경쟁으로 이동함'],
],[3.4,6.3,9.1],8.0)

# ES 3
add_heading(doc,'1.3 연구 규모·성장: 고활동 기반과 최근 재상승')
doc.add_picture(str(FIGS['growth']),width=Cm(18.0)); add_caption(doc,'그림 1-2. 월별 연구활동과 최근 모멘텀',True)
add_bullet(doc,f'월별 논문 수는 165~229편 범위이며 월평균 {MONTHLY_AVG:.1f}편, 변동계수 {CV:.1f}%로 나타나 높은 연구활동이 비교적 안정적으로 유지됨.',True)
add_bullet(doc,f'최근 3개월 평균은 {RECENT3_AVG:.1f}편으로 전체 월평균보다 8.5% 높고 후반 6개월 논문량은 전반기보다 약 {HALF_GROWTH:.1f}% 증가함.')
add_bullet(doc,'2025년 11월과 2026년 2월의 조정 이후 2026년 3~5월 211~220편 수준으로 회복되어 연구 파이프라인의 복원력이 확인됨.')
add_bullet(doc,'성장 해석은 12개월이라는 짧은 기간을 고려하여 급성장으로 단정하지 않고 지속 모니터링이 필요한 안정 확장 신호로 판단함.')

# ES 4
add_heading(doc,'1.4 분야·융합: 강한 양자 코어와 응용 확장의 공존')
add_bullet(doc,f'주 카테고리 기준 양자물리(quant-ph)가 {QUANT_PH_N:,}건·{QUANT_PH_SHARE:.2f}%로 분석대상의 과학적 중심축을 형성함.',True)
add_bullet(doc,f'전체 카테고리는 {ALL_CATEGORY_N}개, 복수 카테고리 논문은 {MULTI_CATEGORY_N:,}건·{MULTI_CATEGORY_SHARE:.2f}%로 나타나 코어 집중 내부에서 다학제 융합이 보편화됨.')
add_bullet(doc,'기계학습(cs.LG), 암호/보안(cs.CR), 광학(physics.optics), 메조스코픽/홀 효과(cond-mat.mes-hall), 통계역학(cond-mat.stat-mech)이 주요 연결축으로 확인됨.')
add_bullet(doc,'학술 카테고리를 산업전략에 활용하려면 양자컴퓨팅·SW/AI, 양자통신·보안, 양자센싱·계측, 양자소자·소재, 양자이론·시뮬레이션, 양자화학·산업응용의 6대 기술축으로 재매핑하는 것이 적합함.')

# ES 5
add_heading(doc,'1.5 연구 주체·협업: 넓은 저변과 국가·허브 집중의 이중 구조')
add_bullet(doc,f'고유 저자 {UNIQUE_AUTHORS:,}명, 제1저자 {UNIQUE_FIRST_AUTHORS:,}명, 유효 기관 {VALID_INST:,}개가 확인되어 개인·기관 단위 연구생태계는 광범위하게 분산됨.',True)
add_bullet(doc,f'공동저자 연구는 {COAUTHOR_N:,}건·{COAUTHOR_SHARE:.1f}%이며 평균 저자 {AVG_AUTHORS:.2f}명, 중앙값 {MEDIAN_AUTHORS}명으로 소·중규모 팀이 양자 연구의 기본 단위로 나타남.')
add_bullet(doc,f'다기관 연구는 {MULTI_INST_N:,}건·{MULTI_INST_SHARE:.1f}%, 국제협력은 {INTL_N:,}건·{INTL_SHARE:.1f}%로 나타나 연구자 협업에 비해 기관·국가 간 협력은 제한적임.')
add_bullet(doc,f'미국은 제1저자 국가 {FIRST_COUNTRY_US}건과 기관 국가 {INST_COUNTRY_US}건으로 연구 주도성과 참여 외연 모두 1위이며 독일·영국은 연구량 대비 국제 허브성이 강함.')
add_bullet(doc,f'한국은 제1저자 국가 {KOREA_FIRST}건과 기관 국가 {KOREA_INST}건으로 12위 수준의 기반을 보유하나 규모와 국제 네트워크 외연의 확대가 요구됨.')

doc.add_picture(str(FIGS['country']),width=Cm(15.7)); add_caption(doc,'그림 1-3. 주요국 연구 주도–참여 포지션',True)

# ES 6
add_heading(doc,'1.6 시장·기술·정책 환경: 연구에서 검증·표준·배치로 이동')
add_bullet(doc,'미국 NQI는 FY2025 양자정보과학 R&D 예산으로 9.98억 달러를 요청하고 과학·인력·산업·인프라·안보·국제협력을 정책축으로 지속 투자함 [R1].',True)
add_bullet(doc,'EU Quantum Europe Strategy는 연구혁신, 양자인프라, 생태계·공급망·산업화, 우주·듀얼유스, 인력의 통합적 강화 방향을 제시함 [R2].')
add_bullet(doc,'2026년 미국 행정명령은 양자컴퓨팅·센싱·네트워킹의 배치·상용화, 성능평가센터, 공급망·제조, 인력, 국제파트너십을 구체적 실행과제로 제시함 [R3].')
add_bullet(doc,'NIST의 FIPS 203·204·205 승인은 양자내성암호가 후보기술에서 실제 전환표준 단계로 진입했음을 의미함 [R4].')
add_bullet(doc,'DARPA QBI는 2033년까지 계산가치가 비용을 초과하는 산업적으로 유용한 양자컴퓨터의 가능성을 독립 검증하는 체계를 운영함 [R5].')
add_bullet(doc,'한국의 제1차 양자종합계획은 2035년 전문인력 1만 명, 양자기업 2천 개, 글로벌 표준 선도 3위권과 5대 양자클러스터 분야를 목표로 제시함 [R6].')

add_caption(doc,'표 1-3. 외부환경 변화와 전략적 의미')
add_table(doc,['외부환경','주요 변화','전략적 의미'],[
['공공투자','미국·EU·영국·한국의 장기전략·인프라·인력 투자','단기 성과만으로 원천·공통기반 투자를 축소하기 어려움'],
['상용화','정부 배치·기업 채택·사용사례 발굴 강화','산업문제와 기존 HPC·AI 워크플로 연계가 중요함'],
['검증','DARPA QBI·성능평가센터·기업 로드맵 구체화','논문 수보다 유용성·비용·재현성·성능검증을 중시함'],
['표준화','NIST PQC·ITU 양자네트워크 표준화 진전','통신·보안 분야는 조기표준 참여와 전환계획이 경쟁력임'],
['공급망','소자·광자·극저온·제어·파운드리 강조','핵심부품과 공동인프라가 기술주권과 사업화 속도를 좌우함'],
['인력·생태계','융합인력·기업·클러스터 목표 확대','물리·공학·SW·산업인력의 순환구조가 필요함'],
],[3.3,7.0,8.5],7.9)

# ES 7
add_heading(doc,'1.7 기술 포트폴리오와 투자 우선순위')
doc.add_picture(str(FIGS['portfolio']),width=Cm(17.5)); add_caption(doc,'그림 1-4. 6대 기술축 전략 포트폴리오',True)
add_bullet(doc,'양자컴퓨팅·SW/AI는 연구기반과 정책·기업 로드맵이 가장 구체적이므로 오류정정·컴파일러·하이브리드·산업문제 실증을 묶어 최우선 추진함.',True)
add_bullet(doc,'양자통신·보안은 PQC 전환과 QKD·양자네트워크가 서로 다른 시간축을 가지므로 즉시 전환형 보안과 중장기 네트워크 실증을 병렬 추진함.')
add_bullet(doc,'양자센싱·계측은 의료·항법·반도체·국방 등 조기 사용사례가 존재하므로 현장성능·시험인증·공공조달을 중심으로 상용화를 가속함.')
add_bullet(doc,'양자소자·소재는 모든 기술축의 성능·자립성을 좌우하는 핵심 기반으로 공동파운드리·극저온·광자·제어·패키징 공급망을 구축함.')
add_bullet(doc,'양자이론·시뮬레이션과 양자화학·산업응용은 장기 원천성과와 고가치 산업문제 발굴을 병행하되 단계적 PoC와 명확한 중간지표를 적용함.')

add_caption(doc,'표 1-4. 기술축별 우선순위와 실행방향')
add_table(doc,['우선군','기술축','핵심 실행방향','주요 KPI'],[
['1순위','양자컴퓨팅·SW/AI','오류정정·디코더·컴파일러·벤치마크·HPC/AI 연계를 통합함','논리오류율·회로성능·재현성·산업 PoC'],
['1순위','양자통신·보안','PQC 전환·QKD·양자망 표준·상호운용 실증을 병렬 추진함','전환율·표준기고·실증망·상호운용'],
['1순위','양자센싱·계측','의료·항법·반도체·국방 사용사례와 시험인증을 연계함','감도·정확도·안정성·현장실증'],
['2순위','양자소자·소재','큐비트·광자·극저온·제어·패키징 공동인프라를 확보함','수율·결함·국산화·시설활용률'],
['장기기반','양자이론·시뮬레이션','개방형 검증문제·고영향 기초연구·후속 응용을 유지함','고영향 논문·오픈코드·응용전환'],
['응용발굴','양자화학·산업응용','제약·소재·에너지의 고가치 문제를 선별해 단계적 PoC를 수행함','예측개선·문제축소·ROI·특허'],
],[2.4,4.3,8.2,5.5],7.5)

# ES 8
add_heading(doc,'1.8 한국형 6대 전략축')
doc.add_picture(str(FIGS['strategy']),width=Cm(17.7)); add_caption(doc,'그림 1-5. 한국형 양자전략 6대 실행축',True)
add_caption(doc,'표 1-5. 전략축별 핵심 과제')
add_table(doc,['전략축','핵심 과제','의사결정 기준'],[
['기술 포트폴리오','6대 기술축·세부기술 정의서·검색식·전문가 검증을 구축함','연구기반·성장·시장·협업·리스크를 통합평가함'],
['공동 인프라','극저온·광원·파운드리·제어·통신 테스트베드·시험평가를 공동 구축함','시설 중복·활용률·접근성·공급망 기여도를 평가함'],
['국제 허브','미국·독일·영국·일본의 허브기관과 브리지 연구자를 기술군별로 연결함','공동 제1저자·시설공유·표준·특허·지속성을 평가함'],
['표준·시장 전환','PQC·양자망·센싱 시험인증과 산업 PoC·조달을 연계함','상호운용·경제성·규제·보안·전환비용을 검증함'],
['데이터 거버넌스','논문–특허–과제–표준–시장–기업 데이터를 통합함','데이터 커버리지·정확도·재현성·업데이트 주기를 관리함'],
['인력·조직','공동박사·포닥·산업파견·융합교육·전담 PMO를 운영함','반복 참여·산학 이동·핵심직무·기업 흡수율을 측정함'],
],[3.4,9.0,8.0],7.6)

# ES 9
add_heading(doc,'1.9 36개월 실행 로드맵')
doc.add_picture(str(FIGS['roadmap']),width=Cm(18.1)); add_caption(doc,'그림 1-6. 36개월 통합 실행 로드맵',True)
add_bullet(doc,'0~6개월은 데이터·기관·기술분류를 정비하고 6대 기술축, 핵심주체, 허브기관, 기준선 KPI를 확정함.',True)
add_bullet(doc,'6~18개월은 논문–특허–표준–시장 연계, 공동 인프라, 독립 벤치마크, 국제 공동과제, 산업 PoC를 수행함.')
add_bullet(doc,'18~36개월은 검증결과를 기반으로 제품·조달·공동표준·공급망·클러스터를 확대하고 성과가 낮은 영역은 전환·중단함.')
add_bullet(doc,'각 단계 종료 시 기술성·시장성·협업성·데이터 신뢰도를 재평가하여 포트폴리오를 동적으로 재배분함.')

add_caption(doc,'표 1-6. 단계별 핵심 산출물과 의사결정')
add_table(doc,['단계','핵심 산출물','주요 의사결정'],[
['0~6개월','통합 DB, 6대 기술축 정의서, 핵심인력·기관·국가맵, KPI 기준선','기술후보의 포함·제외와 우선순위를 확정함'],
['6~18개월','벤치마크, 공동시설 운영안, 국제협력·표준화 계획, 산업 PoC','성능·경제성·파트너 적합성과 계속투자 여부를 검증함'],
['18~36개월','제품·조달·표준·특허·공급망·클러스터, 투자 재배분안','상용화 확대·플랫폼 전환·과제 중단을 결정함'],
],[3.2,9.2,8.0],7.8)

# ES 10
add_heading(doc,'1.10 성과관리 KPI와 주요 리스크')
add_caption(doc,'표 1-7. 통합 성과관리 KPI')
add_table(doc,['영역','핵심 KPI','활용'],[
['연구활동','최근 3개월 모멘텀·고영향 논문·신규 토픽','연구 활력과 부상기술을 판단함'],
['융합·기술','복수분류·교차강도·벤치마크·재현성','융합확장과 실제 기술진전을 구분함'],
['주체·인력','반복 연구자·기관 전문화·산학 이동·기업 참여','핵심인력과 연구거점을 선정함'],
['협업','다기관·국제협력·네트워크 중심성·지속성','허브·브리지·협력공백을 식별함'],
['표준·IP','표준기고·채택·특허군·공동특허·라이선스','시장규칙과 지식재산 선점 여부를 판단함'],
['상용화','PoC·제품배치·조달·투자·매출·산업성과','확대·전환·중단 의사결정에 활용함'],
['데이터품질','DOI·기관·국가·기술분류 커버리지와 오류율','분석 신뢰도와 보완 우선순위를 관리함'],
],[3.4,8.4,8.4],7.8)

add_caption(doc,'표 1-8. 핵심 리스크와 대응')
add_table(doc,['리스크','영향','대응방향'],[
['기술 일정 지연','오류정정·수율·비용 목표가 반복 지연될 수 있음','단계형 투자·독립 검증·다중 플랫폼 포트폴리오를 운영함'],
['시장 과대기대','PoC가 실제 업무·ROI로 연결되지 않을 수 있음','사용사례 가치·통합비용·기회비용을 사전평가함'],
['표준 분절','지역·기관별 프로토콜과 인증이 달라질 수 있음','국제표준 연계·상호운용 시험·전환 로드맵을 강화함'],
['공급망 편중','극저온·광원·공정·제어 장비 의존이 발생할 수 있음','다중공급·공동시설·핵심부품 R&D를 추진함'],
['인력 미스매치','물리·공학·SW·산업인력 간 단절이 발생할 수 있음','융합교육·산업파견·공동박사·재교육을 확대함'],
['연구보안·수출통제','국제협력과 장비·데이터 이동이 지연될 수 있음','협력등급·IP·보안·수출통제 절차를 사전 설계함'],
['데이터 편향','arXiv·검색식·기관표준화에 따라 순위가 변동될 수 있음','다중 DB·민감도 분석·전문가 검증을 병행함'],
],[4.2,7.3,8.7],7.7)

# ES 11
add_heading(doc,'1.11 최종 전략메시지')
add_bullet(doc,'양자기술 경쟁의 중심은 논문 수 확대에서 검증 가능한 성능, 경제성, 표준, 공급망, 산업 배치로 이동하고 있음.',True)
add_bullet(doc,'한국은 양적 전면추격보다 기술축별 선택과 집중, 공용 인프라, 국제 허브 연계, 표준 선점, 산업수요 기반 PoC를 결합해야 함.')
add_bullet(doc,'단기에는 컴퓨팅·SW/AI, 통신·보안, 센싱·계측을 상용화 창구로 활용하고 소자·소재를 핵심 공급망으로 강화함.')
add_bullet(doc,'중장기에는 이론·시뮬레이션의 원천경쟁력을 유지하면서 제약·소재·에너지 등 고가치 산업문제를 통해 응용시장을 발굴함.')
add_bullet(doc,'정책·R&D·산업·표준화의 성과관리는 논문–특허–과제–표준–시장 데이터를 연결한 상시 모니터링으로 전환함.')
add_bullet(doc,'본 Executive Summary의 판단은 후속 통합보고서 작성 시 각 장의 상세 표·그래프·근거자료와 연결하여 최종 의사결정 근거로 활용함.')

add_caption(doc,'표 1-9. 경영·정책 의사결정 우선순위')
add_table(doc,['우선순위','의사결정 과제','즉시 조치'],[
['1','6대 기술축별 투자 포트폴리오 확정','기술 정의서·검색식·정량지표·전문가 평가를 통합함'],
['2','성능·경제성·표준 기반의 검증체계 구축','독립 벤치마크·시험인증·산업 PoC를 설계함'],
['3','핵심 인프라·부품·공급망 확보','공동시설·파운드리·극저온·광자·제어 로드맵을 수립함'],
['4','국제 허브기관·브리지 인력 연계','기술군별 공동과제·시설공유·공동표준을 추진함'],
['5','통합 데이터·KPI 기반 포트폴리오 관리','분기 모니터링과 확대·전환·중단 기준을 운영함'],
],[2.2,8.0,10.0],7.8)

# References
add_heading(doc,'참고자료')
refs=[
('[R1]','National Quantum Initiative, Annual Report FY2025','https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf'),
('[R2]','European Commission, Quantum Europe Strategy, 2025','https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy'),
('[R3]','The White House, Executive Order 14413: Ushering in the Next Frontier of Quantum Innovation, 2026','https://www.whitehouse.gov/presidential-actions/2026/06/ushering-in-the-next-frontier-of-quantum-innovation/'),
('[R4]','NIST CSRC, Post-Quantum Cryptography FIPS Approved','https://csrc.nist.gov/news/2024/postquantum-cryptography-fips-approved'),
('[R5]','DARPA, Quantum Benchmarking Initiative','https://www.darpa.mil/research/programs/quantum-benchmarking-initiative'),
('[R6]','과학기술정보통신부, 제1차 양자종합계획, 2026','https://www.msit.go.kr/eng/bbs/view.do?bbsSeqNo=42&mId=4&mPid=2&nttSeqNo=1222&sCode=eng'),
('[R7]','IBM Quantum, Hardware and Roadmap','https://www.ibm.com/quantum/hardware'),
('[R8]','ITU-T Y.Sup98, Technical Considerations Towards Quantum Networks, 2025','https://www.itu.int/rec/dologin_pub.asp?id=T-REC-Y.Sup98-202511-I!!PDF-E&lang=s&type=items'),
]
for rid,title,url in refs:
    p=doc.add_paragraph(); pformat(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(.42); p.paragraph_format.first_line_indent=Cm(-.42); r=p.add_run(rid+' '); set_run(r,8.8,True,NAVY); add_hyperlink(p,title,url)

cp=doc.core_properties; cp.title='arXiv 양자 분야 제1장 Executive Summary'; cp.subject='제2~7장 분석 결과와 최신 시장·기술·정책·표준화 자료의 통합 요약'; cp.author='OpenAI'; cp.keywords='arXiv, 양자, Executive Summary, 연구동향, 전략, 시장, 정책, 표준화'
settings=doc.settings._element; update=OxmlElement('w:updateFields'); update.set(qn('w:val'),'true'); settings.append(update)
doc.save(OUTPUT)
print(f'Created {OUTPUT}')

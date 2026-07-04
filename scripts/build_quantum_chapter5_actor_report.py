from __future__ import annotations

from pathlib import Path
import math

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path('artifacts')
ASSET_DIR = OUT_DIR / 'chapter5_assets'
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / 'arXiv_양자분야_제5장_연구주체분석_고품질_재작성_20260704.docx'

TOTAL = 2400
AUTHOR_LINKS = 12991
UNIQUE_AUTHORS = 10514
AVG_AUTHORS = 5.41
MEDIAN_AUTHORS = 4
ONE_PAPER_AUTHORS = 8576
REPEAT_AUTHORS = UNIQUE_AUTHORS - ONE_PAPER_AUTHORS
FIVE_PLUS_AUTHORS = 40
UNIQUE_FIRST = 2248
ONE_FIRST = 2111
REPEAT_FIRST = UNIQUE_FIRST - ONE_FIRST
VALID_INST = 1995
INST_LINKS = 3439
FIRST_COUNTRY_VALID = 1951
FIRST_COUNTRY_MISSING = TOTAL - FIRST_COUNTRY_VALID
INST_COUNTRY_VALID = 1998
INST_COUNTRY_MISSING = TOTAL - INST_COUNTRY_VALID
COUNTRY_LINKS = 2298
MULTI_COUNTRY = 241
COUNTRIES = 35

TOP_AUTHORS = [
    ('Ujjwal Sen','India',10),('Saif Al-Kuwari','China',10),('Guang-Can Guo','China',9),
    ('Claudia Linnhoff-Popien','Germany',9),('Henrik Dreyer','Germany',8),('Jonas Stein','Germany',8),
    ('Nouhaila Innan','United States',7),('Muhammad Shafique','United States',7),
    ('Sergi Abadal','United States',7),('Martin Roetteler','United States',7),
]
TOP_FIRST = [
    ('Md Aminur Hossain','India',4),('M. P. Telenkov','Russia',4),('Etienne Granet','Germany',4),
    ('Leo Sünkel','Germany',4),('Ivana Nikoloska','Netherlands',4),('Rakesh Thakur','India',3),
    ('Simon Becker','Italy',3),('Devashish Chaudhary','Australia',3),('Ioannis Krikidis','France',3),
    ('Kuan-Cheng Chen','United Kingdom',3),
]
TOP_INST = [
    ('University of Technology','Netherlands',47),('CNRS','France',34),
    ('Institute for Theoretical Nanoelectronics','Germany',27),
    ('University of Science and Technology','China',24),('University of Maryland','United States',23),
    ('University of California','United States',21),('Normal University','China',21),
    ('University of Tokyo','Japan',21),('Peking University','China',19),('Institute of Physics','Germany',18),
]
FIRST_COUNTRIES = [
    ('미국 (United States)',453),('중국 (China)',277),('독일 (Germany)',161),('영국 (United Kingdom)',113),
    ('일본 (Japan)',112),('인도 (India)',95),('이탈리아 (Italy)',91),('캐나다 (Canada)',71),
    ('호주 (Australia)',59),('프랑스 (France)',56),('스페인 (Spain)',51),('한국 (South Korea)',44),
]
INST_COUNTRIES = [
    ('미국 (United States)',524),('중국 (China)',296),('독일 (Germany)',199),('영국 (United Kingdom)',142),
    ('일본 (Japan)',124),('인도 (India)',111),('이탈리아 (Italy)',106),('프랑스 (France)',77),
    ('캐나다 (Canada)',73),('호주 (Australia)',70),('스페인 (Spain)',66),('한국 (South Korea)',48),
]
DUAL = [
    ('미국',453,524),('중국',277,296),('독일',161,199),('영국',113,142),('일본',112,124),
    ('인도',95,111),('이탈리아',91,106),('프랑스',56,77),('캐나다',71,73),('호주',59,70),
    ('스페인',51,66),('한국',44,48),
]
HHI = [('저자',1.17),('제1저자',4.76),('기관',14.38),('기관 국가',939.44),('제1저자 국가',984.20)]

NAVY='17365D'; BLUE='2F75B5'; SKY='D9EAF7'; GREEN='548235'; LIGHT_GREEN='E2F0D9'; ORANGE='ED7D31'; LIGHT_ORANGE='FCE4D6'; RED='C00000'; LIGHT_RED='F4CCCC'; PURPLE='7030A0'; LIGHT_PURPLE='E4DFEC'; GRAY='6B7280'; LIGHT_GRAY='F2F4F7'; DARK='243447'; WHITE='FFFFFF'; BORDER='CAD4DF'; YELLOW='FFD966'

def c(h): return '#'+h

def rgb(h): return RGBColor.from_string(h)

font_candidates=['/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc','/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc','/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc']
font_path=next((p for p in font_candidates if Path(p).exists()),None)
if font_path:
    fp=font_manager.FontProperties(fname=font_path); plt.rcParams['font.family']=fp.get_name()
else: plt.rcParams['font.family']='DejaVu Sans'
plt.rcParams['axes.unicode_minus']=False
plt.rcParams['font.weight']='bold'

# ---------------- charts ----------------
def base_axes(figsize=(11.6,5.2), grid='x'):
    fig,ax=plt.subplots(figsize=figsize,dpi=220); fig.patch.set_facecolor('white'); ax.set_facecolor('white')
    for s in ['top','right']: ax.spines[s].set_visible(False)
    ax.spines['left'].set_color(c(BORDER)); ax.spines['bottom'].set_color(c(BORDER))
    if grid: ax.grid(axis=grid,color=c(BORDER),lw=.7,alpha=.55)
    ax.set_axisbelow(True); ax.tick_params(colors=c(DARK),labelsize=9)
    return fig,ax

def save_ecosystem(path):
    fig,ax=plt.subplots(figsize=(12,5.2),dpi=220); fig.patch.set_facecolor('white'); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    cards=[
        (.03,.55,.18,.30,SKY,BLUE,'저자','10,514명','넓은 참여 저변'),
        (.225,.55,.18,.30,LIGHT_GREEN,GREEN,'제1저자','2,248명','주도 주체 분산'),
        (.42,.55,.18,.30,LIGHT_ORANGE,ORANGE,'기관','1,995개','분산형 연구 거점'),
        (.615,.55,.18,.30,LIGHT_PURPLE,PURPLE,'국가','35개국','글로벌 연구 축'),
        (.81,.55,.16,.30,LIGHT_GRAY,GRAY,'다국가 논문','241건','10.0%'),
    ]
    for x,y,w,h,fc,ec,title,val,sub in cards:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor=c(fc),edgecolor=c(ec),lw=2))
        ax.text(x+w/2,y+h-.065,title,ha='center',fontsize=11.5,fontweight='bold',color=c(NAVY))
        ax.text(x+w/2,y+.14,val,ha='center',fontsize=18,fontweight='bold',color=c(NAVY))
        ax.text(x+w/2,y+.055,sub,ha='center',fontsize=9.2,fontweight='bold',color=c(GRAY))
    ax.add_patch(FancyBboxPatch((.12,.12),.76,.25,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor=c(NAVY),lw=2.4))
    ax.text(.50,.29,'연구 주체 구조',ha='center',fontsize=11,fontweight='bold',color=c(GRAY))
    ax.text(.50,.22,'개인·기관은 넓은 롱테일  ×  국가는 주요국 중심의 중간 집중',ha='center',fontsize=14,fontweight='bold',color=c(NAVY))
    ax.text(.50,.155,'연구 생산의 분산성과 정책·인프라의 국가 집중이 동시에 존재함',ha='center',fontsize=10.2,fontweight='bold',color=c(DARK))
    for x in [.12,.315,.51,.705,.89]: ax.add_patch(FancyArrowPatch((x,.54),(.50,.38),arrowstyle='-|>',mutation_scale=14,lw=1.2,color=c(GRAY),alpha=.75))
    ax.text(.50,.94,'양자 분야 연구 주체 생태계 개요',ha='center',fontsize=16,fontweight='bold',color=c(NAVY))
    fig.tight_layout(pad=.2); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_author_longtail(path):
    fig,axs=plt.subplots(1,2,figsize=(11.8,5.0),dpi=220); fig.patch.set_facecolor('white')
    vals=[ONE_PAPER_AUTHORS,REPEAT_AUTHORS]; colors=[c(BLUE),c(SKY)]
    axs[0].pie(vals,colors=colors,startangle=90,counterclock=False,wedgeprops=dict(width=.35,edgecolor='white'))
    axs[0].text(0,.08,'81.6%',ha='center',fontsize=22,fontweight='bold',color=c(NAVY)); axs[0].text(0,-.17,'1편 참여 저자',ha='center',fontsize=10.5,fontweight='bold',color=c(DARK)); axs[0].set_title('전체 저자 생산성',fontsize=14,fontweight='bold',color=c(NAVY),pad=15); axs[0].axis('equal')
    vals2=[ONE_FIRST,REPEAT_FIRST]; colors2=[c(GREEN),c(LIGHT_GREEN)]
    axs[1].pie(vals2,colors=colors2,startangle=90,counterclock=False,wedgeprops=dict(width=.35,edgecolor='white'))
    axs[1].text(0,.08,'93.9%',ha='center',fontsize=22,fontweight='bold',color=c(NAVY)); axs[1].text(0,-.17,'1편 제1저자',ha='center',fontsize=10.5,fontweight='bold',color=c(DARK)); axs[1].set_title('제1저자 생산성',fontsize=14,fontweight='bold',color=c(NAVY),pad=15); axs[1].axis('equal')
    fig.suptitle('저자·제1저자 생산성의 롱테일 구조',fontsize=17,fontweight='bold',color=c(NAVY),y=.99)
    fig.text(.5,.02,f'전체 저자 2편 이상 {REPEAT_AUTHORS:,}명·18.4%  |  5편 이상 40명  |  제1저자 2편 이상 {REPEAT_FIRST}명·6.1%',ha='center',fontsize=10,fontweight='bold',color=c(GRAY))
    fig.tight_layout(rect=[0,.05,1,.94]); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def barh_rank(data,title,path,color=BLUE,xmax=None,label_country=True):
    labels=[f'{n} ({co})' if label_country else f'{n} ({co})' for n,co,v in data][::-1]; vals=[v for _,_,v in data][::-1]
    fig,ax=base_axes((12.0,6.0)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=c(SKY if color==BLUE else LIGHT_GREEN if color==GREEN else LIGHT_ORANGE),edgecolor=c(color),lw=1.1)
    for b,v in zip(bars,vals): ax.text(v+.15,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9.6,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.7); ax.set_xlim(0,xmax or max(vals)*1.18); ax.set_xlabel('논문 출현 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title(title,fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_institutions(path):
    labels=[f'{n} ({co})' for n,co,v in TOP_INST][::-1]; vals=[v for _,_,v in TOP_INST][::-1]
    fig,ax=base_axes((12,6.2)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=c(LIGHT_ORANGE),edgecolor=c(ORANGE),lw=1.1)
    for b,v in zip(bars,vals): ax.text(v+.7,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9.4,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.3); ax.set_xlim(0,53); ax.set_xlabel('기관-논문 연결 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('상위 기관 추출값 TOP10: 표준화 전 탐색 신호',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    ax.text(26,9.7,'※ 일반명·축약명·부속기관명이 혼재할 수 있어 확정 순위가 아닌 후보군으로 해석함',ha='center',fontsize=9,fontweight='bold',color=c(RED))
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def country_bar(data,title,path,color=BLUE,xmax=570):
    labels=[n for n,v in data][::-1]; vals=[v for n,v in data][::-1]
    fig,ax=base_axes((12,6.5)); y=np.arange(len(labels)); fc=SKY if color==BLUE else LIGHT_GREEN; bars=ax.barh(y,vals,color=c(fc),edgecolor=c(color),lw=1.0)
    for b,v in zip(bars,vals): ax.text(v+7,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.5); ax.set_xlim(0,xmax); ax.set_xlabel('논문 수 / 국가-논문 연결 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title(title,fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_dual(path):
    labels=[x[0] for x in DUAL][::-1]; a=np.array([x[1] for x in DUAL][::-1]); b=np.array([x[2] for x in DUAL][::-1]); y=np.arange(len(labels))
    fig,ax=base_axes((12,6.5)); ax.barh(y-.18,a,height=.34,color=c(SKY),edgecolor=c(BLUE),label='제1저자 국가'); ax.barh(y+.18,b,height=.34,color=c(LIGHT_GREEN),edgecolor=c(GREEN),label='전체 기관 국가')
    for yy,aa,bb in zip(y,a,b): ax.text(bb+6,yy,f'+{bb-aa}',va='center',fontsize=8.8,fontweight='bold',color=c(PURPLE))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold'); ax.set_xlim(0,570); ax.set_xlabel('논문 수 / 국가-논문 연결 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('국가 분석 이중 기준 비교: 주도성과 참여 외연',fontsize=16,fontweight='bold',color=c(NAVY),pad=18); ax.legend(frameon=False,loc='lower right')
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_scatter(path):
    fig,ax=base_axes((9.2,6.3),grid='both'); x=np.array([v1 for _,v1,v2 in DUAL]); y=np.array([v2 for _,v1,v2 in DUAL]); names=[n for n,_,_ in DUAL]
    ax.scatter(x,y,s=135,color=c(BLUE),edgecolor=c(NAVY),lw=1.0,zorder=3); maxv=550; ax.plot([0,maxv],[0,maxv],ls='--',lw=1.3,color=c(GRAY))
    offsets={'미국':(8,-5),'중국':(7,4),'독일':(7,3),'영국':(7,3),'프랑스':(7,3),'한국':(7,-10),'캐나다':(7,-10)}
    for n,xx,yy in zip(names,x,y):
        dx,dy=offsets.get(n,(6,3)); ax.annotate(n,(xx,yy),xytext=(dx,dy),textcoords='offset points',fontsize=9,fontweight='bold',color=c(DARK))
    ax.set_xlim(0,560); ax.set_ylim(0,560); ax.set_xlabel('제1저자 국가 건수',fontweight='bold',color=c(NAVY)); ax.set_ylabel('전체 기관 국가 건수',fontweight='bold',color=c(NAVY)); ax.set_title('국가별 연구 주도–참여 포지션',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    ax.text(330,490,'대각선 위쪽일수록\n비제1저자·국제협력 참여가 큼',ha='center',fontsize=10,fontweight='bold',color=c(PURPLE))
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_hhi(path):
    labels=[n for n,v in HHI]; vals=np.array([v for n,v in HHI]); logvals=np.log10(vals+1); colors=[c(BLUE),c(GREEN),c(ORANGE),c(LIGHT_PURPLE),c(PURPLE)]
    fig,ax=base_axes((10.8,5.0),grid='y'); x=np.arange(len(labels)); bars=ax.bar(x,logvals,color=colors,edgecolor=c(NAVY),lw=.8,width=.62)
    for b,v in zip(bars,vals): ax.text(b.get_x()+b.get_width()/2,b.get_height()+.08,f'{v:,.2f}',ha='center',fontsize=10,fontweight='bold',color=c(NAVY))
    ax.set_xticks(x); ax.set_xticklabels(labels,fontweight='bold'); ax.set_ylabel('log10(HHI×10,000 + 1)',fontweight='bold',color=c(NAVY)); ax.set_title('연구 주체별 집중도 비교',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    ax.text(2.0,2.95,'개인·기관: 매우 분산',ha='center',fontsize=10,fontweight='bold',color=c(GREEN)); ax.text(4.0,2.55,'국가: 상대적 집중',ha='center',fontsize=10,fontweight='bold',color=c(PURPLE))
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_context(path):
    fig,ax=plt.subplots(figsize=(12,5.5),dpi=220); fig.patch.set_facecolor('white'); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    boxes=[
        (.03,.58,.18,.28,SKY,BLUE,'미국','NQI·센터·산업·인력\n미국 1위 주도·참여'),
        (.225,.58,.18,.28,LIGHT_PURPLE,PURPLE,'EU·영국','인프라·생태계·인력\n독일·영국·프랑스 허브'),
        (.42,.58,.18,.28,LIGHT_ORANGE,ORANGE,'중국·일본','국가 주도 R&D·산업화\n높은 제1저자 비율'),
        (.615,.58,.18,.28,LIGHT_GREEN,GREEN,'한국','2035 인력·기업·표준\nTOP12 연구 기반'),
        (.81,.58,.16,.28,LIGHT_GRAY,GRAY,'시장','300+ 기업 채택·협업\n융합형 인재 수요'),
    ]
    for x,y,w,h,fc,ec,title,body in boxes:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor=c(fc),edgecolor=c(ec),lw=1.9)); ax.text(x+w/2,y+h-.06,title,ha='center',fontsize=11.5,fontweight='bold',color=c(NAVY)); ax.text(x+w/2,y+.105,body,ha='center',fontsize=9.2,fontweight='bold',color=c(DARK),linespacing=1.35)
    ax.add_patch(FancyBboxPatch((.12,.12),.76,.25,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor=c(NAVY),lw=2.4)); ax.text(.50,.29,'정량 결과',ha='center',fontsize=11,fontweight='bold',color=c(GRAY)); ax.text(.50,.22,'연구자·기관은 광범위  |  국가·인프라는 주요국 집중  |  협업 허브는 참여 격차에서 확인',ha='center',fontsize=12.3,fontweight='bold',color=c(NAVY)); ax.text(.50,.155,'연구주체의 양적 규모를 인력·인프라·산업생태계·국제협력 정책과 결합하여 해석함',ha='center',fontsize=9.8,fontweight='bold',color=c(DARK))
    for x in [.12,.315,.51,.705,.89]: ax.add_patch(FancyArrowPatch((x,.57),(.50,.39),arrowstyle='-|>',mutation_scale=14,lw=1.2,color=c(GRAY),alpha=.75))
    ax.text(.50,.95,'연구 주체 분포를 형성하는 정책·시장 환경',ha='center',fontsize=16,fontweight='bold',color=c(NAVY)); fig.tight_layout(pad=.2); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

FIGS={
'ecosystem':ASSET_DIR/'fig_5_1_ecosystem.png','longtail':ASSET_DIR/'fig_5_2_longtail.png','authors':ASSET_DIR/'fig_5_3_authors.png','first':ASSET_DIR/'fig_5_4_first_authors.png','inst':ASSET_DIR/'fig_5_5_institutions.png','fcountry':ASSET_DIR/'fig_5_6_first_country.png','icountry':ASSET_DIR/'fig_5_7_inst_country.png','dual':ASSET_DIR/'fig_5_8_dual.png','scatter':ASSET_DIR/'fig_5_9_scatter.png','hhi':ASSET_DIR/'fig_5_10_hhi.png','context':ASSET_DIR/'fig_5_11_context.png'}
save_ecosystem(FIGS['ecosystem']); save_author_longtail(FIGS['longtail']); barh_rank(TOP_AUTHORS,'상위 반복 참여 저자 TOP10: 저자명(국가) 병기',FIGS['authors'],BLUE,12); barh_rank(TOP_FIRST,'상위 제1저자 TOP10: 저자명(국가) 병기',FIGS['first'],GREEN,5.2); save_institutions(FIGS['inst']); country_bar(FIRST_COUNTRIES,'제1저자 국가 기준 TOP12: 연구 주도 축',FIGS['fcountry'],BLUE,500); country_bar(INST_COUNTRIES,'전체 기관 국가 기준 TOP12: 연구 참여 범위',FIGS['icountry'],GREEN,570); save_dual(FIGS['dual']); save_scatter(FIGS['scatter']); save_hhi(FIGS['hhi']); save_context(FIGS['context'])

# ---------------- doc helpers ----------------
def set_cell_shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_border(cell,color=BORDER,size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ['top','left','bottom','right']:
        el=borders.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)

def set_cell_margins(cell,top=80,start=90,bottom=80,end=90):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_run(run,size=10,bold=False,color=DARK):
    run.font.name='맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=rgb(color)

def pformat(p,before=0,after=4,line=1.25,keep=False):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line; pf.keep_with_next=keep

def add_page_number(p):
    run=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text='PAGE'; e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); run._r.extend([b,i,e])

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); pformat(p,before=12 if level==1 else 8,after=6,line=1.0,keep=True); r=p.add_run(text); set_run(r,15 if level==1 else 12,True,NAVY)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'3'); bottom.set(qn('w:color'),BLUE); pBdr.append(bottom); pPr.append(pBdr)
    return p

def add_bullet(doc,text,major=False):
    p=doc.add_paragraph(); pformat(p,after=3,line=1.28); p.paragraph_format.left_indent=Cm(.48); p.paragraph_format.first_line_indent=Cm(-.34); r=p.add_run('□ ' if major else '· '); set_run(r,10.2,True,BLUE if major else GRAY); r=p.add_run(text); set_run(r,10.2,major,DARK); return p

def add_caption(doc,text,figure=False):
    p=doc.add_paragraph(); pformat(p,before=3,after=6,line=1.0,keep=not figure); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); set_run(r,9.2,True,GRAY if figure else NAVY)

def add_table(doc,headers,rows,widths,font=8.25):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,NAVY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,after=0,line=1.05); r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cell=cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,WHITE if ri%2==0 else LIGHT_GRAY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pformat(p,after=0,line=1.08); r=p.add_run(str(v)); set_run(r,font,i==0,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0); return table

def add_callout(doc,title,lines,fill=SKY,accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,accent,'8'); set_cell_margins(cell,130,150,130,150); p=cell.paragraphs[0]; pformat(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=cell.add_paragraph(); pformat(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.first_line_indent=Cm(-.25); r=p.add_run('• '); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_hyperlink(p,text,url):
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rPr.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)

# ---------------- document ----------------
doc=Document(); sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35); sec.header_distance=Cm(.55); sec.footer_distance=Cm(.55)
style=doc.styles['Normal']; style.font.name='맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); style.font.size=Pt(10)
header=sec.header.paragraphs[0]; r=header.add_run('arXiv 양자 분야 논문 분석 | 제5장 연구 주체 분석'); set_run(r,8.4,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('분석 기준: 2025.06~2026.05, 12개월·2,400건  |  '); set_run(r,8.2,False,GRAY); add_page_number(footer)

cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; cell=cover.cell(0,0); set_cell_shading(cell,NAVY); set_cell_margins(cell,720,300,720,300); set_cell_border(cell,NAVY,'0'); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CHAPTER 5'); set_run(r,14,True,'9DC3E6'); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('연구 주체 분석'); set_run(r,27,True,WHITE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('2025년 6월~2026년 5월 | 12개월·2,400건'); set_run(r,12,True,SKY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,before=16); r=p.add_run('저자·기관의 롱테일 × 국가별 주도성·참여 외연 × 정책·인력 생태계'); set_run(r,10.8,True,YELLOW)

doc.add_paragraph(); kpis=[('10,514명','고유 저자'),('2,248명','고유 제1저자'),('1,995개','유효 기관'),('35개국','관찰 국가'),('10.0%','다국가 논문')]; t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(val,label) in enumerate(kpis):
    cell=t.cell(0,i); set_cell_shading(cell,[SKY,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]); set_cell_border(cell,[BLUE,GREEN,ORANGE,PURPLE,GRAY][i],'7'); set_cell_margins(cell,100,70,100,70); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,15,True,NAVY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_run(r,8.6,True,GRAY)
add_callout(doc,'본 장의 핵심 판단',[
    '개인과 기관 단위에서는 상위 집중도가 매우 낮아 폭넓은 롱테일 연구 생태계가 형성된 것으로 보임.',
    '국가 단위에서는 미국·중국·독일·영국·일본 중심의 상대적 집중이 나타나 연구 인프라와 정책 역량이 주요국에 집중된 것으로 판단함.',
    '제1저자 국가와 전체 기관 국가를 병렬 적용하여 연구 주도성과 공동연구 참여 외연을 구분하고, 미국·독일·영국·프랑스를 국제 협업 허브 후보로 식별함.',
],SKY,BLUE)
doc.add_page_break()

add_heading(doc,'5.1 분석 목적 및 앞선 장과의 연결')
add_bullet(doc,'제3장에서 확인한 월평균 200편의 연구 생산을 실제로 누가 만들고 있는지 저자·제1저자·기관·국가 단위로 분석함.',True)
add_bullet(doc,'제4장에서 확인한 양자물리(quant-ph) 코어와 컴퓨터과학·응집물질·광학·암호/보안 확장 구조가 어떤 연구 주체에서 형성되는지 연결하여 해석함.')
add_bullet(doc,'저자 분석은 참여 저변과 반복 참여자를, 제1저자 분석은 연구 주도 기회를, 기관 분석은 수행 거점과 인프라 후보를 파악함.')
add_bullet(doc,'국가 분석은 제1저자 국가 기준과 전체 기관 국가 기준을 병렬 적용하여 연구 주도성과 국제 참여 외연을 구분함.')
add_bullet(doc,'본 장의 상위 저자·기관·국가는 제6장 공저·기관·국가 협업 네트워크의 허브 후보로 활용하고 제7장 국제협력·인력·R&D 전략으로 확장함.')
add_caption(doc,'표 5-1. 연구 주체 분석 단위와 해석 목적')
add_table(doc,['분석 단위','집계 기준','핵심 질문','후속 활용'],[
    ['전체 저자','논문별 저자 출현 및 고유 저자','연구 참여 저변과 반복 참여자는 누구인가','공저 네트워크·신흥 연구자 탐지'],
    ['제1저자','논문별 첫 번째 저자','연구 주도 기회가 특정 소수에게 집중되는가','주도 연구자·핵심 인력 후보'],
    ['기관','정제된 기관명과 기관-논문 연결','어떤 연구 거점과 인프라가 반복적으로 등장하는가','기관 협업·벤치마킹·인프라 맵'],
    ['제1저자 국가','논문별 1개 국가','어느 국가가 연구 생산을 주도하는가','국가별 주도성·정책 역량 비교'],
    ['전체 기관 국가','논문별 고유 국가 복수 집계','어느 국가가 공동연구에 폭넓게 참여하는가','국제협력 허브·파트너 탐색'],
], [3.1,5.1,6.0,5.1],8.1)

add_heading(doc,'5.2 연구 주체 생태계 개요')
doc.add_picture(str(FIGS['ecosystem']),width=Cm(18.0)); add_caption(doc,'그림 5-1. 양자 분야 연구 주체 생태계 개요',True)
add_bullet(doc,f'2,400건에서 저자-논문 연결 {AUTHOR_LINKS:,}건, 고유 저자 {UNIQUE_AUTHORS:,}명, 논문당 평균 {AVG_AUTHORS:.2f}명·중앙값 {MEDIAN_AUTHORS}명이 확인됨.',True)
add_bullet(doc,f'고유 제1저자 {UNIQUE_FIRST:,}명과 유효 기관 {VALID_INST:,}개가 확인되어 연구 주도 인력과 수행기관이 광범위하게 분산된 구조를 보임.')
add_bullet(doc,f'제1저자 국가는 {FIRST_COUNTRY_VALID:,}건·81.3%, 전체 기관 국가는 {INST_COUNTRY_VALID:,}건·83.3%에서 확인되어 국가 분석의 해석 가능 범위를 확보함.')
add_bullet(doc,f'전체 기관 국가 기준 국가-논문 연결은 {COUNTRY_LINKS:,}건이며 다국가 논문은 {MULTI_COUNTRY}건·10.0%로 확인되어 국가 간 참여 외연과 실제 국제협력 범위를 분리해 볼 필요가 있음.')
add_caption(doc,'표 5-2. 연구 주체 핵심 지표 요약')
add_table(doc,['구분','지표','값','해석'],[
    ['저자','고유 저자 / 저자-논문 연결',f'{UNIQUE_AUTHORS:,}명 / {AUTHOR_LINKS:,}건','연구 참여 저변이 매우 넓음'],['저자','평균 / 중앙값',f'{AVG_AUTHORS:.2f}명 / {MEDIAN_AUTHORS}명','팀 연구가 일반적이며 대형 팀 일부가 평균을 높임'],['제1저자','고유 제1저자',f'{UNIQUE_FIRST:,}명','논문 주도 기회가 광범위하게 분산됨'],['기관','유효 기관 / 기관-논문 연결',f'{VALID_INST:,}개 / {INST_LINKS:,}건','대학·연구소·기업 등 다수 거점이 참여함'],['국가','제1저자 국가 확인',f'{FIRST_COUNTRY_VALID:,}건·81.3%','논문별 연구 주도 국가를 파악함'],['국가','기관 국가 확인 / 연결',f'{INST_COUNTRY_VALID:,}건·83.3% / {COUNTRY_LINKS:,}건','전체 연구 참여 범위와 공동연구 외연을 파악함'],['국제협력','다국가 논문',f'{MULTI_COUNTRY}건·10.0%','제6장 국제협업 분석의 기준값임'],
], [2.6,4.7,4.4,7.0],8.0)

add_heading(doc,'5.3 저자 생산성: 신규 유입 중심의 롱테일 구조')
doc.add_picture(str(FIGS['longtail']),width=Cm(17.3)); add_caption(doc,'그림 5-2. 저자·제1저자 생산성의 롱테일 구조',True)
add_bullet(doc,f'1편에만 참여한 저자는 {ONE_PAPER_AUTHORS:,}명·81.6%로 나타나 분석기간 동안 다수의 신규·일회 참여 연구자가 유입된 것으로 보임.',True)
add_bullet(doc,f'2편 이상 반복 참여 저자는 {REPEAT_AUTHORS:,}명·18.4%, 5편 이상은 {FIVE_PLUS_AUTHORS}명으로 확인되어 반복 참여 허브는 소수이나 전체 저변은 매우 넓음.')
add_bullet(doc,'롱테일 구조는 기술 진입장벽이 낮다는 의미가 아니라 이론·실험·소자·알고리즘·응용 등 다양한 역할과 기술군이 병렬로 참여하는 다층 생태계로 해석해야 함.')
add_bullet(doc,'상위 반복 저자만 추적할 경우 신흥 연구자와 신규 기술군을 놓칠 수 있으므로 최근 등장, 반복 참여, 기술군 이동, 공동저자 중심성을 병렬 분석해야 함.')
add_caption(doc,'표 5-3. 저자 생산성 구조와 전략적 의미')
add_table(doc,['지표','값','정량 해석','전략 활용'],[
    ['1편 참여 저자',f'{ONE_PAPER_AUTHORS:,}명·81.6%','단기·신규 참여층이 대다수를 구성함','신흥 인력·신규 기관 유입을 추적함'],['2편 이상 저자',f'{REPEAT_AUTHORS:,}명·18.4%','반복 참여 연구자 풀이 존재함','기술군별 핵심 연구자 후보를 도출함'],['5편 이상 저자',f'{FIVE_PLUS_AUTHORS}명','고빈도 연구자는 매우 제한적임','공동연구 허브·연속 과제 후보를 검토함'],['저자 HHI×10,000','1.17','개인 단위 집중도가 극히 낮음','단순 순위보다 네트워크·영향력 지표를 결합함'],
], [3.5,3.8,6.0,6.3],8.3)

add_heading(doc,'5.4 상위 반복 참여 저자: 허브 후보와 해석 한계')
doc.add_picture(str(FIGS['authors']),width=Cm(18.1)); add_caption(doc,'그림 5-3. 상위 반복 참여 저자 TOP10: 저자명(국가) 병기',True)
add_bullet(doc,'상위 저자는 7~10편 범위이며 상위 10명의 논문 출현 합계는 82건·전체 논문 대비 3.42%로 확인됨.',True)
add_bullet(doc,'저자-논문 연결 대비 비중은 0.63%에 불과하여 양자 분야 전체가 특정 스타 연구자에게 의존하는 구조로 보기는 어려움.')
add_bullet(doc,'중국·독일·미국·인도 소속 연구자가 반복적으로 나타나며, 이는 국가별 핵심 연구그룹과 공동연구 네트워크의 입구 후보로 활용할 수 있음.')
add_bullet(doc,'논문 수는 영향력·기술수준의 대체지표가 아니므로 피인용, 참고문헌 영향력, 기술군별 반복 참여, 공동저자 중심성, 특허·과제 연계를 추가해야 함.')
add_caption(doc,'표 5-4. 상위 반복 참여 저자 TOP10')
add_table(doc,['순위','저자(국가)','논문 수','전체 논문 대비','활용 관점'],[[i+1,f'{n} ({co})',v,f'{v/TOTAL*100:.2f}%','반복 참여·네트워크 허브 후보'] for i,(n,co,v) in enumerate(TOP_AUTHORS)],[1.3,7.5,2.5,3.0,6.0],8.0)

add_heading(doc,'5.5 제1저자: 연구 주도 기회의 광범위한 분산')
add_bullet(doc,f'고유 제1저자는 {UNIQUE_FIRST:,}명이며 논문 1건당 고유 제1저자 평균 생산성은 {TOTAL/UNIQUE_FIRST:.2f}건으로 확인됨.',True)
add_bullet(doc,f'1편의 논문에서만 제1저자로 등장한 연구자는 {ONE_FIRST:,}명·93.9%, 2편 이상은 {REPEAT_FIRST}명·6.1%로 나타남.')
add_bullet(doc,'제1저자 순위는 분야별 저자 순서 관행을 고려해야 하며 교신저자·책임저자·실험 리더를 완전하게 대체하지 않음.')
doc.add_picture(str(FIGS['first']),width=Cm(18.0)); add_caption(doc,'그림 5-4. 상위 제1저자 TOP10: 저자명(국가) 병기',True)
add_bullet(doc,'상위 제1저자 10명의 합계는 35건·전체의 1.46%로 나타나 연구 주도권도 특정 소수에게 집중되지 않음.',True)
add_bullet(doc,'상위군은 인도·러시아·독일·네덜란드·이탈리아·호주·프랑스·영국으로 분산되어 미국·중국 중심의 국가 총량과 개인 단위 주도 연구자의 분포가 동일하지 않음을 보임.')
add_caption(doc,'표 5-5. 상위 제1저자 TOP10')
add_table(doc,['순위','제1저자(국가)','논문 수','전체 논문 대비','해석'],[[i+1,f'{n} ({co})',v,f'{v/TOTAL*100:.2f}%','반복 제1저자·연구 주도 후보'] for i,(n,co,v) in enumerate(TOP_FIRST)],[1.3,7.5,2.5,3.0,6.0],8.0)

add_heading(doc,'5.6 기관: 연구 인프라 거점과 표준화 전 탐색 신호')
doc.add_picture(str(FIGS['inst']),width=Cm(18.1)); add_caption(doc,'그림 5-5. 상위 기관 추출값 TOP10: 기관명 표준화 전 기준',True)
add_bullet(doc,f'유효 기관은 {VALID_INST:,}개, 기관-논문 연결은 {INST_LINKS:,}건으로 확인되어 대학·국가연구기관·전문연구소가 폭넓게 참여함.',True)
add_bullet(doc,'상위 기관은 University of Technology, CNRS, Institute for Theoretical Nanoelectronics, University of Maryland, University of Tokyo 등으로 나타남.')
add_bullet(doc,'상위 10개 기관 연결은 255건·전체 기관 연결의 7.41%로 낮아 대표 허브 기관과 광범위한 분산 기관군이 공존함.')
add_bullet(doc,'University of Technology, Normal University, Institute of Physics는 국가·도시·상위기관 정보가 생략된 일반명일 수 있어 확정 순위가 아닌 표준화 전 후보로 관리함.')
add_bullet(doc,'최종 기관 경쟁력 분석에서는 ROR·OpenAlex·GRID·Crossref·ORCID를 활용해 동일기관·부속기관·약어를 통합하고 분야별 전문화·협업 중심성을 재산출해야 함.')
add_caption(doc,'표 5-6. 상위 기관 추출값 TOP10')
add_table(doc,['순위','기관명(추정 국가)','연결 수','기관 연결 대비','데이터 상태'],[[i+1,f'{n} ({co})',v,f'{v/INST_LINKS*100:.2f}%','표준화 전 탐색값'] for i,(n,co,v) in enumerate(TOP_INST)],[1.3,8.0,2.3,3.0,5.7],8.0)

add_heading(doc,'5.7 국가 분석 체계: 제1저자 국가와 전체 기관 국가')
add_caption(doc,'표 5-7. 국가 분석의 이중 기준')
add_table(doc,['기준','사용 정보','집계 단위','주요 의미','유의사항'],[
    ['제1저자 국가','제1저자 소속기관 국가','논문별 1개 국가','연구 생산의 주도 축과 핵심 인력 기반','저자 순서 관행·미확인 449건을 고려함'],
    ['전체 기관 국가','모든 저자 소속기관 국가','논문별 국가 중복 제거 후 복수 집계','공동연구 참여 범위와 국제 연결성','다국가 논문은 복수 국가에 반영되어 합계가 논문 수를 초과함'],
], [3.2,4.2,4.4,5.6,5.3],8.1)
add_bullet(doc,'두 기준을 동일한 분모로 직접 비교하지 않고 제1저자 건수, 기관 국가 연결 수, 전체 논문 대비 비중, 각 유효 집합 대비 비중을 구분하여 제시함.',True)
add_bullet(doc,'제1저자 기준만 적용하면 공동저자 참여국의 외연을 과소평가하고, 기관 국가만 적용하면 연구 주도 역할을 구분하기 어려우므로 병렬 분석이 필요함.')

add_heading(doc,'5.8 제1저자 국가: 글로벌 연구 주도 축')
doc.add_picture(str(FIGS['fcountry']),width=Cm(18.1)); add_caption(doc,'그림 5-6. 제1저자 국가 기준 TOP12: 연구 주도 축',True)
add_bullet(doc,f'제1저자 국가가 확인된 논문은 {FIRST_COUNTRY_VALID:,}건·81.3%이며 총 {COUNTRIES}개국이 확인됨.',True)
add_bullet(doc,'미국 453건, 중국 277건, 독일 161건이 상위 3개국을 형성하며 영국·일본·인도·이탈리아가 뒤를 이음.')
add_bullet(doc,'상위 5개국은 1,116건·전체의 46.5%, 상위 10개국은 1,488건·전체의 62.0%이자 국가 확인 논문의 76.3%를 차지함.')
add_bullet(doc,'개인·기관 단위 분산성과 달리 국가 단위에서는 장기 투자·인프라·인력·기업 생태계를 보유한 주요국 중심 구조가 뚜렷함.')
add_caption(doc,'표 5-8. 제1저자 국가 TOP12')
add_table(doc,['순위','국가','논문 수','전체 대비','국가 확인 대비'],[[i+1,n,v,f'{v/TOTAL*100:.2f}%',f'{v/FIRST_COUNTRY_VALID*100:.2f}%'] for i,(n,v) in enumerate(FIRST_COUNTRIES)],[1.3,6.8,2.5,3.2,3.5],8.0)

add_heading(doc,'5.9 전체 기관 국가: 연구 참여 범위와 국제 허브성')
doc.add_picture(str(FIGS['icountry']),width=Cm(18.1)); add_caption(doc,'그림 5-7. 전체 기관 국가 기준 TOP12: 연구 참여 범위',True)
add_bullet(doc,f'기관 국가가 하나 이상 확인된 논문은 {INST_COUNTRY_VALID:,}건·83.3%, 국가-논문 연결은 {COUNTRY_LINKS:,}건으로 확인됨.',True)
add_bullet(doc,'미국 524건, 중국 296건, 독일 199건, 영국 142건, 일본 124건 순으로 나타남.')
add_bullet(doc,'상위 10개국은 1,722건으로 전체 국가-논문 연결의 74.9%를 차지하여 국제 참여 외연도 주요국 중심으로 형성됨.')
add_bullet(doc,f'국가 확인 논문당 평균 국가는 {COUNTRY_LINKS/INST_COUNTRY_VALID:.2f}개이며 다국가 논문은 {MULTI_COUNTRY}건·10.0%로 나타나 다수의 국내·단일국 연구와 선택적 국제협력이 공존함.')
add_caption(doc,'표 5-9. 전체 기관 국가 TOP12')
add_table(doc,['순위','국가','국가-논문 연결','전체 논문 대비','국가 연결 대비'],[[i+1,n,v,f'{v/TOTAL*100:.2f}%',f'{v/COUNTRY_LINKS*100:.2f}%'] for i,(n,v) in enumerate(INST_COUNTRIES)],[1.3,6.4,3.1,3.2,3.5],7.9)

add_heading(doc,'5.10 두 국가 기준 비교: 주도성–참여 외연의 국가 유형')
doc.add_picture(str(FIGS['dual']),width=Cm(18.1)); add_caption(doc,'그림 5-8. 국가 분석 이중 기준 비교: 제1저자 국가와 전체 기관 국가',True)
add_bullet(doc,'미국은 453건에서 524건으로 +71건, 독일은 +38건, 영국은 +29건, 프랑스는 +21건 증가하여 비제1저자·국제 공동연구 참여 외연이 상대적으로 큼.',True)
add_bullet(doc,'중국·일본·캐나다·한국은 제1저자/기관 국가 비율이 90% 이상으로 나타나 연구 참여가 제1저자 역할과 비교적 밀접하게 연결됨.')
add_bullet(doc,'참여 확대 폭은 국제협력 허브성의 보조 신호이며 저자 순서 관행, 공동소속, 국가정보 커버리지의 영향을 고려해야 함.')
add_caption(doc,'표 5-10. 제1저자 국가와 기관 국가 비교 TOP12')
add_table(doc,['순위','국가','제1저자 국가','기관 국가','참여 확대','제1저자/기관 비율'],[[i+1,n,a,b,f'+{b-a}',f'{a/b*100:.1f}%'] for i,(n,a,b) in enumerate(DUAL)],[1.2,3.0,2.7,2.7,2.7,4.4],8.0)

doc.add_picture(str(FIGS['scatter']),width=Cm(14.8)); add_caption(doc,'그림 5-9. 국가별 연구 주도–참여 포지션',True)
add_bullet(doc,'대각선 위쪽으로 멀리 위치하는 미국·독일·영국·프랑스는 국제 공동연구와 비제1저자 참여를 통해 연구 네트워크 외연을 확장하는 허브 후보로 보임.',True)
add_bullet(doc,'한국은 44건에서 48건으로 +4건, 캐나다는 +2건으로 나타나 자체 주도 논문 비중이 높으나 국제 참여 외연은 상대적으로 제한적일 수 있음.')
add_bullet(doc,'제6장에서는 국가쌍·기관쌍과 네트워크 중심성을 결합하여 허브성 가설을 직접 검증해야 함.')

add_heading(doc,'5.11 집중도 비교: 개인·기관 분산과 국가 집중의 이중 구조')
doc.add_picture(str(FIGS['hhi']),width=Cm(17.0)); add_caption(doc,'그림 5-10. 연구 주체별 HHI 집중도 비교',True)
add_bullet(doc,'저자 HHI 1.17, 제1저자 HHI 4.76, 기관 HHI 14.38로 나타나 개인·기관 단위 생산은 매우 분산됨.',True)
add_bullet(doc,'제1저자 국가 HHI 984.20, 기관 국가 HHI 939.44로 나타나 국가 단위에서는 상대적으로 높은 집중 구조가 확인됨.')
add_bullet(doc,'기관 국가 HHI가 제1저자 국가보다 낮은 것은 공동연구 참여국이 추가되면서 분포가 소폭 분산되기 때문으로 보임.')
add_bullet(doc,'양자 연구는 다수 연구자·기관이 참여하지만 대규모 인프라·공공투자·기업 생태계를 갖춘 국가가 연구량을 견인하는 다층 구조로 판단함.')
add_caption(doc,'표 5-11. 연구 주체별 집중도와 해석')
add_table(doc,['분석 단위','HHI×10,000','보조 지표','구조 진단'],[
    ['저자','1.17','상위 10명 82건','매우 낮은 집중도·넓은 참여 저변'],['제1저자','4.76','상위 10명 35건','주도 연구자도 분산됨'],['기관','14.38','상위 10개 기관 255건','주요 허브와 롱테일 기관이 공존함'],['기관 국가','939.44','상위 10개국 1,722 연결','참여 외연은 주요국 중심이나 주도 기준보다 소폭 분산됨'],['제1저자 국가','984.20','상위 10개국 1,488건','연구 주도는 주요국 중심으로 집중됨'],
], [3.4,3.2,5.2,7.0],8.2)

add_heading(doc,'5.12 최신 시장·기술·정책 환경과 연구 주체 분포의 통합 해석')
doc.add_picture(str(FIGS['context']),width=Cm(18.0)); add_caption(doc,'그림 5-11. 연구 주체 분포를 형성하는 정책·시장 환경',True)
add_bullet(doc,'미국 NQI는 기초과학, 연구센터, 양자경제개발컨소시엄, 인력, 인프라, 국제협력의 다층 체계를 운영함. 미국이 제1저자·기관 국가 모두 1위를 차지하는 결과는 이러한 광범위한 생태계와 구조적으로 부합함 [R1].',True)
add_bullet(doc,'EU Quantum Europe Strategy는 연구혁신, 양자인프라, 생태계·공급망·산업화, 우주·듀얼유스, 기술·인력 이동성을 핵심축으로 제시함. 독일·영국·프랑스의 기관 참여 확대는 유럽의 다국가 인프라·컨소시엄 구조와 연결해 해석할 수 있음 [R2].')
add_bullet(doc,'영국 National Quantum Strategy는 10년 장기전략과 연구·산업·인력·국제협력을 결합함. 영국의 제1저자 113건 대비 기관 국가 142건은 국내 주도성과 국제 네트워크 참여가 병행되는 신호로 보임 [R3].')
add_bullet(doc,'일본은 Integrated Innovation Strategy 2025에서 양자 산업화 원년과 Quantum-Ready Japan을 제시하고 RIKEN·AIST·NICT·QST 등 연구기반과 산학관 협력을 강조함. 일본의 높은 제1저자/기관 비율은 자체 주도 기반의 강도를 보여주는 보조 신호임 [R4].')
add_bullet(doc,'한국의 제1차 양자종합계획은 2035년 전문인력 1만 명, 양자기업 2천 개, 표준 선도국 목표와 컴퓨팅·통신·센서·소재부품·알고리즘 클러스터를 제시함. 현재 TOP12 연구 기반을 분야 특화·국제협력·인력 유입으로 확장할 필요가 있음 [R5].')
add_bullet(doc,'McKinsey는 300개 이상의 기업이 양자컴퓨팅을 채택·협업하고 있으며 산업화에는 물리학자뿐 아니라 엔지니어·소프트웨어·비즈니스 인력이 필요하다고 제시함. 본 데이터의 넓은 저자·기관 저변은 융합 인재 생태계의 잠재적 공급 기반으로 해석함 [R6].')
add_callout(doc,'정량–정성 통합 판단',[
    '연구자·기관의 분산성은 기술 다변화와 신규 참여가 가능한 개방형 생태계를 의미함.',
    '국가 집중은 양자컴퓨터·센싱·네트워크에 필요한 대규모 인프라·공공투자·인력정책의 영향을 반영하는 것으로 보임.',
    '국가별 순위만으로 기술수준을 단정하지 않고 고영향 논문, 특허, 연구비, 실증 인프라, 표준화 참여, 국제협력 중심성을 함께 평가함.',
],LIGHT_ORANGE,ORANGE)

add_heading(doc,'5.13 한국의 연구 주체 포지션과 실행 과제')
add_bullet(doc,'한국은 제1저자 국가 44건·전체 1.83%, 전체 기관 국가 48건·전체 2.00%로 두 기준 모두 12위에 위치함.',True)
add_bullet(doc,'제1저자/기관 국가 비율은 91.7%로 높아 관찰된 참여가 제1저자 역할과 밀접하나 국제 공동연구 참여 확대 폭은 +4건으로 주요 유럽국보다 제한적임.')
add_bullet(doc,'단순 양적 확대보다 제4장에서 확인한 양자센싱·광학, 암호/보안, 기계학습, 소자·소재 교차 분야에서 국내 기관·연구자 특화도를 분석해야 함.')
add_bullet(doc,'정책적으로 인력 1만 명 목표를 논문 저자 수의 단순 증가가 아니라 반복 참여율, 박사후·산업 인력 이동, 국제 공동연구, 기업·연구소 유입 지표로 세분화할 필요가 있음.')
add_caption(doc,'표 5-12. 한국 연구 주체 분석 결과와 권장 실행과제')
add_table(doc,['진단 항목','현재 신호','해석','권장 과제'],[
    ['연구 주도','제1저자 국가 44건·12위','기본 연구 생산 기반은 확보함','기술군별 국내 제1저자·기관 특화도를 산출함'],['연구 참여','기관 국가 48건·12위','참여와 주도 역할이 유사함','미국·독일·영국·프랑스 허브와 공동과제를 확대함'],['국제 외연','참여 확대 +4건','국제 연결의 추가 확장 여지가 큼','국가쌍·기관쌍·브리지 연구자를 발굴함'],['인력 기반','전체 저자 데이터 내 국내 인력 분리 미완료','국가 총량만으로 인력 질을 판단하기 어려움','ORCID·기관표준화·분야별 경력·산학 이동을 연계함'],['전략 연계','2035 인력·기업·표준 목표','연구·산업·표준 인력을 연결할 필요가 있음','논문–특허–표준–기업 생태계 통합지표를 구축함'],
], [3.2,4.1,5.8,6.4],8.1)

add_heading(doc,'5.14 전략적 시사점 및 제6·7장 연결')
add_bullet(doc,'핵심 연구자 선정은 단순 논문 수가 아니라 기술군별 반복 참여, 최근 성장, 공동저자 중심성, 영향력, 소속기관 인프라를 결합하여 수행함.',True)
add_bullet(doc,'기관 벤치마킹은 표준화 전 순위를 확정값으로 사용하지 않고 기관 식별자 통합 후 분야 전문화지수·국제협력 중심성·기업 연계를 재산출함.')
add_bullet(doc,'국가별 전략은 제1저자 국가를 주도성 지표, 기관 국가를 협력 외연 지표로 활용하여 ‘연구 선도국’, ‘국제 허브국’, ‘자체 주도형 국가’를 구분함.')
add_bullet(doc,'제6장에서는 미국–독일, 미국–영국, 미국–인도 등 국가쌍과 상위 기관쌍을 네트워크로 분석하여 본 장의 허브 가설을 검증함.')
add_bullet(doc,'제7장에서는 연구자·기관·국가 구조를 유망기술 후보, 국제협력 대상, 인력정책, 연구인프라 투자, 상시 모니터링 체계로 통합함.')
add_caption(doc,'표 5-13. 본 장의 발견과 후속 분석 연결')
add_table(doc,['핵심 발견','후속 분석','전략 활용'],[
    ['저자 81.6%가 1편 참여','신규·반복 참여 연구자 코호트와 기술군 이동 분석','신흥 인력과 지속 연구 허브를 구분함'],['제1저자 93.9%가 1편','교신저자·책임저자·소속기관과 결합','연구 리더십 후보를 정교화함'],['기관 TOP10 비중 7.41%','기관명 표준화 후 분야별·네트워크 중심성 분석','연구 인프라·협력 거점을 선정함'],['미·독·영·프 참여 확대 큼','국가쌍·기관쌍 네트워크 분석','국제 협력 허브와 파트너를 발굴함'],['한국 두 기준 12위','국내 기술군 특화도·협력 외연 분석','선택과 집중형 R&D·인력 전략을 설계함'],
], [5.1,7.1,7.3],8.2)

add_heading(doc,'5.15 분석 한계 및 데이터 보완 과제')
add_bullet(doc,'저자명 동명이인·이명·이니셜 표기가 존재하므로 고유 저자 수와 생산성은 ORCID·소속·공동저자 기반 식별 보완이 필요함.',True)
add_bullet(doc,'제1저자는 분야별 관행에 따라 실제 연구책임자와 다를 수 있으므로 교신저자·마지막 저자·기여문 정보를 추가해야 함.')
add_bullet(doc,'기관명은 자동 추출 기반이므로 동일기관 분산과 일반명 병합 오류가 존재할 수 있어 ROR·OpenAlex 등 외부 식별자 매핑이 필요함.')
add_bullet(doc,'국가정보 미확인 논문은 제1저자 449건, 기관 국가 402건으로 국가 순위와 국제협력률이 과소추정될 가능성을 병기함.')
add_bullet(doc,'논문 수는 기술수준·영향력의 직접 지표가 아니므로 피인용, 저널게재, 특허, 연구비, 실증, 표준화, 기업투자와 결합해야 함.')
add_caption(doc,'표 5-14. 연구 주체 데이터의 한계와 보완 방향')
add_table(doc,['한계','영향','현재 통제','보완 방향'],[
    ['저자 식별','동명이인·이명으로 고유 저자와 순위가 왜곡될 수 있음','정제된 문자열 기준으로 집계함','ORCID·소속·공동저자 그래프를 활용함'],['제1저자 관행','연구 리더십이 과대·과소 평가될 수 있음','주도성의 보조 지표로 한정함','교신·마지막 저자와 기여문을 결합함'],['기관명 표준화','기관 순위·협업쌍이 분산 또는 병합될 수 있음','비기관성 값을 제외함','ROR·OpenAlex·GRID를 매핑함'],['국가 결측','국가별 건수·국제협력률이 과소 추정될 수 있음','커버리지와 미확인 건수를 병기함','소속문자열·도메인·외부기관 DB로 보완함'],['양적 지표 중심','연구 품질·기술수준을 직접 설명하지 못함','정성 정책·시장 근거를 교차함','인용·특허·과제·실증·표준 지표를 통합함'],
], [3.6,5.3,4.7,6.0],8.0)

add_heading(doc,'5.16 본 장 소결')
add_bullet(doc,'12개월 2,400건에서 고유 저자 10,514명, 제1저자 2,248명, 유효 기관 1,995개가 확인되어 개인·기관 단위의 넓은 연구 생태계가 형성됨.',True)
add_bullet(doc,'저자·제1저자·기관 HHI는 매우 낮아 특정 개인·기관 독점형이 아니며, 상위 주체는 전체 지배자보다 반복 참여·네트워크 허브 후보로 해석함.')
add_bullet(doc,'국가 단위에서는 미국·중국·독일·영국·일본 중심의 상대적 집중이 나타나 연구 인프라·공공투자·산업생태계의 영향이 큼.')
add_bullet(doc,'제1저자 국가와 기관 국가의 차이에서 미국·독일·영국·프랑스의 국제 참여 외연이 크게 나타나 제6장 협업 네트워크의 핵심 검증 대상으로 선정함.')
add_bullet(doc,'한국은 두 기준 모두 12위의 기반을 보유하나 국제 참여 확대 폭은 제한적이므로 분야 특화와 허브국 연계를 병행하는 전략이 요구됨.')
add_bullet(doc,'다음 제6장에서는 평균 저자 5.41명, 다국가 논문 10.0%, 상위 기관·국가 후보를 바탕으로 공저·다기관·국제협력 네트워크를 분석함.')

add_heading(doc,'참고자료')
refs=[
    ('[R1]','National Quantum Initiative, Annual Report FY2025','https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf'),
    ('[R2]','European Commission, Quantum Europe Strategy, 2025','https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy'),
    ('[R3]','UK Government, National Quantum Strategy','https://www.gov.uk/government/publications/national-quantum-strategy'),
    ('[R4]','Japan Cabinet Office, Integrated Innovation Strategy 2025','https://www8.cao.go.jp/cstp/tougosenryaku/togo2025_honbun_eiyaku.pdf'),
    ('[R5]','과학기술정보통신부, 제1차 양자종합계획, 2026','https://www.msit.go.kr/eng/bbs/view.do?bbsSeqNo=42&mId=4&mPid=2&nttSeqNo=1222&sCode=eng'),
    ('[R6]','McKinsey, Quantum Technology Monitor 2026: A Commercial Tipping Point','https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/mckinsey-quantum-technology-monitor-2026-a-commercial-tipping-point'),
]
for rid,title,url in refs:
    p=doc.add_paragraph(); pformat(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(.42); p.paragraph_format.first_line_indent=Cm(-.42); r=p.add_run(rid+' '); set_run(r,8.8,True,NAVY); add_hyperlink(p,title,url)

cp=doc.core_properties; cp.title='arXiv 양자 분야 제5장 연구 주체 분석'; cp.subject='2025년 6월~2026년 5월, 12개월·2,400건 기준 정량·정성 통합 분석'; cp.author='OpenAI'; cp.keywords='arXiv, 양자, 연구 주체, 저자, 기관, 국가, 정책, 인력'
settings=doc.settings._element; update=OxmlElement('w:updateFields'); update.set(qn('w:val'),'true'); settings.append(update)
doc.save(OUTPUT); print(f'Created {OUTPUT}')

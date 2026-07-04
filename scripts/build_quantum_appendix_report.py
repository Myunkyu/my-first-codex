from __future__ import annotations

from pathlib import Path
import zipfile

import numpy as np
import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR=Path('artifacts'); ASSET_DIR=OUT_DIR/'appendix_assets'; OUT_DIR.mkdir(parents=True,exist_ok=True); ASSET_DIR.mkdir(parents=True,exist_ok=True)
OUTPUT=OUT_DIR/'arXiv_양자분야_부록_다양한_분석자료_정리_고품질_20260704.docx'

RAW_TOTAL=2427; PARTIAL=27; TOTAL=2400
MONTHS=['2025-06','2025-07','2025-08','2025-09','2025-10','2025-11','2025-12','2026-01','2026-02','2026-03','2026-04','2026-05']
MONTHLY=np.array([182,187,183,219,229,171,227,186,165,211,220,220])
RAW_AUTHORS=10585; AUTHORS=10514; RAW_FIRST=2273; FIRST=2248; RAW_INST=2022; INST=1995; COUNTRIES=35
MAIN_CAT=87; ALL_CAT=117; QUANT_PH=1680; QUANT_SHARE=70.00; MULTI_CAT=1138; MULTI_SHARE=47.42
COAUTHOR=2113; COAUTHOR_SHARE=88.0; MULTI_INST=710; MULTI_INST_SHARE=29.6; INTL=241; INTL_SHARE=10.0
DOI=15.99; PDF=100.00

REFS=[
('M1','시장·산업','McKinsey','Quantum Technology Monitor 2026: A commercial tipping point','2026','300개 이상 기업 채택, 매출·투자·경제가치 전망','시장성·기업채택·하이브리드 해석','B','https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/mckinsey-quantum-technology-monitor-2026-a-commercial-tipping-point'),
('M2','시장·산업','Boston Consulting Group','Forecast for Quantum Computing Still Looks Bright','2024','장기 경제가치와 하드웨어·SW 시장 전망','장기 시장 시나리오 비교','B','https://www.bcg.com/publications/2024/long-term-forecast-for-quantum-computing-still-looks-bright'),
('M3','시장·산업','GSMA','Post-Quantum Cryptography Documents','지속 갱신','이동통신망·로밍·보안의 양자안전 전환자료','통신·보안 산업의 PQC 전환','A','https://www.gsma.com/solutions-and-impact/technologies/security/post-quantum-cryptography-documents/'),
('P1','정책','U.S. National Quantum Initiative','Annual Report FY2025','2024','예산·과학·인력·산업·인프라·안보·국제협력','미국 연구주체·장기투자 해석','A','https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf'),
('P2','정책','European Commission','Quantum Europe Strategy','2025','연구혁신·인프라·공급망·산업화·인력','유럽 협업·산업화 해석','A','https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy'),
('P3','정책','UK Government','National Quantum Strategy','2023','10년 전략과 연구·산업·인력·국제협력','영국 허브성과 산업연계','A','https://www.gov.uk/government/publications/national-quantum-strategy'),
('P4','정책','Japan Cabinet Office','Integrated Innovation Strategy 2025','2025','양자산업화·연구기관·인력·산학관 협력','일본 자체주도형 기반 해석','A','https://www8.cao.go.jp/cstp/tougosenryaku/togo2025_honbun_eiyaku.pdf'),
('P5','정책','과학기술정보통신부','제1차 양자종합계획','2026','2035년 인력·기업·표준·클러스터 목표','한국형 실행전략·인력·산업','A','https://www.msit.go.kr/eng/bbs/view.do?bbsSeqNo=42&mId=4&mPid=2&nttSeqNo=1222&sCode=eng'),
('P6','정책','The White House','Ushering in the Next Frontier of Quantum Innovation','2026','배치·상용화·평가·제조·인력 정책','연구에서 검증·배치로의 전환','A','https://www.whitehouse.gov/presidential-actions/2026/06/ushering-in-the-next-frontier-of-quantum-innovation/'),
('S1','표준화','NIST CSRC','Post-Quantum Cryptography FIPS Approved','2024','FIPS 203·204·205 승인과 전환기준','암호·보안 즉시 전환신호','A','https://csrc.nist.gov/news/2024/postquantum-cryptography-fips-approved'),
('S2','표준화','ITU-T','Y.Sup98 Technical considerations towards quantum networks','2025','양자네트워크 아키텍처·프로토콜·성능·보안','양자통신 표준화·협업','A','https://www.itu.int/rec/dologin_pub.asp?id=T-REC-Y.Sup98-202511-I!!PDF-E&lang=s&type=items'),
('S3','표준화','ETSI','Quantum-Safe Cryptography','지속 갱신','양자안전 암호와 전환·상호운용 활동','유럽 PQC 모니터링','A','https://www.etsi.org/technologies/quantum-safe-cryptography'),
('S4','표준화','IETF','Post-Quantum Use in Protocols','지속 갱신','인터넷 프로토콜의 PQC·하이브리드 논의','TLS·PKI 전환 이슈','A','https://datatracker.ietf.org/wg/pquip/about/'),
('S5','표준화','ISO/IEC JTC 1','Quantum computing standardization activities','지속 갱신','용어·성능·상호운용·보안 국제표준 기반','국제표준·시험인증 로드맵','A','https://www.iso.org/committee/45020.html'),
('T1','기술 로드맵','IBM Quantum','Hardware and roadmap','지속 갱신','모듈형 시스템·오류정정·내결함성 로드맵','컴퓨팅 상용화 시점 해석','B','https://www.ibm.com/quantum/hardware'),
('T2','기술 검증','DARPA','Quantum Benchmarking Initiative','지속 갱신','산업적으로 유용한 양자컴퓨터 독립검증','성능·경제성 평가설계','A','https://www.darpa.mil/research/programs/quantum-benchmarking-initiative'),
('T3','기술 로드맵','Quantum Flagship','Strategic Research and Industry Agenda 2030','2024','컴퓨팅·통신·센싱·기초과학·산업화 로드맵','유럽 기술축 비교','A','https://qt.eu/about-quantum-flagship/strategic-research-and-industry-agenda-2030'),
('D1','데이터 인프라','arXiv','arXiv API and metadata','지속 갱신','논문 메타데이터·카테고리·버전·저자','원천 논문 수집','A','https://info.arxiv.org/help/api/index.html'),
('D2','데이터 인프라','Crossref','REST API','지속 갱신','DOI·발행정보·참고문헌·기관 메타데이터','DOI·서지정보 보완','A','https://www.crossref.org/documentation/retrieve-metadata/rest-api/'),
('D3','데이터 인프라','OpenAlex','OpenAlex API','지속 갱신','논문·저자·기관·인용·주제 메타데이터','인용·기관·주체 검증','A','https://docs.openalex.org/'),
('D4','데이터 인프라','ROR','Research Organization Registry','지속 갱신','글로벌 연구기관 식별자·명칭·국가','기관명 통합·네트워크 재산출','A','https://ror.org/'),
('D5','데이터 인프라','ORCID','ORCID Registry and API','지속 갱신','연구자 고유식별자·소속·성과 연결','동명이인·이명 보완','A','https://info.orcid.org/documentation/features/orcid-registry/'),
]

NAVY='17365D'; BLUE='2F75B5'; LIGHT_BLUE='D9EAF7'; GREEN='548235'; LIGHT_GREEN='E2F0D9'; ORANGE='ED7D31'; LIGHT_ORANGE='FCE4D6'; PURPLE='7030A0'; LIGHT_PURPLE='E4DFEC'; GRAY='6B7280'; LIGHT_GRAY='F2F4F7'; DARK='243447'; WHITE='FFFFFF'; BORDER='CAD4DF'; YELLOW='FFD966'

def rgb(h): return RGBColor.from_string(h)
font_candidates=['/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc','/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc','/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc']
font_path=next((p for p in font_candidates if Path(p).exists()),None)
if font_path:
    fp=font_manager.FontProperties(fname=font_path); plt.rcParams['font.family']=fp.get_name()
else: plt.rcParams['font.family']='DejaVu Sans'
plt.rcParams['axes.unicode_minus']=False; plt.rcParams['font.weight']='bold'

# --- figures ---
def finish(fig,path): fig.tight_layout(pad=1.0); fig.savefig(path,dpi=220,bbox_inches='tight'); plt.close(fig)
def data_flow(path):
    fig,ax=plt.subplots(figsize=(11.2,4.6)); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1); colors=plt.rcParams['axes.prop_cycle'].by_key()['color']; boxes=[(.04,.32,.25,.36,'원수집 DB',f'{RAW_TOTAL:,}건\n2025.06~2026.06'),(.38,.32,.24,.36,'부분월 제외',f'2026.06\n{PARTIAL}건 제외'),(.71,.32,.25,.36,'분석용 DB',f'{TOTAL:,}건\n2025.06~2026.05')]
    for i,(x,y,w,h,title,body) in enumerate(boxes): ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.02,rounding_size=.025',facecolor='white',edgecolor=colors[i],linewidth=2.2)); ax.text(x+w/2,y+h*.70,title,ha='center',fontsize=13,fontweight='bold'); ax.text(x+w/2,y+h*.33,body,ha='center',fontsize=12,fontweight='bold',linespacing=1.4)
    ax.add_patch(FancyArrowPatch((.30,.50),(.37,.50),arrowstyle='-|>',mutation_scale=18,linewidth=1.6)); ax.add_patch(FancyArrowPatch((.63,.50),(.70,.50),arrowstyle='-|>',mutation_scale=18,linewidth=1.6)); ax.text(.50,.92,'원수집 데이터에서 비교 가능한 완전월 분석 DB로의 전환',ha='center',fontsize=16,fontweight='bold'); ax.text(.50,.13,'모든 수치·성장률·국가·협업 비교는 12개월·2,400건을 공통 분모로 적용함',ha='center',fontsize=10,fontweight='bold'); finish(fig,path)
def dashboard(path):
    fig,ax=plt.subplots(figsize=(12,5.6)); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1); colors=plt.rcParams['axes.prop_cycle'].by_key()['color']; cards=[(.03,.56,.22,.31,'연구 규모',f'{TOTAL:,}건\n월평균 200편',colors[0]),(.27,.56,.22,.31,'분야·융합',f'quant-ph {QUANT_SHARE:.2f}%\n복수분류 {MULTI_SHARE:.2f}%',colors[1]),(.51,.56,.22,.31,'연구 주체',f'저자 {AUTHORS:,}명\n기관 {INST:,}개',colors[2]),(.75,.56,.22,.31,'협업 구조',f'공저 {COAUTHOR_SHARE:.1f}%\n국제협력 {INTL_SHARE:.1f}%',colors[3])]
    for x,y,w,h,title,body,color in cards: ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor='white',edgecolor=color,linewidth=2.1)); ax.text(x+w/2,y+h*.70,title,ha='center',fontsize=12.2,fontweight='bold'); ax.text(x+w/2,y+h*.33,body,ha='center',fontsize=11.5,fontweight='bold',linespacing=1.5)
    ax.add_patch(FancyBboxPatch((.09,.13),.82,.25,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor='black',linewidth=2.2)); ax.text(.50,.29,'부록 활용 목적',ha='center',fontsize=11,fontweight='bold'); ax.text(.50,.22,'본문 결과의 재현·검증·업데이트를 위한 원자료·지표·출처·품질기준을 통합 관리함',ha='center',fontsize=13.2,fontweight='bold'); ax.text(.50,.16,'정량 DB와 정성 근거를 분리 보관하되 분석 단계에서 교차검증하도록 구성함',ha='center',fontsize=9.8,fontweight='bold'); ax.text(.50,.94,'본문 핵심지표와 부록의 연결',ha='center',fontsize=16,fontweight='bold'); finish(fig,path)
def source_chart(path):
    cats=['시장·산업','정책','표준화','기술 로드맵/검증','데이터 인프라']; vals=[sum(1 for r in REFS if r[1]=='시장·산업'),sum(1 for r in REFS if r[1]=='정책'),sum(1 for r in REFS if r[1]=='표준화'),sum(1 for r in REFS if r[1] in ['기술 로드맵','기술 검증']),sum(1 for r in REFS if r[1]=='데이터 인프라')]; fig,ax=plt.subplots(figsize=(10.2,5.2)); bars=ax.bar(cats,vals); ax.set_ylabel('수록 자료 수(건)',fontweight='bold'); ax.set_title('부록 참고자료 포트폴리오',fontsize=16,fontweight='bold',pad=16); ax.grid(axis='y',alpha=.25); ax.set_ylim(0,max(vals)+2); ax.tick_params(axis='x',labelrotation=0)
    for b,v in zip(bars,vals): ax.text(b.get_x()+b.get_width()/2,v+.15,str(v),ha='center',fontsize=11,fontweight='bold')
    finish(fig,path)
def evidence(path):
    items=[('정부·국제기구',4.9,4.7,4.2),('표준기관',4.8,4.9,4.5),('공식 기술 로드맵',4.0,4.3,4.6),('컨설팅 시장보고서',3.4,3.5,4.7),('서지·식별자 DB',4.5,4.8,4.8)]; fig,ax=plt.subplots(figsize=(9.4,6.2)); colors=plt.rcParams['axes.prop_cycle'].by_key()['color']
    for i,(name,a,d,t) in enumerate(items): ax.scatter(a,d,s=260+t*120,alpha=.70,edgecolors='black',linewidths=1,color=colors[i]); ax.text(a+.05,d+.05,name,fontsize=9.7,fontweight='bold')
    ax.set_xlim(3,5.2); ax.set_ylim(3,5.2); ax.set_xlabel('공신력·권위성(5점)',fontweight='bold'); ax.set_ylabel('분석목적 직접성(5점)',fontweight='bold'); ax.set_title('자료 유형별 근거 신뢰도 매트릭스',fontsize=16,fontweight='bold',pad=16); ax.grid(alpha=.25); ax.text(3.08,3.12,'※ 원 크기: 최신성·업데이트 가능성',fontsize=8.8,fontweight='bold'); finish(fig,path)
def cycle(path):
    fig,ax=plt.subplots(figsize=(11.6,5.6)); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1); colors=plt.rcParams['axes.prop_cycle'].by_key()['color']; stages=[(.05,.58,.20,.25,'월간','논문 수집·신규 토픽\n기관·저자·국가 변화'),(.29,.58,.20,.25,'분기','기술축 성장·협업\n시장·투자·기업 신호'),(.53,.58,.20,.25,'반기','특허·표준·정책\n벤치마크·실증성과'),(.77,.58,.18,.25,'연간','포트폴리오 평가\n확대·전환·중단')]
    for i,(x,y,w,h,title,body) in enumerate(stages): ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor='white',edgecolor=colors[i],linewidth=2)); ax.text(x+w/2,y+h*.70,title,ha='center',fontsize=12,fontweight='bold'); ax.text(x+w/2,y+h*.30,body,ha='center',fontsize=9.3,fontweight='bold',linespacing=1.4); 
    for i in range(3): ax.add_patch(FancyArrowPatch((stages[i][0]+stages[i][2],.705),(stages[i+1][0],.705),arrowstyle='-|>',mutation_scale=14,linewidth=1.2))
    ax.add_patch(FancyBboxPatch((.15,.13),.70,.25,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor='black',linewidth=2.1)); ax.text(.50,.29,'연속 개선 루프',ha='center',fontsize=11,fontweight='bold'); ax.text(.50,.22,'수집 → 정비 → 분석 → 정성검증 → 전략판단 → DB·분류체계 업데이트',ha='center',fontsize=12.5,fontweight='bold'); ax.text(.50,.16,'변경 이력·산식·출처·전문가 판단을 버전 단위로 보존함',ha='center',fontsize=9.5,fontweight='bold'); ax.text(.50,.94,'분석자료 상시 업데이트 주기',ha='center',fontsize=16,fontweight='bold'); finish(fig,path)
def chapter_map(path):
    fig,ax=plt.subplots(figsize=(12,6.4)); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1); colors=plt.rcParams['axes.prop_cycle'].by_key()['color']; chapters=[(.04,.68,.20,.19,'제2장','대상·방법론','검색식·기간·품질관리'),(.29,.68,.20,.19,'제3장','규모·성장','월별·반기·모멘텀'),(.54,.68,.20,.19,'제4장','분야·카테고리','코어·융합·기술축'),(.79,.68,.17,.19,'제5장','연구 주체','저자·기관·국가'),(.17,.35,.24,.19,'제6장','협업 구조','공저·다기관·국제망'),(.59,.35,.24,.19,'제7장','종합 전략','포트폴리오·로드맵')]
    for i,(x,y,w,h,ch,title,sub) in enumerate(chapters): ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor='white',edgecolor=colors[i],linewidth=2)); ax.text(x+w/2,y+h*.72,ch,ha='center',fontsize=9.5,fontweight='bold'); ax.text(x+w/2,y+h*.46,title,ha='center',fontsize=11.5,fontweight='bold'); ax.text(x+w/2,y+h*.20,sub,ha='center',fontsize=8.8,fontweight='bold')
    ax.add_patch(FancyBboxPatch((.25,.08),.50,.17,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor='black',linewidth=2.2)); ax.text(.50,.19,'부록: 원자료·산식·출처·품질관리·업데이트 템플릿',ha='center',fontsize=12.5,fontweight='bold'); ax.text(.50,.125,'각 장의 분석을 재현하고 향후 데이터 갱신 시 동일 구조로 확장함',ha='center',fontsize=9.5,fontweight='bold'); ax.text(.50,.95,'본문–부록 분석자료 연결 구조',ha='center',fontsize=16,fontweight='bold'); finish(fig,path)
FIGS={'flow':ASSET_DIR/'fig_app_1_flow.png','dash':ASSET_DIR/'fig_app_2_dashboard.png','source':ASSET_DIR/'fig_app_3_source.png','evidence':ASSET_DIR/'fig_app_4_evidence.png','cycle':ASSET_DIR/'fig_app_5_cycle.png','map':ASSET_DIR/'fig_app_6_map.png'}
data_flow(FIGS['flow']); dashboard(FIGS['dash']); source_chart(FIGS['source']); evidence(FIGS['evidence']); cycle(FIGS['cycle']); chapter_map(FIGS['map'])

# --- doc helpers ---
def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)
def border(cell,color=BORDER,size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); b=tcPr.first_child_found_in('w:tcBorders')
    if b is None: b=OxmlElement('w:tcBorders'); tcPr.append(b)
    for edge in ['top','left','bottom','right']:
        el=b.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); b.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)
def margins(cell,top=80,start=90,bottom=80,end=90):
    tcPr=cell._tc.get_or_add_tcPr(); m=tcPr.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tcPr.append(m)
    for n,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=m.find(qn(f'w:{n}'))
        if node is None: node=OxmlElement(f'w:{n}'); m.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')
def set_run(run,size=10,bold=False,color=DARK):
    run.font.name='맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=rgb(color)
def pf(p,before=0,after=4,line=1.25,keep=False):
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line; p.paragraph_format.keep_with_next=keep
def page_num(p):
    run=p.add_run(); a=OxmlElement('w:fldChar'); a.set(qn('w:fldCharType'),'begin'); b=OxmlElement('w:instrText'); b.set(qn('xml:space'),'preserve'); b.text='PAGE'; c=OxmlElement('w:fldChar'); c.set(qn('w:fldCharType'),'end'); run._r.extend([a,b,c])
def heading(doc,text,level=1):
    p=doc.add_paragraph(); pf(p,12 if level==1 else 8,6,1.0,True); r=p.add_run(text); set_run(r,15 if level==1 else 12,True,NAVY)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'12'); bot.set(qn('w:space'),'3'); bot.set(qn('w:color'),BLUE); pBdr.append(bot); pPr.append(pBdr)
def bullet(doc,text,major=False):
    p=doc.add_paragraph(); pf(p,after=3,line=1.28); p.paragraph_format.left_indent=Cm(.48); p.paragraph_format.first_line_indent=Cm(-.34); r=p.add_run('□ ' if major else '· '); set_run(r,10.2,True,BLUE if major else GRAY); r=p.add_run(text); set_run(r,10.2,major,DARK)
def caption(doc,text,fig=False):
    p=doc.add_paragraph(); pf(p,before=3,after=6,line=1.0,keep=not fig); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if fig else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); set_run(r,9.2,True,GRAY if fig else NAVY)
def table(doc,headers,rows,widths,font=8.0):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.width=Cm(widths[i]); shade(c,NAVY); border(c); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pf(p,after=0,line=1.05); r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            c=cells[i]; c.width=Cm(widths[i]); shade(c,WHITE if ri%2==0 else LIGHT_GRAY); border(c); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pf(p,after=0,line=1.08); r=p.add_run(str(v)); set_run(r,font,i==0,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def callout(doc,title,lines,fill=LIGHT_BLUE,accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,fill); border(c,accent,'8'); margins(c,130,150,130,150); p=c.paragraphs[0]; pf(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=c.add_paragraph(); pf(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.first_line_indent=Cm(-.25); r=p.add_run('• '); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
def link(p,text,url):
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rp=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rp.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rp.append(u); r.append(rp); tx=OxmlElement('w:t'); tx.text=text; r.append(tx); h.append(r); p._p.append(h)

# --- build document ---
doc=Document(); sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35); sec.header_distance=Cm(.55); sec.footer_distance=Cm(.55)
st=doc.styles['Normal']; st.font.name='맑은 고딕'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); st.font.size=Pt(10)
h=sec.header.paragraphs[0]; r=h.add_run('arXiv 양자 분야 논문 분석 | 부록. 다양한 분석 자료들 정리'); set_run(r,8.4,True,NAVY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=f.add_run('분석 기준: 2025.06~2026.05, 12개월·2,400건  |  '); set_run(r,8.2,False,GRAY); page_num(f)
cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; c=cover.cell(0,0); shade(c,NAVY); margins(c,750,300,750,300); border(c,NAVY,'0'); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('APPENDIX'); set_run(r,14,True,'9DC3E6'); p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('다양한 분석 자료들 정리'); set_run(r,27,True,WHITE); p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('원자료·지표·출처·품질관리·업데이트 체계'); set_run(r,13,True,LIGHT_BLUE); p=c.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pf(p,before=18); r=p.add_run('2025년 6월~2026년 5월 | 12개월·2,400건'); set_run(r,11,True,YELLOW)
doc.add_paragraph(); kpis=[('2,427건','원수집'),('2,400건','분석용 DB'),('117개','전체 카테고리'),(f'{len(REFS)}건','핵심 참고자료'),('A등급 중심','근거 신뢰도')]; kt=doc.add_table(rows=1,cols=5); kt.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(val,lab) in enumerate(kpis):
    cc=kt.cell(0,i); shade(cc,[LIGHT_BLUE,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]); border(cc,[BLUE,GREEN,ORANGE,PURPLE,GRAY][i],'7'); margins(cc,100,70,100,70); p=cc.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,14,True,NAVY); p=cc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(lab); set_run(r,8.5,True,GRAY)
callout(doc,'부록의 역할',['본문 제1~7장의 정량분석 결과를 재현·검증할 수 있도록 분석범위, 핵심 원자료, 지표 정의와 산식을 정리함.','시장·기술·정책·표준화 자료를 분석목적별로 분류하고 공신력·최신성·직접성을 기준으로 활용수준을 제시함.','향후 데이터 갱신 시 동일한 구조로 재분석할 수 있도록 품질점검표, 업데이트 주기, 산출물 템플릿을 제공함.']); doc.add_page_break()

heading(doc,'부록 A. 분석자료 구성 및 본문 연결'); doc.add_picture(str(FIGS['map']),width=Cm(18)); caption(doc,'그림 A-1. 본문–부록 분석자료 연결 구조',True); bullet(doc,'부록은 본문에서 사용한 핵심 원자료와 산식뿐 아니라 정성분석의 근거자료와 품질관리 기준까지 통합함.',True); bullet(doc,'각 표와 그림은 본문 장과 직접 연결되도록 구성하여 최종 통합보고서 결합 시 중복을 최소화함.'); bullet(doc,'정량근거와 정성근거는 별도 관리하되 종합 판단에서는 근거 유형과 적용범위를 명시하여 교차검증함.')
caption(doc,'표 A-1. 본문 장별 부록 자료 매핑'); table(doc,['본문 장','핵심 분석','부록 보완자료','재활용 목적'],[['제1장','Executive Summary','핵심지표·전략메시지·근거자료 목록','경영진·정책결정자용 요약 갱신'],['제2장','대상·방법론','검색식·기간·필드·제외기준·품질관리','동일 조건 재수집·재현'],['제3장','규모·성장','월별 원자료·산식·변동지표','시계열 갱신·이상치 검증'],['제4장','분야·카테고리','카테고리 국문명·코드·재분류·융합지표','기술분류체계 확장'],['제5장','연구 주체','저자·기관·국가 기준·표준화 규칙','주체 순위·네트워크 재산출'],['제6장','협업 구조','공저·기관·국가 협업 판정·네트워크 지표','허브·브리지·지속협력 분석'],['제7장','종합 전략','시장·정책·표준·기술 로드맵과 KPI','포트폴리오·로드맵 재평가']],[3,4.5,7.4,5.7],7.9)

heading(doc,'부록 B. 분석대상 데이터와 핵심 원자료'); doc.add_picture(str(FIGS['flow']),width=Cm(17.8)); caption(doc,'그림 B-1. 원수집 DB에서 12개월 분석용 DB로의 전환',True)
caption(doc,'표 B-1. 원수집 DB와 분석용 DB 비교'); table(doc,['지표','원수집 DB','분석용 DB','차이·처리기준'],[['분석기간','2025.06~2026.06','2025.06~2026.05','2026년 6월 부분월 제외'],['논문 수',f'{RAW_TOTAL:,}건',f'{TOTAL:,}건',f'{PARTIAL}건 제외'],['고유 저자',f'{RAW_AUTHORS:,}명',f'{AUTHORS:,}명','부분월 저자 제외 후 재산출'],['고유 제1저자',f'{RAW_FIRST:,}명',f'{FIRST:,}명','동일 기준 재산출'],['기관 수',f'{RAW_INST:,}개',f'{INST:,}개','비기관성 문자열 제외'],['주 카테고리',f'{MAIN_CAT}개',f'{MAIN_CAT}개','체계 동일'],['전체 카테고리',f'{ALL_CAT}개',f'{ALL_CAT}개','복수분류 포함'],['DOI 커버리지',f'{DOI:.2f}%','원수집 기준 참고','최종 출판본 연결 보완 필요'],['PDF 확보율',f'{PDF:.2f}%','동일','전문 분석 기반 확보']],[3.5,4.2,4.2,8],8)
caption(doc,'표 B-2. 월별 논문 수 원자료'); rows=[]
for i,(m,n) in enumerate(zip(MONTHS,MONTHLY),1): rows.append([i,m,int(n),'' if i==1 else f'{(n-MONTHLY[i-2])/MONTHLY[i-2]*100:+.1f}%','완전월'])
table(doc,['순번','월','논문 수','전월 대비','데이터 상태'],rows,[1.5,3.5,3,3.4,7],8.2)
doc.add_picture(str(FIGS['dash']),width=Cm(18)); caption(doc,'그림 B-2. 본문 핵심지표와 부록의 연결',True)
caption(doc,'표 B-3. 분야·연구주체·협업 핵심 원자료'); table(doc,['영역','지표','값','유의사항'],[['분야','양자물리(quant-ph)',f'{QUANT_PH:,}건·{QUANT_SHARE:.2f}%','arXiv 주 카테고리 기준'],['분야','복수 카테고리',f'{MULTI_CAT:,}건·{MULTI_SHARE:.2f}%','카테고리 수와 연구주제 수는 다름'],['연구주체','고유 저자',f'{AUTHORS:,}명','동명이인·이명 정비 필요'],['연구주체','고유 제1저자',f'{FIRST:,}명','저자순서 관행 고려'],['연구주체','유효 기관',f'{INST:,}개','ROR 표준화 전 탐색값'],['연구주체','국가',f'{COUNTRIES}개국','국가정보 결측 고려'],['협업','공동저자',f'{COAUTHOR:,}건·{COAUTHOR_SHARE:.1f}%','저자 2명 이상'],['협업','다기관',f'{MULTI_INST:,}건·{MULTI_INST_SHARE:.1f}%','정제 기관 2개 이상'],['협업','국제협력',f'{INTL:,}건·{INTL_SHARE:.1f}%','기관 국가 2개 이상']],[3,4.2,4,8.7],8)

heading(doc,'부록 C. 검색식·데이터 필드·분석 기준'); search='((((ti:QUANTUM OR ti:QUANTOMETER) OR ti:QUANTIZATION) AND (((((((((((((((ti:SENSOR OR ti:SENSING) OR ti:SENSE) OR ti:DETECT) OR ti:MEASUREMENT) OR ti:INERTIA) OR ti:TIME) OR ti:FREQUENCY) OR ti:"MAGNETIC FIELD") OR ti:"ELECTRIC FIELD") OR ti:"LIGHT BASED") OR ti:COMPUTER) OR ti:COMPUTING) OR ti:COMPUTATION) OR ti:COMPUTATIONAL) OR ti:COMMUNICATION))) AND (submittedDate:[202506060000 TO 202606062359])'
caption(doc,'표 C-1. arXiv 수집 검색식'); table(doc,['구분','내용'],[['검색식',search],['검색필드','제목(ti)과 제출일(submittedDate) 기준'],['검색대상','양자와 센싱·검출·계측·컴퓨팅·통신 용어 동시 포함'],['장점','기술연관성이 높은 최신 연구를 신속 수집'],['한계','QUANTIZATION에 따른 AI 모델 양자화 포함과 제목 미표기 핵심논문 누락 가능'],['보완','초록·전문 정합성 분류, 제외키워드, 기술축별 검색식 추가']],[3.5,16.4],7.8)
caption(doc,'표 C-2. 권장 데이터 필드와 활용'); table(doc,['필드군','핵심 필드','활용'],[['서지정보','arXiv ID, 제목, 초록, 제출일, 수정일, DOI, 저널','중복제거·시계열·출판전환'],['분류정보','주/전체 카테고리, 실무 기술축','분야 규모·융합·성장'],['저자정보','저자, 제1저자, 교신저자, ORCID','생산성·리더십·인력 이동'],['기관정보','원문/표준기관명, ROR, 국가','기관 순위·협업·국가'],['협업정보','저자 수, 기관 수, 국가 수, 협업쌍','공저·다기관·국제망'],['텍스트정보','키워드, 명사구, 임베딩, 토픽','부상기술·정성분석'],['품질정보','결측·정합성·표준화 상태·검증일','품질관리·변경이력']],[3.3,7.5,9.1],7.8)

heading(doc,'부록 D. 지표 정의·계산식·해석 가이드'); caption(doc,'표 D-1. 연구 규모·성장 지표'); table(doc,['지표','정의·계산식','해석','주의사항'],[['월별 논문 수','해당 월 제출 건수','연구활동 절대 규모','부분월·수집지연 제외'],['전월 대비 증가율','(당월-전월)/전월×100','단기 변화 방향','기저효과·계절성 고려'],['전·후반기 성장률','(후반6개월-전반6개월)/전반×100','기간 내 구조적 확대','장기추세 대체 불가'],['최근 3개월 모멘텀','최근3개월 평균/전체 월평균-1','최근 가속도','이벤트성 급증 점검'],['변동계수(CV)','표준편차/평균×100','활동량 안정성','평균 작은 집단에 부적합']],[3.6,6,5.8,5.6],7.7)
caption(doc,'표 D-2. 분야·집중·융합 지표'); table(doc,['지표','정의·계산식','해석','주의사항'],[['카테고리 점유율','카테고리 논문/전체×100','상대 연구규모','복수분류 합계 100% 초과 가능'],['복수 카테고리 비율','2개 이상 카테고리 논문/전체×100','다학제 융합 외형','실질 융합 깊이와 다름'],['교차강도','두 카테고리 동시출현/기준 카테고리','연결 강도','표본규모 민감'],['HHI','점유율 제곱합','집중도','집계단위·분모 명시'],['전문화지수','주체 특정분야 비중/전체 평균','상대 기술특화','소규모 주체 변동성 큼']],[3.6,6,5.8,5.6],7.7)
caption(doc,'표 D-3. 연구 주체·협업 지표'); table(doc,['지표','정의·계산식','해석','주의사항'],[['저자 생산성','저자별 논문 출현 수','반복 참여·활동성','동명이인·이명 정비'],['제1저자 국가','제1저자 소속국가 1개 집계','연구 주도성','저자순서 관행 고려'],['기관 국가','논문 내 고유 국가 복수 집계','참여 외연','합계가 논문 수 초과 가능'],['공동저자 비율','저자 2명 이상/전체×100','팀 연구 보편성','기여도 차이 미반영'],['다기관 비율','정제 기관 2개 이상/전체×100','기관경계 협업','기관표준화 민감'],['국제협력 비율','기관국가 2개 이상/전체×100','국제 개방성','결측 시 과소추정'],['네트워크 중심성','degree·betweenness·eigenvector','허브·브리지','단순 논문 수와 구분']],[3.5,6.3,5.8,5.4],7.6)

heading(doc,'부록 E. 시장·정책·표준·기술 로드맵 참고자료'); doc.add_picture(str(FIGS['source']),width=Cm(16.5)); caption(doc,'그림 E-1. 부록 참고자료 포트폴리오',True); doc.add_picture(str(FIGS['evidence']),width=Cm(15.4)); caption(doc,'그림 E-2. 자료 유형별 근거 신뢰도 매트릭스',True); bullet(doc,'공식 정부·국제기구·표준기관 자료를 1차 근거로 사용하고 시장 전망은 복수 자료를 비교하여 활용함.',True); bullet(doc,'기업 로드맵은 기술목표와 시간표를 파악하는 데 활용하되 독립 검증자료와 구분함.'); bullet(doc,'A등급은 정책·표준·공식 DB 등 직접성과 권위성이 높은 자료, B등급은 시장·기업 전망 등 보조자료로 정의함.')
for cat,title in [('시장·산업','표 E-1. 시장·산업 참고자료'),('정책','표 E-2. 주요국 정책 참고자료'),('표준화','표 E-3. 표준화 참고자료'),('기술','표 E-4. 기술 로드맵·검증 참고자료'),('데이터','표 E-5. 데이터·식별자 인프라')]:
    subset=[r for r in REFS if (r[1] in ['기술 로드맵','기술 검증'] if cat=='기술' else r[1]=='데이터 인프라' if cat=='데이터' else r[1]==cat)]; caption(doc,title); table(doc,['ID','기관','자료명','연도','핵심 내용','보고서 활용','등급'],[[r[0],r[2],r[3],r[4],r[5],r[6],r[7]] for r in subset],[1.2,3.4,5.2,1.8,6,5.5,1.3],7.1)
heading(doc,'부록 E-6. 참고자료 링크 목록',2)
for r in REFS:
    p=doc.add_paragraph(); pf(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(.45); p.paragraph_format.first_line_indent=Cm(-.45); rr=p.add_run(f'[{r[0]}] '); set_run(rr,8.8,True,NAVY); link(p,f'{r[2]}. {r[3]} ({r[4]})',r[8])

heading(doc,'부록 F. 분석자료 품질관리 및 검증 체크리스트'); caption(doc,'표 F-1. 데이터 품질 점검표'); table(doc,['점검영역','점검 항목','판정 기준','현재 상태','우선순위'],[['수집 완전성','기간·페이지·중복·버전','기간 내 누락·중복·최신버전','부분월 제외 완료','상'],['검색 정합성','양자 핵심/비핵심','초록·전문 정합성 검증률','제목 검색식 기준','최상'],['서지정보','DOI·저널·발행상태','DOI·출판정보 연결률',f'DOI {DOI:.2f}%','상'],['저자 식별','동명이인·이명·ORCID','식별자 매핑률·오류율','문자열 기준 집계','상'],['기관 표준화','약어·부속기관·일반명','ROR 매핑률·중복률','표준화 전 탐색값','최상'],['국가정보','제1저자·기관국가 결측','국가 확인률·결측편향','결측 존재','상'],['분류 정확도','학술분류·실무기술축','전문가 일치율·다중분류 정합성','학술분류 중심','최상'],['협업 네트워크','허위쌍·복수쌍·중심성','기관표준화 후 재산출','탐색 네트워크','최상'],['정성근거','출처·연도·직접성','A등급 비중·갱신일','공식자료 중심','중'],['재현성','산식·코드·버전·변경이력','동일 입력 시 동일 결과','분석기 기반','상']],[3.2,5.5,5.4,4.3,3],7.6)
caption(doc,'표 F-2. 오류 발견 시 변경관리 기준'); table(doc,['오류 유형','수정 기준','영향범위 확인','변경기록'],[['기간·분모 오류','부분월·중복·제외기준 재적용','모든 비율·성장률·순위 재산출','버전·수정일·사유 기록'],['기관·저자 오류','원문·식별자·소속·공동저자 대조','주체 순위·국가·협업망 재산출','전·후 매핑표 보존'],['카테고리 오류','원문 카테고리와 기술정의서 재검토','분야규모·융합·전략 검토','분류근거·전문가 판단 기록'],['외부자료 갱신','공식 최신판·게시일·개정내용 확인','시장·정책·표준 해석 수정','이전 자료와 차이 요약'],['그래프·표 불일치','원데이터·필터·단위·반올림 대조','본문·Executive Summary 동시 수정','검수자·검수일 기록']],[4,6,6.2,5.2],7.8)

heading(doc,'부록 G. 상시 모니터링 및 업데이트 체계'); doc.add_picture(str(FIGS['cycle']),width=Cm(18)); caption(doc,'그림 G-1. 분석자료 상시 업데이트 주기',True); caption(doc,'표 G-1. 업데이트 주기와 산출물'); table(doc,['주기','업데이트 대상','핵심 산출물','의사결정'],[['월간','arXiv 신규논문·저자·기관·카테고리','월별 규모·신규 토픽·신규 주체·이상치','부상신호·DB 오류 조기탐지'],['분기','기술축 성장·협업·특허·기업·투자','기술축 스코어카드·허브 변화·시장신호','심층분석 후보·과제 조정'],['반기','정책·표준·로드맵·실증성과','정성근거 업데이트·표준화·실증 브리프','협력·표준·실증 우선순위'],['연간','전체 DB·분류체계·성과·시장전망','통합보고서·포트폴리오·로드맵','투자 확대·전환·중단']],[2.8,7,6.8,5],7.8)
caption(doc,'표 G-2. 버전·변경이력 관리 필드'); table(doc,['필드','작성 내용','예시'],[['분석 버전','연도·월·빌드번호','2026.07-v1.0'],['데이터 기준일','수집 종료일과 분석기간','수집 2026-06-06 / 분석 2025.06~2026.05'],['검색식 버전','검색식 해시·변경사항','QNT-TITLE-v1.2'],['분류체계 버전','대·중·소분류 버전과 변경내역','Quantum-Tech-v2.0'],['표준화 버전','저자·기관·국가 매핑 버전','ROR-map-2026Q3'],['외부자료 기준일','시장·정책·표준 최종 확인일','2026-07-04'],['검수자·검수일','정량·정성·편집 검수','분석팀 / 2026-07-04'],['변경 사유','수치·분류·해석 변경 이유','부분월 제외·기관 통합·최신 정책 반영']],[3.6,8.4,7.6],7.8)

heading(doc,'부록 H. 재사용 가능한 분석 산출물 템플릿'); caption(doc,'표 H-1. 산업·기술군 확장용 모듈 구조'); table(doc,['모듈','입력','처리','출력'],[['수집 모듈','검색식·기간·API·파일','수집·버전통합·중복제거','원수집 DB·로그'],['정비 모듈','저자·기관·국가·DOI','식별자 매핑·표준화·결측관리','정제 DB·매핑표'],['분류 모듈','카테고리·초록·전문·정의서','규칙·임베딩·전문가 검증','대–중–소 분류·신뢰도'],['정량 모듈','정제 DB·기간·기술축','규모·성장·점유·집중·협업','표·그래프·지표 시트'],['정성 모듈','시장·정책·표준·로드맵','출처평가·기술축 매핑·교차검증','정성 브리프·근거표'],['전략 모듈','정량·정성·평가기준','스코어링·포트폴리오·리스크','우선순위·로드맵·KPI'],['보고 모듈','표·그래프·문장 템플릿','Word·Excel·PPT 자동화','통합보고서·부록·대시보드']],[3,5.4,6.2,6.6],7.8)
caption(doc,'표 H-2. 최종 통합보고서 결합 체크리스트'); table(doc,['체크 항목','완료 기준'],[['장별 분모 일치','모든 수치 분석에 12개월·2,400건을 적용함'],['부분월 설명','방법론·본문·주석·Executive Summary에 2026년 6월 제외를 명시함'],['표·그림 번호','표는 위쪽, 그림은 아래쪽에 장별 일련번호를 적용함'],['분류 명칭','국문명을 메인으로 하고 영문 코드·명칭을 괄호 병기함'],['저자 표기','그래프의 저자명에 국가를 괄호로 병기함'],['국가 기준','제1저자 국가와 전체 기관 국가의 의미와 분모를 구분함'],['정량–정성 연결','각 핵심 수치에 시장·기술·정책·표준 근거를 연결함'],['데이터 한계','DOI·기관·국가·검색식·arXiv 편향을 명시함'],['참고자료','출처 ID·기관·연도·URL·활용내용을 일관되게 관리함'],['변경이력','수정 수치·분류·해석과 사유를 버전별로 보존함']],[5,15],8.1)
heading(doc,'부록 소결'); bullet(doc,'본 부록은 단순 자료집이 아니라 원자료·산식·정성근거·품질관리·업데이트 체계를 통합한 재현성 관리체계로 구성함.',True); bullet(doc,'원수집 2,427건과 분석용 2,400건의 범위를 명확히 구분하여 모든 장의 정량분모를 일치시킴.'); bullet(doc,'시장·정책·표준·기술 로드맵은 공식자료를 우선하고 시장·기업 자료는 시나리오와 보조근거로 제한하여 활용함.'); bullet(doc,'기관·저자·기술분류 표준화를 최우선 데이터 보완과제로 설정하고 월간–분기–반기–연간 갱신체계를 제시함.'); bullet(doc,'향후 다른 산업·기술군 분석에서도 동일 모듈을 재사용할 수 있도록 입력–처리–출력 구조와 통합보고서 체크리스트를 제공함.')
cp=doc.core_properties; cp.title='arXiv 양자 분야 부록: 다양한 분석 자료들 정리'; cp.subject='원자료·지표·출처·품질관리·상시 업데이트 체계'; cp.author='OpenAI'; cp.keywords='arXiv, 양자, 부록, 원자료, 지표, 시장, 정책, 표준화, 품질관리'
settings=doc.settings._element; upd=OxmlElement('w:updateFields'); upd.set(qn('w:val'),'true'); settings.append(upd)
doc.save(OUTPUT); print(f'Created {OUTPUT}'); print({'valid':zipfile.is_zipfile(OUTPUT),'size_mb':round(OUTPUT.stat().st_size/1024/1024,2),'figures':len(list(ASSET_DIR.glob('*.png'))),'tables':len(doc.tables)})

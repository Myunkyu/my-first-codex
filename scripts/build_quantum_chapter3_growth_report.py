from __future__ import annotations

from pathlib import Path
import math
import statistics

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path("artifacts")
ASSET_DIR = OUT_DIR / "chapter3_assets"
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / "arXiv_양자분야_제3장_연구규모_성장분석_고품질_재작성_20260704.docx"

MONTHS = ["2025-06","2025-07","2025-08","2025-09","2025-10","2025-11","2025-12","2026-01","2026-02","2026-03","2026-04","2026-05"]
COUNTS = np.array([182,187,183,219,229,171,227,186,165,211,220,220], dtype=float)

TOTAL = int(COUNTS.sum())
MEAN = float(COUNTS.mean())
MEDIAN = float(np.median(COUNTS))
STD = float(COUNTS.std(ddof=0))
CV = STD / MEAN * 100
RANGE = int(COUNTS.max() - COUNTS.min())
MOM = np.r_[np.nan, (COUNTS[1:] - COUNTS[:-1]) / COUNTS[:-1] * 100]
DIFF = np.r_[np.nan, COUNTS[1:] - COUNTS[:-1]]
ROLL3 = np.array([np.nan, np.nan] + [COUNTS[i-2:i+1].mean() for i in range(2, len(COUNTS))])
FIRST6 = int(COUNTS[:6].sum())
SECOND6 = int(COUNTS[6:].sum())
HALF_GROWTH = (SECOND6 - FIRST6) / FIRST6 * 100
PRIOR3 = int(COUNTS[6:9].sum())
RECENT3 = int(COUNTS[9:12].sum())
MOMENTUM = (RECENT3 - PRIOR3) / PRIOR3 * 100
HIGH_MONTHS = int((COUNTS >= MEAN).sum())
HIGH_TOTAL = int(COUNTS[COUNTS >= MEAN].sum())
HIGH_SHARE = HIGH_TOTAL / TOTAL * 100
x = np.arange(1, 13, dtype=float)
slope, intercept = np.polyfit(x, COUNTS, 1)
pred = slope*x + intercept
ss_res = float(((COUNTS-pred)**2).sum())
ss_tot = float(((COUNTS-MEAN)**2).sum())
r2 = 1 - ss_res/ss_tot

NAVY = "17365D"
BLUE = "2F75B5"
SKY = "D9EAF7"
GREEN = "548235"
LIGHT_GREEN = "E2F0D9"
ORANGE = "ED7D31"
LIGHT_ORANGE = "FCE4D6"
RED = "C00000"
LIGHT_RED = "F4CCCC"
PURPLE = "7030A0"
LIGHT_PURPLE = "E4DFEC"
GRAY = "6B7280"
LIGHT_GRAY = "F2F4F7"
DARK = "243447"
WHITE = "FFFFFF"
BORDER = "CAD4DF"
YELLOW = "FFD966"

# ------------------------ font and chart defaults ------------------------
font_candidates = [
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc",
]
font_path = next((p for p in font_candidates if Path(p).exists()), None)
if font_path:
    prop = font_manager.FontProperties(fname=font_path)
    plt.rcParams["font.family"] = prop.get_name()
else:
    plt.rcParams["font.family"] = "DejaVu Sans"
plt.rcParams["axes.unicode_minus"] = False
plt.rcParams["font.weight"] = "bold"


def c(h: str) -> str:
    return "#" + h


def rgb(h: str) -> RGBColor:
    return RGBColor.from_string(h)

# ------------------------ chart creation ------------------------
def base_axes(figsize=(11.6, 5.0)):
    fig, ax = plt.subplots(figsize=figsize, dpi=220)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    for s in ["top", "right"]:
        ax.spines[s].set_visible(False)
    ax.spines["left"].set_color(c(BORDER))
    ax.spines["bottom"].set_color(c(BORDER))
    ax.tick_params(colors=c(DARK), labelsize=9)
    ax.grid(axis="y", color=c(BORDER), linewidth=0.7, alpha=0.55)
    ax.set_axisbelow(True)
    return fig, ax


def save_monthly_trend(path: Path):
    fig, ax = base_axes((12.0, 5.2))
    idx = np.arange(len(MONTHS))
    colors = [c(BLUE) if v >= MEAN else c(SKY) for v in COUNTS]
    bars = ax.bar(idx, COUNTS, color=colors, edgecolor=c(BLUE), linewidth=0.8, width=0.66, zorder=3)
    ax.plot(idx, ROLL3, color=c(ORANGE), linewidth=2.8, marker="o", markersize=5.5, label="3개월 이동평균", zorder=4)
    ax.axhline(MEAN, color=c(GREEN), linewidth=2.0, linestyle="--", label=f"12개월 평균 {MEAN:.1f}편")
    ax.plot(idx, pred, color=c(PURPLE), linewidth=1.8, linestyle=":", label=f"선형 추세 +{slope:.1f}편/월")
    for b, v in zip(bars, COUNTS):
        ax.text(b.get_x()+b.get_width()/2, v+4.0, f"{int(v)}", ha="center", va="bottom", fontsize=9.3, fontweight="bold", color=c(NAVY), clip_on=False)
    ax.set_xticks(idx)
    ax.set_xticklabels([m.replace("2025-", "25.").replace("2026-", "26.") for m in MONTHS], fontweight="bold")
    ax.set_ylim(145, 250)
    ax.set_ylabel("논문 수(편)", fontsize=10.5, fontweight="bold", color=c(NAVY))
    ax.set_title("월별 논문 제출량과 3개월 이동평균", fontsize=16, fontweight="bold", color=c(NAVY), pad=18)
    ax.legend(loc="upper left", ncol=3, frameon=False, fontsize=9.3)
    ax.annotate("1차 고점", xy=(4,229), xytext=(3.0,245), arrowprops=dict(arrowstyle="->", color=c(ORANGE), lw=1.5), color=c(ORANGE), fontweight="bold", fontsize=9.5)
    ax.annotate("최근 고점 유지", xy=(11,220), xytext=(9.1,244), arrowprops=dict(arrowstyle="->", color=c(GREEN), lw=1.5), color=c(GREEN), fontweight="bold", fontsize=9.5)
    fig.tight_layout(pad=1.0)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_mom_change(path: Path):
    fig, ax = base_axes((12.0, 4.8))
    idx = np.arange(1, len(MONTHS))
    diffs = DIFF[1:]
    cols = [c(GREEN) if v > 0 else c(RED) if v < 0 else c(GRAY) for v in diffs]
    bars = ax.bar(idx, diffs, color=cols, width=0.66, zorder=3)
    ax.axhline(0, color=c(NAVY), linewidth=1.2)
    for b, d, p in zip(bars, diffs, MOM[1:]):
        offset = 3.2 if d >= 0 else -4.0
        va = "bottom" if d >= 0 else "top"
        ax.text(b.get_x()+b.get_width()/2, d+offset, f"{int(d):+d}편\n({p:+.1f}%)", ha="center", va=va, fontsize=8.8, fontweight="bold", color=c(DARK), clip_on=False)
    ax.set_xticks(idx)
    ax.set_xticklabels([m.replace("2025-", "25.").replace("2026-", "26.") for m in MONTHS[1:]], fontweight="bold")
    ax.set_ylim(-72, 70)
    ax.set_ylabel("전월 대비 증감(편)", fontsize=10.5, fontweight="bold", color=c(NAVY))
    ax.set_title("월별 증감과 조정·반등 구간", fontsize=16, fontweight="bold", color=c(NAVY), pad=18)
    fig.tight_layout(pad=1.0)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_rolling(path: Path):
    fig, ax = base_axes((12.0, 4.6))
    idx = np.arange(2, len(MONTHS))
    vals = ROLL3[2:]
    ax.plot(idx, vals, color=c(BLUE), linewidth=3.0, marker="o", markersize=7, zorder=3)
    ax.fill_between(idx, vals, MEAN, where=vals>=MEAN, color=c(LIGHT_GREEN), alpha=0.75)
    ax.fill_between(idx, vals, MEAN, where=vals<MEAN, color=c(LIGHT_ORANGE), alpha=0.75)
    ax.axhline(MEAN, color=c(GREEN), linewidth=1.8, linestyle="--", label=f"평균 {MEAN:.1f}편")
    for i, v in zip(idx, vals):
        ax.text(i, v+2.5, f"{v:.1f}", ha="center", va="bottom", fontsize=9.2, fontweight="bold", color=c(NAVY))
    ax.set_xticks(idx)
    ax.set_xticklabels([m.replace("2025-", "25.").replace("2026-", "26.") for m in MONTHS[2:]], fontweight="bold")
    ax.set_ylim(176, 224)
    ax.set_ylabel("3개월 평균(편)", fontsize=10.5, fontweight="bold", color=c(NAVY))
    ax.set_title("3개월 이동평균으로 본 연구활동 모멘텀", fontsize=16, fontweight="bold", color=c(NAVY), pad=18)
    ax.legend(frameon=False, loc="upper left")
    ax.annotate("저점 통과", xy=(7,187.3), xytext=(5.8,180), arrowprops=dict(arrowstyle="->", color=c(RED), lw=1.5), color=c(RED), fontweight="bold")
    ax.annotate("최근 최고", xy=(11,217.0), xytext=(9.3,221), arrowprops=dict(arrowstyle="->", color=c(GREEN), lw=1.5), color=c(GREEN), fontweight="bold")
    fig.tight_layout(pad=1.0)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_period_comparison(path: Path):
    fig, ax = base_axes((10.8, 4.8))
    labels = ["전반 6개월\n25.06~25.11", "후반 6개월\n25.12~26.05", "직전 3개월\n25.12~26.02", "최근 3개월\n26.03~26.05"]
    values = [FIRST6/6, SECOND6/6, PRIOR3/3, RECENT3/3]
    cols = [c(SKY), c(BLUE), c(LIGHT_ORANGE), c(GREEN)]
    bars = ax.bar(np.arange(4), values, color=cols, edgecolor=[c(BLUE),c(BLUE),c(ORANGE),c(GREEN)], linewidth=1.2, width=0.62, zorder=3)
    ax.axhline(MEAN, color=c(PURPLE), linestyle="--", linewidth=1.8, label=f"12개월 평균 {MEAN:.1f}편")
    for b, v in zip(bars, values):
        ax.text(b.get_x()+b.get_width()/2, v+3, f"{v:.1f}편", ha="center", fontsize=11, fontweight="bold", color=c(NAVY))
    ax.text(0.5, 226, f"전·후반 +{HALF_GROWTH:.1f}%", ha="center", fontsize=10.5, fontweight="bold", color=c(BLUE))
    ax.text(2.5, 226, f"최근 모멘텀 +{MOMENTUM:.1f}%", ha="center", fontsize=10.5, fontweight="bold", color=c(GREEN))
    ax.set_xticks(np.arange(4))
    ax.set_xticklabels(labels, fontweight="bold")
    ax.set_ylim(160, 235)
    ax.set_ylabel("월평균 논문 수(편)", fontsize=10.5, fontweight="bold", color=c(NAVY))
    ax.set_title("전·후반기 및 최근 3개월 모멘텀 비교", fontsize=16, fontweight="bold", color=c(NAVY), pad=18)
    ax.legend(frameon=False, loc="upper left")
    fig.tight_layout(pad=1.0)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_activity_matrix(path: Path):
    fig, ax = base_axes((11.2, 5.6))
    y = np.nan_to_num(DIFF, nan=0.0)
    ax.axvline(MEAN, color=c(GRAY), linestyle="--", linewidth=1.2)
    ax.axhline(0, color=c(GRAY), linestyle="--", linewidth=1.2)
    cols = [c(BLUE) if i < 6 else c(GREEN) for i in range(12)]
    ax.scatter(COUNTS, y, s=135, c=cols, edgecolors=c(NAVY), linewidths=1.1, zorder=3)
    offsets = [(4,5),(4,5),(-26,-14),(4,5),(4,5),(4,-16),(4,5),(4,-16),(4,-16),(4,5),(4,5),(4,-16)]
    for i,(xx,yy) in enumerate(zip(COUNTS,y)):
        ax.annotate(MONTHS[i].replace("20", ""), (xx,yy), xytext=offsets[i], textcoords="offset points", fontsize=8.4, fontweight="bold", color=c(DARK))
    ax.text(226, 54, "고규모·확장", ha="center", va="center", fontsize=11, fontweight="bold", color=c(GREEN))
    ax.text(176, 54, "저규모·회복", ha="center", va="center", fontsize=11, fontweight="bold", color=c(ORANGE))
    ax.text(226, -50, "고규모·조정", ha="center", va="center", fontsize=11, fontweight="bold", color=c(PURPLE))
    ax.text(176, -50, "저규모·조정", ha="center", va="center", fontsize=11, fontweight="bold", color=c(RED))
    ax.set_xlim(155, 239)
    ax.set_ylim(-68, 68)
    ax.set_xlabel("월별 논문 수(편)", fontsize=10.5, fontweight="bold", color=c(NAVY))
    ax.set_ylabel("전월 대비 증감(편)", fontsize=10.5, fontweight="bold", color=c(NAVY))
    ax.set_title("연구활동 규모–증감 포지션 맵", fontsize=16, fontweight="bold", color=c(NAVY), pad=18)
    fig.tight_layout(pad=1.0)
    fig.savefig(path, bbox_inches="tight")
    plt.close(fig)


def save_external_drivers(path: Path):
    fig, ax = plt.subplots(figsize=(11.8,5.2), dpi=220)
    ax.axis("off")
    ax.set_xlim(0,1); ax.set_ylim(0,1)
    boxes = [
        (0.03,0.58,0.22,0.28,SKY,BLUE,"시장·산업",["2025년 민간투자 확대",">300개 기업의 채택·협업", "하이브리드 활용 확대"]),
        (0.28,0.58,0.22,0.28,LIGHT_GREEN,GREEN,"기술 성숙",["오류정정 below-threshold", "FTQC 로드맵 구체화", "독립 벤치마킹 강화"]),
        (0.53,0.58,0.22,0.28,LIGHT_ORANGE,ORANGE,"정책·재정",["미국 NQI 지속 투자", "EU 5대 전략축", "영국 10년 전략"]),
        (0.78,0.58,0.19,0.28,LIGHT_PURPLE,PURPLE,"후속 확인",["2026년 美 EO 14413", "상용화·제조·배치", "연구량 지속성 재확인"]),
    ]
    for x0,y0,w,h,fc,ec,title,lines in boxes:
        ax.add_patch(FancyBboxPatch((x0,y0),w,h,boxstyle="round,pad=0.015,rounding_size=0.025",facecolor=c(fc),edgecolor=c(ec),linewidth=2))
        ax.text(x0+w/2,y0+h-0.055,title,ha="center",va="center",fontsize=12.5,fontweight="bold",color=c(NAVY))
        ax.text(x0+w/2,y0+0.105,"\n".join(lines),ha="center",va="center",fontsize=9.6,fontweight="bold",color=c(DARK),linespacing=1.35)
    ax.add_patch(FancyBboxPatch((0.12,0.13),0.76,0.25,boxstyle="round,pad=0.02,rounding_size=0.03",facecolor="white",edgecolor=c(NAVY),linewidth=2.4))
    ax.text(0.50,0.29,"정량 신호",ha="center",va="center",fontsize=12,fontweight="bold",color=c(GRAY))
    ax.text(0.50,0.22,"월평균 200.0편  ·  후반기 +5.0%  ·  최근 3개월 217.0편",ha="center",va="center",fontsize=15,fontweight="bold",color=c(NAVY))
    ax.text(0.50,0.16,"단기 급증보다 정책·기술·산업 파이프라인에 의해 유지되는 구조적 고활동 연구영역으로 해석함",ha="center",va="center",fontsize=10.2,fontweight="bold",color=c(DARK))
    for x0 in [0.14,0.39,0.64,0.875]:
        ax.add_patch(FancyArrowPatch((x0,0.57),(0.5,0.39),arrowstyle="-|>",mutation_scale=15,linewidth=1.4,color=c(GRAY),alpha=0.8))
    fig.tight_layout(pad=0.2)
    fig.savefig(path,bbox_inches="tight",facecolor="white")
    plt.close(fig)

fig_paths = {
    "trend": ASSET_DIR / "fig_3_1_monthly_trend.png",
    "mom": ASSET_DIR / "fig_3_2_mom.png",
    "rolling": ASSET_DIR / "fig_3_3_rolling.png",
    "compare": ASSET_DIR / "fig_3_4_compare.png",
    "matrix": ASSET_DIR / "fig_3_5_matrix.png",
    "drivers": ASSET_DIR / "fig_3_6_drivers.png",
}
save_monthly_trend(fig_paths["trend"])
save_mom_change(fig_paths["mom"])
save_rolling(fig_paths["rolling"])
save_period_comparison(fig_paths["compare"])
save_activity_matrix(fig_paths["matrix"])
save_external_drivers(fig_paths["drivers"])

# ------------------------ Word helpers ------------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="5"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ["top","left","bottom","right"]:
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m,v in [("top",top),("start",start),("bottom",bottom),("end",end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_run(run, size=10, bold=False, color=DARK):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = rgb(color)


def pformat(p, before=0, after=4, line=1.25, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = line
    pf.keep_with_next = keep


def add_page_number(p):
    run = p.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin,instr,end])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(); pformat(p, before=12 if level==1 else 8, after=6, line=1.0, keep=True)
    r = p.add_run(text); set_run(r, 15 if level==1 else 12, True, NAVY)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"),"single"); bottom.set(qn("w:sz"),"12"); bottom.set(qn("w:space"),"3"); bottom.set(qn("w:color"),BLUE)
        pBdr.append(bottom); pPr.append(pBdr)
    return p


def add_bullet(doc, text, major=False):
    p = doc.add_paragraph(); pformat(p, after=3, line=1.28)
    p.paragraph_format.left_indent = Cm(0.48); p.paragraph_format.first_line_indent = Cm(-0.34)
    marker = "□ " if major else "· "
    r = p.add_run(marker); set_run(r, 10.2, True, BLUE if major else GRAY)
    r = p.add_run(text); set_run(r, 10.2, major, DARK)
    return p


def add_caption(doc, text, figure=False):
    p = doc.add_paragraph(); pformat(p, before=3, after=6, line=1.0, keep=not figure)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); set_run(r, 9.2, True, GRAY if figure else NAVY)


def add_table(doc, headers, rows, widths, font=8.5):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for row in table.rows:
        for i,w in enumerate(widths): row.cells[i].width = Cm(w)
    for i,h in enumerate(headers):
        cell = table.rows[0].cells[i]; set_cell_shading(cell,NAVY); set_cell_border(cell); set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,after=0,line=1.05)
        r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cell=cells[i]; set_cell_shading(cell,WHITE if ri%2==0 else LIGHT_GRAY); set_cell_border(cell); set_cell_margins(cell)
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pformat(p,after=0,line=1.08)
            r=p.add_run(str(v)); set_run(r,font, i==0, DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return table


def add_callout(doc, title, lines, fill=SKY, accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,accent,"8"); set_cell_margins(cell,130,150,130,150)
    p=cell.paragraphs[0]; pformat(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=cell.add_paragraph(); pformat(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(0.35); p.paragraph_format.first_line_indent=Cm(-0.25)
        r=p.add_run("• "); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)


def add_hyperlink(paragraph, text, url):
    rid=paragraph.part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    h=OxmlElement("w:hyperlink"); h.set(qn("r:id"),rid)
    r=OxmlElement("w:r"); rPr=OxmlElement("w:rPr")
    color=OxmlElement("w:color"); color.set(qn("w:val"),"0563C1"); rPr.append(color)
    u=OxmlElement("w:u"); u.set(qn("w:val"),"single"); rPr.append(u)
    r.append(rPr); t=OxmlElement("w:t"); t.text=text; r.append(t); h.append(r); paragraph._p.append(h)

# ------------------------ build document ------------------------
doc=Document()
sec=doc.sections[0]
sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35)
sec.header_distance=Cm(0.55); sec.footer_distance=Cm(0.55)
style=doc.styles["Normal"]; style.font.name="맑은 고딕"; style._element.rPr.rFonts.set(qn("w:eastAsia"),"맑은 고딕"); style.font.size=Pt(10)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.LEFT
r=header.add_run("arXiv 양자 분야 논문 분석 | 제3장 연구 규모·성장 분석"); set_run(r,8.4,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=footer.add_run("분석 기준: 2025.06~2026.05, 12개월·2,400건  |  "); set_run(r,8.2,False,GRAY); add_page_number(footer)

# cover
cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; cover.autofit=False
cell=cover.cell(0,0); set_cell_shading(cell,NAVY); set_cell_margins(cell,720,300,720,300); set_cell_border(cell,NAVY,"0")
p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("CHAPTER 3"); set_run(r,14,True,"9DC3E6")
p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("연구 규모·성장 분석"); set_run(r,27,True,WHITE)
p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("2025년 6월~2026년 5월 | 12개월·2,400건"); set_run(r,12,True,"D9EAF7")
p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,before=16); r=p.add_run("정량 추세 × 시장·기술·정책 환경 통합 해석"); set_run(r,11,True,"FFD966")

doc.add_paragraph()
# KPI cards as table
kpis=[("2,400건","분석 논문"),("200.0편","월평균"),("+5.0%","후반기 증가"),("217.0편","최근 3개월 평균"),("11.1%","변동계수")]
t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
for i,(val,label) in enumerate(kpis):
    cell=t.cell(0,i); set_cell_shading(cell,[SKY,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]); set_cell_border(cell,[BLUE,GREEN,ORANGE,PURPLE,GRAY][i],"7"); set_cell_margins(cell,100,70,100,70)
    p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,15,True,NAVY)
    p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_run(r,8.7,True,GRAY)

add_callout(doc,"본 장의 핵심 판단",[
    "2026년 6월 부분월 27건을 제외함으로써 최근 흐름은 ‘급감’이 아니라 2026년 3~5월의 재상승·고점 유지로 재해석함.",
    "월평균 200.0편, 중앙값 199.0편, 변동계수 11.1%로 확인되어 양자 연구가 높은 활동성과 제한적 변동성을 동시에 보임.",
    "시장 상용화 신호, 오류정정 진전, 국가전략·공공투자 확대가 학술 연구의 지속성을 뒷받침하는 외부환경으로 작용하는 것으로 판단함.",
],SKY,BLUE)

doc.add_page_break()

add_heading(doc,"3.1 분석 목적 및 본 장의 위치")
add_bullet(doc,"본 장은 제2장에서 확정한 분석 DB를 시간축으로 재구성하여 연구활동의 규모, 변동성, 성장 모멘텀, 피크·저점의 반복성을 분석함.",True)
add_bullet(doc,"연구 규모·성장 분석은 이후 제4장 분야·카테고리 분포 분석에서 ‘어떤 분야가 고점과 회복을 주도했는가’를 확인하기 위한 기준축으로 활용함.")
add_bullet(doc,"제5장 연구 주체 분석에서는 피크 구간의 저자·기관·국가 기여도를 연결하고, 제6장 협업 구조 분석에서는 연구량 증가가 공저·다기관·국제협력 확대와 동행하는지 검증함.")
add_bullet(doc,"정량 추세는 시장·기술·정책 자료와 교차해석하되 외부환경과 논문 수 사이의 직접 인과관계는 단정하지 않고 구조적 부합성의 근거로 활용함.")

add_caption(doc,"표 3-1. 연구 규모·성장 분석의 기준과 적용 범위")
add_table(doc,["구분","적용 기준","분석상 의미","유의사항"],[
    ["분석 기간","2025.06~2026.05","연속 12개월의 월별 연구활동을 비교함","2025년 6월은 6월 6일부터 수집되어 소폭 과소집계 가능성이 있음"],
    ["분석 논문 수","2,400건","모든 규모·성장 지표의 공통 분모로 사용함","2026년 6월 부분월 27건은 제외함"],
    ["시간 단위","월·3개월·6개월","단월 변동과 구조적 모멘텀을 함께 확인함","1개 월의 급등락만으로 성장·감소를 판단하지 않음"],
    ["성장 해석","전월비·이동평균·구간비교","조정·반등·고점 유지의 흐름을 구분함","학회 일정·휴가·연말연초 효과 가능성을 함께 고려함"],
], [2.7,4.2,5.2,5.6],8.4)

add_heading(doc,"3.2 핵심 지표 및 연구활동 규모")
add_bullet(doc,f"12개월 분석 논문은 총 {TOTAL:,}건이며 월평균 {MEAN:.1f}편, 중앙값 {MEDIAN:.1f}편으로 확인됨.",True)
add_bullet(doc,f"최고 월은 2025년 10월 {int(COUNTS.max())}편, 최저 월은 2026년 2월 {int(COUNTS.min())}편이며 최고–최저 변동폭은 {RANGE}편으로 산출됨.")
add_bullet(doc,f"표준편차는 {STD:.1f}편, 변동계수는 {CV:.1f}%로 나타나 월별 조정은 존재하나 고활동 연구영역 내부의 제한적 변동으로 판단함.")
add_bullet(doc,f"평균 이상 월은 {HIGH_MONTHS}개월이며 해당 논문은 {HIGH_TOTAL:,}건, 전체의 {HIGH_SHARE:.1f}%로 확인되어 고활동 구간이 특정 1개월에 집중되지 않음.")
add_bullet(doc,f"선형 추세의 기울기는 월 +{slope:.1f}편이나 설명력은 R²={r2:.3f}로 낮아 단순 직선 성장보다 조정과 반등이 반복되는 계단형 흐름으로 해석함.")

add_caption(doc,"표 3-2. 연구 규모·성장 핵심 지표")
add_table(doc,["지표","값","해석"],[
    ["분석 논문 수",f"{TOTAL:,}건","2026년 6월 부분월을 제외한 공통 분석 모집단임"],
    ["월평균 / 중앙값",f"{MEAN:.1f}편 / {MEDIAN:.1f}편","평균과 중앙값이 유사하여 극단 월의 왜곡이 제한적임"],
    ["최고 / 최저 월",f"2025-10 {int(COUNTS.max())}편 / 2026-02 {int(COUNTS.min())}편","고점과 저점 사이에 단기 조정이 존재함"],
    ["표준편차 / 변동계수",f"{STD:.1f}편 / {CV:.1f}%","고활동 수준 안에서 중간 정도의 월별 변동을 보임"],
    ["전·후반기 증감",f"{FIRST6:,}건 → {SECOND6:,}건 / +{HALF_GROWTH:.1f}%","후반 6개월의 연구량이 소폭 확대됨"],
    ["최근 3개월 모멘텀",f"{PRIOR3:,}건 → {RECENT3:,}건 / +{MOMENTUM:.1f}%","2026년 3월 이후 회복 강도가 확인됨"],
], [3.8,4.1,9.8],8.6)

add_heading(doc,"3.3 월별 논문 제출 추이")
doc.add_picture(str(fig_paths["trend"]),width=Cm(18.0)); add_caption(doc,"그림 3-1. 월별 논문 제출량과 3개월 이동평균",True)
add_bullet(doc,"2025년 6~8월은 182~187편의 초기 안정 구간을 형성하며 평균 이하이나 변동폭이 작게 유지됨.",True)
add_bullet(doc,"2025년 9월 219편, 10월 229편으로 1차 확장 구간이 형성되며 분석기간 최고 수준에 도달함.")
add_bullet(doc,"2025년 11월 171편으로 조정된 뒤 12월 227편으로 즉시 반등하여 일방향 하락이 아닌 고점–조정–재고점의 반복 구조를 보임.")
add_bullet(doc,"2026년 1~2월은 186편과 165편으로 둔화되나, 3월 211편, 4~5월 각각 220편으로 회복하여 최근 구간은 고활동 상태로 종료됨.")
add_bullet(doc,"3개월 이동평균은 2026년 3월 187.3편을 저점으로 4월 198.7편, 5월 217.0편까지 상승하여 단월 반등이 아닌 연속적 회복으로 확인됨.")

add_caption(doc,"표 3-3. 월별 논문 수와 전월 대비 변화")
monthly_rows=[]
for i,m in enumerate(MONTHS):
    if i==0:
        d="-"; p="-"; phase="기준월"
    else:
        d=f"{int(DIFF[i]):+d}편"; p=f"{MOM[i]:+.1f}%"
        phase="확장" if DIFF[i]>=20 else "조정" if DIFF[i]<=-20 else "안정"
    monthly_rows.append([m,f"{int(COUNTS[i])}",d,p,phase])
add_table(doc,["월","논문 수","증감","증감률","활동 국면"],monthly_rows,[2.6,3.0,3.0,3.0,5.9],8.3)

add_heading(doc,"3.4 월별 변동과 조정·반등 구조")
doc.add_picture(str(fig_paths["mom"]),width=Cm(18.0)); add_caption(doc,"그림 3-2. 월별 증감과 조정·반등 구간",True)
add_bullet(doc,"가장 큰 증가폭은 2025년 12월 +56편(+32.7%), 가장 큰 감소폭은 2025년 11월 -58편(-25.3%)로 확인됨.",True)
add_bullet(doc,"대폭 감소 다음 달에 대폭 반등하는 패턴이 나타나므로 개별 월의 급락을 구조적 쇠퇴로 해석하는 것은 부적절함.")
add_bullet(doc,"2026년 2월 -21편(-11.3%) 이후 3월 +46편(+27.9%)으로 회복되어 연초 저점은 기술 관심의 약화보다 제출 일정과 연구 사이클의 영향 가능성이 큼.")
add_bullet(doc,"증가 6개월, 감소 4개월, 보합 1개월로 나타나 방향성은 완만한 증가 우위이나 변동성이 혼재하는 성장형 연구영역으로 판단함.")

doc.add_picture(str(fig_paths["matrix"]),width=Cm(17.6)); add_caption(doc,"그림 3-3. 연구활동 규모–증감 포지션 맵",True)
add_bullet(doc,"2025년 9~10월과 2026년 3~4월은 평균 이상의 연구량과 양(+)의 증감을 동시에 보이는 ‘고규모·확장’ 구간으로 확인됨.",True)
add_bullet(doc,"2025년 12월은 직전 월 급감에 대한 반등 성격이 강하지만 227편으로 고규모 구간에 재진입하여 연구 파이프라인의 복원력이 확인됨.")
add_bullet(doc,"2026년 5월은 220편을 유지하면서 전월 대비 보합을 기록하여 급상승보다 고점 안정화 국면으로 해석함.")

add_heading(doc,"3.5 3개월 이동평균과 최근 모멘텀")
doc.add_picture(str(fig_paths["rolling"]),width=Cm(18.0)); add_caption(doc,"그림 3-4. 3개월 이동평균으로 본 연구활동 모멘텀",True)
add_bullet(doc,"3개월 이동평균은 2025년 10월 210.3편, 12월 209.0편으로 고점이 반복되어 2025년 하반기 연구활동의 지속성을 보여줌.",True)
add_bullet(doc,"2026년 3월 기준 이동평균은 187.3편까지 낮아졌으나 4월 198.7편, 5월 217.0편으로 빠르게 회복함.")
add_bullet(doc,"최근 3개월 평균 217.0편은 전체 평균보다 17.0편(+8.5%) 높아 후속 분야·카테고리 분석에서 부상 영역을 식별할 필요가 큼.")
add_bullet(doc,"단일 월이 아니라 3개월 이동평균이 최고치로 종료된 점은 최근 상승이 일시적 이벤트보다 연속적 연구활동의 결과일 가능성을 높임.")

add_heading(doc,"3.6 전·후반기 및 구간별 성장 비교")
doc.add_picture(str(fig_paths["compare"]),width=Cm(17.2)); add_caption(doc,"그림 3-5. 전·후반기 및 최근 3개월 모멘텀 비교",True)
add_bullet(doc,f"전반 6개월은 {FIRST6:,}건·월평균 {FIRST6/6:.1f}편, 후반 6개월은 {SECOND6:,}건·월평균 {SECOND6/6:.1f}편으로 후반기가 +{HALF_GROWTH:.1f}% 증가함.",True)
add_bullet(doc,f"직전 3개월은 {PRIOR3:,}건·월평균 {PRIOR3/3:.1f}편, 최근 3개월은 {RECENT3:,}건·월평균 {RECENT3/3:.1f}편으로 +{MOMENTUM:.1f}% 증가함.")
add_bullet(doc,"후반기 전체의 증가는 크지 않지만 최근 3개월에서 증가폭이 확대되어 분석기간 말에 모멘텀이 강화된 것으로 보임.")
add_bullet(doc,"따라서 성장 단계는 ‘급격한 팽창’보다 ‘높은 기반 규모 위에서 조정 후 재상승하는 안정 성장’으로 정의하는 것이 적절함.")

add_caption(doc,"표 3-4. 연구활동 단계별 구간 분석")
phases=[
    ["초기 안정","2025.06~08",int(COUNTS[:3].sum()),f"{COUNTS[:3].mean():.1f}","180편대의 안정적 기준선 형성"],
    ["1차 확장","2025.09~10",int(COUNTS[3:5].sum()),f"{COUNTS[3:5].mean():.1f}","평균을 크게 상회하며 1차 고점 형성"],
    ["조정·재반등","2025.11~12",int(COUNTS[5:7].sum()),f"{COUNTS[5:7].mean():.1f}","급감 뒤 즉시 고점 회복"],
    ["연초 조정","2026.01~02",int(COUNTS[7:9].sum()),f"{COUNTS[7:9].mean():.1f}","연초 저점 구간이나 구조적 하락 근거는 제한적"],
    ["최근 재상승","2026.03~05",int(COUNTS[9:12].sum()),f"{COUNTS[9:12].mean():.1f}","평균 상회·고점 유지로 분석기간 종료"],
]
add_table(doc,["국면","기간","합계","월평균","해석"],phases,[3.0,3.4,2.8,3.1,8.3],8.3)

add_heading(doc,"3.7 정량 추세와 최신 시장·기술·정책 환경의 통합 해석")
doc.add_picture(str(fig_paths["drivers"]),width=Cm(18.0)); add_caption(doc,"그림 3-6. 연구활동 규모를 뒷받침하는 시장·기술·정책 환경",True)
add_bullet(doc,"시장 측면에서는 McKinsey의 2026년 모니터가 300개 이상의 기업이 양자컴퓨팅을 활용·협업하고 있으며 2025년 양자컴퓨팅 기업 매출이 10억 달러를 넘었다고 제시함. 이는 학술 연구량의 안정성이 상용화 준비와 병행되는 흐름으로 해석할 수 있는 보조 근거임 [R3].",True)
add_bullet(doc,"기술 측면에서는 Google Quantum AI의 below-threshold 표면코드 실증이 2025년 Nature에 게재되었고, IBM은 2026년 양자 이점과 2029년 대규모 내결함성 시스템을 목표로 제시함. 이러한 이정표는 오류정정·디코더·제어·하이브리드 컴퓨팅 연구를 지속시키는 기술 수요로 연결될 수 있음 [R4][R5].")
add_bullet(doc,"정책 측면에서는 미국 NQI가 FY2025 약 9.98억 달러의 QIS R&D 예산을 요청하고 장기간 연방 연구투자를 유지하였으며, EU는 2025년 연구혁신·인프라·생태계·우주/듀얼유스·인력의 5개 전략축을 제시함 [R1][R2].")
add_bullet(doc,"DARPA QBI는 2033년 산업적으로 유용한 양자컴퓨터의 실현 가능성을 독립 검증하고 있으며, 2026년 Stage A와 IV&V 프로그램을 확대함. 이는 단순 큐비트 수보다 유용성·비용·검증 중심 연구가 강화되는 환경을 보여줌 [R6].")
add_bullet(doc,"분석기간 종료 이후인 2026년 6월 미국 EO 14413은 양자컴퓨팅·센싱·네트워킹의 상용화와 배치를 가속하는 정부 차원의 후속 신호임. 본 장에서는 분석기간 내 논문 증가의 직접 원인으로 사용하지 않고 연구활동 지속 가능성을 확인하는 사후 정책 근거로 활용함 [R7].")

add_caption(doc,"표 3-5. 정량 신호와 외부환경 근거의 결합 해석")
add_table(doc,["정량 신호","정성 근거","통합 해석","후속 분석 질문"],[
    ["월평균 200.0편","미국·EU·영국의 장기 국가전략과 공공투자","단기 유행보다 장기 연구 인프라에 기반한 고활동 영역으로 보임","국가별 연구량과 정책투자의 부합성이 나타나는가"],
    ["후반기 +5.0%","민간투자·기업 채택과 양자서비스 시장 확대","연구 생태계가 기초연구와 상용화 준비를 동시에 확대하는 것으로 보임","컴퓨팅·센싱·통신 중 어느 분야가 증가를 주도하는가"],
    ["최근 3개월 217.0편","오류정정 실증·FTQC·하이브리드 로드맵 구체화","기술 병목 해결을 위한 후속 연구가 최근 모멘텀을 지지할 가능성이 있음","오류정정·디코더·제어 키워드가 실제 증가했는가"],
    ["조정 후 즉시 반등","DARPA 등 독립 검증·벤치마킹 강화","기술 기대와 검증 수요가 함께 존재하여 연구량이 저점에서 복원되는 구조로 보임","피크월의 연구 주체와 협업 네트워크가 달라지는가"],
], [3.4,5.0,6.3,5.0],8.1)

add_callout(doc,"정성 통합 해석의 주의사항",[
    "외부 시장·정책 자료와 월별 논문 수의 시간적 동행은 구조적 부합성을 설명하지만 직접적인 인과관계를 입증하지 않음.",
    "McKinsey·기업 로드맵은 시장·전략 신호로 활용하되 정부 공식자료, 동료평가 논문, 독립 벤치마킹 자료와 교차검증함.",
    "2026년 6월 이후 공개된 자료는 분석기간 내 수치의 원인이 아니라 향후 연구량 지속 가능성을 확인하는 사후 환경자료로 구분함.",
],LIGHT_ORANGE,ORANGE)

add_heading(doc,"3.8 전략적 시사점 및 후속 장 연결")
add_bullet(doc,"양자 분야는 월평균 200편의 안정적 고활동 연구영역이므로 연 1회 정적 보고보다 월별 업데이트와 3개월 이동평균 기반 상시 모니터링이 적합함.",True)
add_bullet(doc,"최근 3개월 평균이 전체 평균을 상회하므로 제4장에서는 2026년 3~5월 증가분을 주 카테고리, 전체 카테고리, 복수 카테고리 조합으로 분해할 필요가 있음.")
add_bullet(doc,"1차 고점(2025년 9~10월), 재반등(2025년 12월), 최근 고점(2026년 3~5월)의 기술 구성을 비교하면 일회성 이슈와 지속 성장 기술을 구분할 수 있음.")
add_bullet(doc,"연구량 증가가 특정 저자·기관의 집중인지 글로벌 저변 확대인지 확인하기 위해 제5장의 연구 주체 분석과 월별 교차표를 구축해야 함.")
add_bullet(doc,"다기관·국제협력 비율이 피크 구간에서 상승하는지 확인하여 연구량 증가가 네트워크 확장에 기반하는지 제6장에서 검증해야 함.")

add_caption(doc,"표 3-6. 본 장의 발견과 제4~7장 연결 과제")
add_table(doc,["본 장의 발견","후속 분석","전략 활용"],[
    ["최근 3개월 평균 217.0편","2026.03~05의 증가 카테고리·키워드·토픽을 추출함","부상 기술 후보군을 도출함"],
    ["피크가 여러 시점에 반복됨","2025.09~10, 2025.12, 2026.03~05의 공통 기술군을 비교함","지속 성장 기술과 이벤트성 기술을 구분함"],
    ["저점 뒤 반등이 빠름","저점·고점 월의 저자·기관·국가 구성을 비교함","핵심 연구 허브와 회복 주체를 파악함"],
    ["변동계수 11.1%","분야별·국가별 변동계수와 성장률을 비교함","안정 코어와 고변동 부상영역을 분리함"],
    ["정책·시장 환경이 동시 강화됨","정책 발표·표준화·기술 이정표와 토픽 증가 시점을 매핑함","정량–정성 통합 모니터링 체계를 설계함"],
], [5.4,7.0,6.3],8.3)

add_heading(doc,"3.9 분석 한계 및 검증 과제")
add_bullet(doc,"2025년 6월은 6월 6일부터 수집되어 완전월보다 최대 5일이 누락될 수 있으므로 전반기 증가율이 소폭 과대평가될 가능성을 명시함.",True)
add_bullet(doc,"12개월 데이터만으로 계절성을 확정할 수 없으므로 연말연초·학회 일정 효과는 가설로 제시하고 최소 24~36개월 확장 분석으로 검증해야 함.")
add_bullet(doc,"arXiv 제출량은 연구활동의 선행 신호이나 동료평가 완료, 상용화 성과, 기술 품질을 직접 나타내지 않으므로 인용·DOI·특허·정책성과와 별도 연계해야 함.")
add_bullet(doc,"제목 검색식에 의해 양자화(quantization) 기반 컴퓨터과학 논문이 일부 포함될 수 있으므로 성장 월의 증가분은 초록·전문 기반 기술분류로 재확인해야 함.")

add_caption(doc,"표 3-7. 한계요인과 보완 방향")
add_table(doc,["한계요인","영향","현재 통제","보완 방향"],[
    ["2025년 6월 부분성","전반기 규모가 다소 작게 집계될 수 있음","지정 분석기간의 기준월로 포함하고 주석 처리함","일단위 보정 또는 2025.07~2026.05 민감도 분석을 병행함"],
    ["12개월의 짧은 기간","계절성과 구조적 추세를 분리하기 어려움","3개월 이동평균과 6개월 비교를 사용함","24~36개월 장기 시계열로 확장함"],
    ["프리프린트 중심성","품질·상용화 성과와 차이가 존재함","연구활동성 지표로 범위를 제한함","인용·저널게재·특허·과제 데이터와 연계함"],
    ["검색식 범위","비핵심 양자화 논문이 유입될 수 있음","카테고리와 전문을 이용한 후속 점검을 전제함","기술분류 모델과 제외규칙을 구축함"],
], [3.2,5.0,5.4,5.8],8.2)

add_heading(doc,"3.10 본 장 소결")
add_bullet(doc,"2026년 6월 부분월을 제외한 결과, 연구 규모는 12개월 2,400건·월평균 200.0편으로 확정되며 양자 분야는 안정적 고활동 연구영역으로 판단함.",True)
add_bullet(doc,"후반 6개월은 전반기 대비 +5.0%, 최근 3개월은 직전 3개월 대비 +12.6% 증가하여 분석기간 말의 모멘텀이 강화된 것으로 나타남.")
add_bullet(doc,"월별 흐름은 선형 급성장보다 고점–조정–재반등이 반복되는 계단형 성장 구조를 보이며, 2026년 4~5월은 각 220편으로 고점 안정화가 확인됨.")
add_bullet(doc,"시장 상용화, 오류정정 진전, 국가전략과 공공투자, 독립 벤치마킹 확대가 안정적 연구량과 부합하는 외부환경으로 작용함.")
add_bullet(doc,"다음 제4장에서는 최근 재상승 구간과 반복 피크 구간을 중심으로 양자물리(quant-ph), 컴퓨터과학, 응집물질, 광학, 암호·보안 등 분야·카테고리의 성장 기여도를 분석함.")

add_heading(doc,"참고자료")
refs=[
    ("[R1]","National Quantum Initiative, Annual Report FY2025","https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf"),
    ("[R2]","European Commission, Quantum Europe Strategy, 2025","https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy"),
    ("[R3]","McKinsey, Quantum Technology Monitor 2026: A Commercial Tipping Point","https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/mckinsey-quantum-technology-monitor-2026-a-commercial-tipping-point"),
    ("[R4]","Google Quantum AI and Collaborators, Quantum Error Correction Below the Surface Code Threshold, Nature 638, 2025","https://www.nature.com/articles/s41586-024-08449-y"),
    ("[R5]","IBM Quantum, Hardware and Roadmap","https://www.ibm.com/quantum/hardware"),
    ("[R6]","DARPA, Quantum Benchmarking Initiative","https://www.darpa.mil/research/programs/quantum-benchmarking-initiative"),
    ("[R7]","The White House, Executive Order 14413: Ushering in the Next Frontier of Quantum Innovation, 2026","https://www.whitehouse.gov/presidential-actions/2026/06/ushering-in-the-next-frontier-of-quantum-innovation/"),
    ("[R8]","UK Government, National Quantum Strategy","https://www.gov.uk/government/publications/national-quantum-strategy"),
]
for rid,title,url in refs:
    p=doc.add_paragraph(); pformat(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(0.42); p.paragraph_format.first_line_indent=Cm(-0.42)
    r=p.add_run(rid+" "); set_run(r,8.8,True,NAVY); add_hyperlink(p,title,url)

# core properties
cp=doc.core_properties
cp.title="arXiv 양자 분야 제3장 연구 규모·성장 분석"
cp.subject="2025년 6월~2026년 5월, 12개월·2,400건 기준 정량·정성 통합 분석"
cp.author="OpenAI"
cp.keywords="arXiv, 양자, 연구 규모, 성장 분석, 시장, 기술, 정책"
settings=doc.settings._element; update=OxmlElement("w:updateFields"); update.set(qn("w:val"),"true"); settings.append(update)

doc.save(OUTPUT)
print(f"Created {OUTPUT}")

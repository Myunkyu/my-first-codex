from __future__ import annotations

from pathlib import Path
import math

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path('artifacts')
ASSET_DIR = OUT_DIR / 'chapter4_assets'
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / 'arXiv_양자분야_제4장_분야_카테고리_분포분석_고품질_재작성_20260704.docx'

TOTAL = 2400
PRIMARY_CATS = [
    ('양자물리 (quant-ph)',1680,70.00),('메조스코픽/홀 효과 (cond-mat.mes-hall)',69,2.88),
    ('기계학습 (cs.LG)',44,1.83),('암호/보안 (cs.CR)',43,1.79),('컴퓨터 비전 (cs.CV)',34,1.42),
    ('고에너지 현상론 (hep-ph)',31,1.29),('강상관 전자 (cond-mat.str-el)',30,1.25),
    ('고에너지 물리 이론 (hep-th)',28,1.17),('광학 (physics.optics)',26,1.08),
    ('화학물리 (physics.chem-ph)',25,1.04),('일반상대론/양자중력 (gr-qc)',23,0.96),
    ('통계역학 (cond-mat.stat-mech)',22,0.92),('신호처리 (eess.SP)',21,0.88),
    ('신흥기술 (cs.ET)',19,0.79),('재료과학 (cond-mat.mtrl-sci)',18,0.75),
]
FIELD_GROUPS = [
    ('양자물리 코어 (Quantum physics core)',1680,70.00),('컴퓨터과학 (Computer science)',225,9.38),
    ('응집물질 (Condensed matter)',169,7.04),('응용/원자/화학물리 (Applied/atomic/chemical physics)',115,4.79),
    ('고에너지·중력·핵 (HEP/gravity/nuclear)',106,4.42),('수학·통계·비선형 (Math/stat/nonlinear)',48,2.00),
    ('전기전자·시스템 (Electrical engineering & systems)',38,1.58),('기타 융합 (Other interdisciplinary)',19,0.79),
]
ALL_CATS = [
    ('양자물리 (quant-ph)',1901,79.21),('기계학습 (cs.LG)',143,5.96),('메조스코픽/홀 효과 (cond-mat.mes-hall)',123,5.12),
    ('광학 (physics.optics)',103,4.29),('통계역학 (cond-mat.stat-mech)',100,4.17),('고에너지 물리 이론 (hep-th)',95,3.96),
    ('강상관 전자 (cond-mat.str-el)',94,3.92),('암호/보안 (cs.CR)',90,3.75),('인공지능 (cs.AI)',83,3.46),
    ('수리물리 (math-ph)',82,3.42),('신흥기술 (cs.ET)',76,3.17),('화학물리 (physics.chem-ph)',67,2.79),
    ('고에너지 현상론 (hep-ph)',60,2.50),('일반상대론/양자중력 (gr-qc)',58,2.42),('재료과학 (cond-mat.mtrl-sci)',53,2.21),
]
CAT_COUNTS = [('1개',1262,52.58),('2개',719,29.96),('3개',289,12.04),('4개',105,4.38),('5개',25,1.04)]
PAIR_EXACT = [
    ('양자물리 + 광학',44,1.83),('양자물리 + 수리물리',34,1.42),('양자물리 + 암호/보안',31,1.29),
    ('양자물리 + 통계역학',28,1.17),('양자물리 + 기계학습',27,1.12),('양자물리 + 강상관 전자',24,1.00),
    ('양자물리 + 신흥기술',22,0.92),('양자물리 + 화학물리',22,0.92),('양자물리 + 메조스코픽/홀 효과',18,0.75),
    ('양자물리 + 원자물리',16,0.67),
]
COOCCUR = [
    ('통계역학 (cond-mat.stat-mech)',90),('기계학습 (cs.LG)',86),('광학 (physics.optics)',80),
    ('수리물리 (math-ph)',69),('강상관 전자 (cond-mat.str-el)',68),('고에너지 물리 이론 (hep-th)',63),
    ('메조스코픽/홀 효과 (cond-mat.mes-hall)',62),('신흥기술 (cs.ET)',53),('화학물리 (physics.chem-ph)',51),
    ('암호/보안 (cs.CR)',47),
]
GROWTH = [
    ('양자물리 (quant-ph)',824,856,32,3.9),('통계역학 (cond-mat.stat-mech)',6,16,10,166.7),
    ('정보이론 (cs.IT)',3,10,7,233.3),('암호/보안 (cs.CR)',20,23,3,15.0),
    ('일반상대론/양자중력 (gr-qc)',10,13,3,30.0),('컴퓨터 비전 (cs.CV)',16,18,2,12.5),
    ('메조스코픽/홀 효과 (cond-mat.mes-hall)',34,35,1,2.9),('격자 고에너지물리 (hep-lat)',6,7,1,16.7),
    ('소프트웨어공학 (cs.SE)',5,6,1,20.0),('원자물리 (physics.atom-ph)',5,6,1,20.0),
]

HHI = 0.493
EFFECTIVE_CATS = 1 / HHI
ASSIGNMENTS = sum(int(k.replace('개','')) * n for k,n,_ in CAT_COUNTS)
AVG_CATS = ASSIGNMENTS / TOTAL
MULTI_N = TOTAL - CAT_COUNTS[0][1]
MULTI_AVG = (ASSIGNMENTS - CAT_COUNTS[0][1]) / MULTI_N
TOP10_PAIRS_N = sum(v for _,v,_ in PAIR_EXACT)
TOP10_PAIRS_MULTI_SHARE = TOP10_PAIRS_N / MULTI_N * 100

NAVY='17365D'; BLUE='2F75B5'; SKY='D9EAF7'; GREEN='548235'; LIGHT_GREEN='E2F0D9'; ORANGE='ED7D31'; LIGHT_ORANGE='FCE4D6'; RED='C00000'; LIGHT_RED='F4CCCC'; PURPLE='7030A0'; LIGHT_PURPLE='E4DFEC'; GRAY='6B7280'; LIGHT_GRAY='F2F4F7'; DARK='243447'; WHITE='FFFFFF'; BORDER='CAD4DF'; YELLOW='FFD966'

def c(h): return '#'+h

def rgb(h): return RGBColor.from_string(h)

font_candidates=['/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc','/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc','/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc']
font_path=next((p for p in font_candidates if Path(p).exists()),None)
if font_path:
    fp=font_manager.FontProperties(fname=font_path); plt.rcParams['font.family']=fp.get_name()
else: plt.rcParams['font.family']='DejaVu Sans'
plt.rcParams['axes.unicode_minus']=False
plt.rcParams['font.weight']='bold'

# ---------------- charts ----------------
def base_axes(figsize=(11.5,5.2)):
    fig,ax=plt.subplots(figsize=figsize,dpi=220); fig.patch.set_facecolor('white'); ax.set_facecolor('white')
    for s in ['top','right']: ax.spines[s].set_visible(False)
    ax.spines['left'].set_color(c(BORDER)); ax.spines['bottom'].set_color(c(BORDER)); ax.grid(axis='x',color=c(BORDER),lw=.7,alpha=.55); ax.set_axisbelow(True)
    ax.tick_params(colors=c(DARK),labelsize=9)
    return fig,ax

def save_core_share(path):
    fig,ax=plt.subplots(figsize=(8.8,5.4),dpi=220); fig.patch.set_facecolor('white')
    vals=[70,30]; colors=[c(BLUE),c(SKY)]
    ax.pie(vals,colors=colors,startangle=90,counterclock=False,wedgeprops=dict(width=.34,edgecolor='white'))
    ax.text(0,0.08,'70.00%',ha='center',va='center',fontsize=22,fontweight='bold',color=c(NAVY))
    ax.text(0,-0.18,'양자물리\n(quant-ph)',ha='center',va='center',fontsize=11,fontweight='bold',color=c(DARK))
    ax.text(0,-1.28,'주 카테고리 1,680건 · 기타 720건',ha='center',fontsize=10.5,fontweight='bold',color=c(GRAY))
    ax.set_title('주 카테고리의 코어 집중 구조',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    ax.axis('equal'); fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_primary_zoom(path):
    data=PRIMARY_CATS[1:]; labels=[x[0] for x in data][::-1]; vals=[x[1] for x in data][::-1]
    fig,ax=base_axes((12.0,6.7)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=c(SKY),edgecolor=c(BLUE),lw=1.0)
    for b,v in zip(bars,vals): ax.text(v+1.2,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.6); ax.set_xlim(0,76); ax.set_xlabel('논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('주 카테고리 TOP2~15 확대 분포',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_field_groups(path):
    labels=[x[0] for x in FIELD_GROUPS][::-1]; vals=[x[1] for x in FIELD_GROUPS][::-1]
    cols=[c(SKY),c(LIGHT_PURPLE),c(LIGHT_ORANGE),c(LIGHT_GREEN),c(LIGHT_GRAY),c(LIGHT_PURPLE),c(LIGHT_ORANGE),c(BLUE)][::-1]
    fig,ax=base_axes((12.0,5.8)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=cols,edgecolor=c(NAVY),lw=.8)
    for b,v in zip(bars,vals): ax.text(v+18,b.get_y()+b.get_height()/2,f'{v:,}건 ({v/TOTAL*100:.2f}%)',va='center',fontsize=9.2,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.8); ax.set_xlim(0,1900); ax.set_xlabel('논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('실무 해석용 상위 분야군 분포',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_all_categories(path):
    labels=[x[0] for x in ALL_CATS[1:]][::-1]; vals=[x[1] for x in ALL_CATS[1:]][::-1]
    fig,ax=base_axes((12.0,6.7)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=c(LIGHT_GREEN),edgecolor=c(GREEN),lw=1.0)
    for b,v in zip(bars,vals): ax.text(v+2,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.6); ax.set_xlim(0,160); ax.set_xlabel('포함 논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('전체 카테고리 포함 기준 TOP2~15',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    ax.text(78,13.8,'양자물리(quant-ph) 1,901건·79.21%는 별도 코어 지표로 관리함',ha='center',fontsize=9.4,fontweight='bold',color=c(PURPLE))
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_category_counts(path):
    labels=[x[0] for x in CAT_COUNTS]; vals=[x[1] for x in CAT_COUNTS]; shares=[x[2] for x in CAT_COUNTS]
    fig,ax=base_axes((10.5,5.0)); x=np.arange(len(labels)); colors=[c(SKY),c(BLUE),c(GREEN),c(ORANGE),c(PURPLE)]; bars=ax.bar(x,vals,color=colors,edgecolor=c(NAVY),lw=.8,width=.63)
    for b,v,s in zip(bars,vals,shares): ax.text(b.get_x()+b.get_width()/2,v+34,f'{v:,}건\n{s:.2f}%',ha='center',fontsize=10,fontweight='bold',color=c(NAVY))
    ax.set_xticks(x); ax.set_xticklabels(labels,fontweight='bold'); ax.set_ylim(0,1430); ax.set_ylabel('논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('논문별 부여 카테고리 수 분포',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    ax.text(2.5,1310,f'복수 카테고리 1,138건·47.42%  |  논문당 평균 {AVG_CATS:.2f}개',ha='center',fontsize=10.5,fontweight='bold',color=c(GREEN))
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_exact_pairs(path):
    labels=[x[0] for x in PAIR_EXACT][::-1]; vals=[x[1] for x in PAIR_EXACT][::-1]
    fig,ax=base_axes((12.0,5.8)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=c(LIGHT_PURPLE),edgecolor=c(PURPLE),lw=1)
    for b,v in zip(bars,vals): ax.text(v+.8,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9.2,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.7); ax.set_xlim(0,49); ax.set_xlabel('정확 조합 논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('양자물리 중심 정확 복수 카테고리 조합 TOP10',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_cooccur(path):
    labels=[x[0] for x in COOCCUR][::-1]; vals=[x[1] for x in COOCCUR][::-1]
    fig,ax=base_axes((12.0,5.8)); y=np.arange(len(labels)); bars=ax.barh(y,vals,color=c(LIGHT_ORANGE),edgecolor=c(ORANGE),lw=1)
    for b,v in zip(bars,vals): ax.text(v+1.2,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9.2,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.7); ax.set_xlim(0,99); ax.set_xlabel('quant-ph 동시 부여 논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('양자물리(quant-ph) 교차 카테고리 TOP10',fontsize=16,fontweight='bold',color=c(NAVY),pad=18)
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_growth(path):
    labels=[x[0] for x in GROWTH][::-1]; before=[x[1] for x in GROWTH][::-1]; after=[x[2] for x in GROWTH][::-1]
    y=np.arange(len(labels)); fig,ax=base_axes((12.0,6.0)); ax.barh(y-.18,before,height=.34,color=c(SKY),edgecolor=c(BLUE),label='전반 6개월'); ax.barh(y+.18,after,height=.34,color=c(LIGHT_GREEN),edgecolor=c(GREEN),label='후반 6개월')
    for yy,b,a in zip(y,before,after): ax.text(max(b,a)+1.0,yy,f'{a-b:+d}',va='center',fontsize=8.8,fontweight='bold',color=c(NAVY))
    ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.5); ax.set_xlim(0,900); ax.set_xlabel('논문 수(건)',fontweight='bold',color=c(NAVY)); ax.set_title('주요 카테고리 전·후반 6개월 변화',fontsize=16,fontweight='bold',color=c(NAVY),pad=18); ax.legend(frameon=False,loc='lower right')
    fig.tight_layout(pad=1.0); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

def save_mapping(path):
    fig,ax=plt.subplots(figsize=(12,5.6),dpi=220); fig.patch.set_facecolor('white'); ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    ax.add_patch(FancyBboxPatch((.36,.38),.28,.26,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor=c(NAVY),lw=2.5))
    ax.text(.50,.56,'양자물리 코어',ha='center',fontsize=15,fontweight='bold',color=c(NAVY)); ax.text(.50,.47,'quant-ph 70.00%',ha='center',fontsize=11.5,fontweight='bold',color=c(BLUE)); ax.text(.50,.41,'학술 중심축',ha='center',fontsize=10,fontweight='bold',color=c(GRAY))
    items=[
        (.03,.67,.27,.21,SKY,BLUE,'양자컴퓨팅·SW/AI','cs.LG · cs.AI · cs.ET · eess.SP'),
        (.70,.67,.27,.21,LIGHT_PURPLE,PURPLE,'양자통신·보안','cs.CR · physics.optics · QKD/PQC'),
        (.03,.16,.27,.21,LIGHT_GREEN,GREEN,'양자센싱·계측','physics.optics · atom-ph · eess.SP'),
        (.70,.16,.27,.21,LIGHT_ORANGE,ORANGE,'양자소자·소재','mes-hall · str-el · mtrl-sci'),
        (.36,.72,.28,.18,LIGHT_GRAY,GRAY,'양자이론·시뮬레이션','math-ph · stat-mech · hep-th'),
        (.36,.08,.28,.18,'FFF2CC','BF9000','양자화학·응용','physics.chem-ph · material/energy'),
    ]
    for x,y,w,h,fc,ec,title,sub in items:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor=c(fc),edgecolor=c(ec),lw=1.8))
        ax.text(x+w/2,y+h*.64,title,ha='center',fontsize=11.2,fontweight='bold',color=c(NAVY)); ax.text(x+w/2,y+h*.28,sub,ha='center',fontsize=8.8,fontweight='bold',color=c(DARK))
    for start,end in [((.36,.52),(.30,.75)),((.64,.52),(.70,.75)),((.36,.43),(.30,.27)),((.64,.43),(.70,.27)),((.50,.64),(.50,.72)),((.50,.38),(.50,.26))]:
        ax.add_patch(FancyArrowPatch(start,end,arrowstyle='<->',mutation_scale=14,lw=1.4,color=c(GRAY)))
    ax.text(.50,.96,'arXiv 학술 카테고리의 실무형 양자기술 재매핑 프레임',ha='center',fontsize=16,fontweight='bold',color=c(NAVY))
    ax.text(.50,.015,'단일 귀속이 아닌 N:M 매핑을 적용하여 융합기술과 산업 응용 경로를 보존함',ha='center',fontsize=10.3,fontweight='bold',color=c(GRAY))
    fig.tight_layout(pad=.2); fig.savefig(path,bbox_inches='tight'); plt.close(fig)

FIGS={
'core':ASSET_DIR/'fig_4_1_core.png','primary':ASSET_DIR/'fig_4_2_primary_zoom.png','groups':ASSET_DIR/'fig_4_3_groups.png','all':ASSET_DIR/'fig_4_4_all_categories.png','counts':ASSET_DIR/'fig_4_5_counts.png','pairs':ASSET_DIR/'fig_4_6_pairs.png','cooccur':ASSET_DIR/'fig_4_7_cooccur.png','growth':ASSET_DIR/'fig_4_8_growth.png','mapping':ASSET_DIR/'fig_4_9_mapping.png'}
save_core_share(FIGS['core']); save_primary_zoom(FIGS['primary']); save_field_groups(FIGS['groups']); save_all_categories(FIGS['all']); save_category_counts(FIGS['counts']); save_exact_pairs(FIGS['pairs']); save_cooccur(FIGS['cooccur']); save_growth(FIGS['growth']); save_mapping(FIGS['mapping'])

# ---------------- doc helpers ----------------
def set_cell_shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_border(cell,color=BORDER,size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ['top','left','bottom','right']:
        el=borders.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)

def set_cell_margins(cell,top=80,start=90,bottom=80,end=90):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_run(run,size=10,bold=False,color=DARK):
    run.font.name='맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=rgb(color)

def pformat(p,before=0,after=4,line=1.25,keep=False):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line; pf.keep_with_next=keep

def add_page_number(p):
    run=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text='PAGE'; e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); run._r.extend([b,i,e])

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); pformat(p,before=12 if level==1 else 8,after=6,line=1.0,keep=True); r=p.add_run(text); set_run(r,15 if level==1 else 12,True,NAVY)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'3'); bottom.set(qn('w:color'),BLUE); pBdr.append(bottom); pPr.append(pBdr)
    return p

def add_bullet(doc,text,major=False):
    p=doc.add_paragraph(); pformat(p,after=3,line=1.28); p.paragraph_format.left_indent=Cm(.48); p.paragraph_format.first_line_indent=Cm(-.34); r=p.add_run('□ ' if major else '· '); set_run(r,10.2,True,BLUE if major else GRAY); r=p.add_run(text); set_run(r,10.2,major,DARK); return p

def add_caption(doc,text,figure=False):
    p=doc.add_paragraph(); pformat(p,before=3,after=6,line=1.0,keep=not figure); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); set_run(r,9.2,True,GRAY if figure else NAVY)

def add_table(doc,headers,rows,widths,font=8.4):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,NAVY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,after=0,line=1.05); r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cell=cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,WHITE if ri%2==0 else LIGHT_GRAY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pformat(p,after=0,line=1.08); r=p.add_run(str(v)); set_run(r,font,i==0,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0); return table

def add_callout(doc,title,lines,fill=SKY,accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,accent,'8'); set_cell_margins(cell,130,150,130,150); p=cell.paragraphs[0]; pformat(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=cell.add_paragraph(); pformat(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.first_line_indent=Cm(-.25); r=p.add_run('• '); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_hyperlink(p,text,url):
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rPr.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)

# ---------------- document ----------------
doc=Document(); sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35); sec.header_distance=Cm(.55); sec.footer_distance=Cm(.55)
style=doc.styles['Normal']; style.font.name='맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); style.font.size=Pt(10)
header=sec.header.paragraphs[0]; r=header.add_run('arXiv 양자 분야 논문 분석 | 제4장 분야·카테고리 분포 분석'); set_run(r,8.4,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('분석 기준: 2025.06~2026.05, 12개월·2,400건  |  '); set_run(r,8.2,False,GRAY); add_page_number(footer)

cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; cell=cover.cell(0,0); set_cell_shading(cell,NAVY); set_cell_margins(cell,720,300,720,300); set_cell_border(cell,NAVY,'0'); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CHAPTER 4'); set_run(r,14,True,'9DC3E6'); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('분야·카테고리 분포 분석'); set_run(r,26,True,WHITE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('2025년 6월~2026년 5월 | 12개월·2,400건'); set_run(r,12,True,SKY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,before=16); r=p.add_run('양자물리 코어 × 다학제 확장 × 기술분류 재매핑'); set_run(r,11,True,YELLOW)

doc.add_paragraph(); kpis=[('87개','주 카테고리'),('117개','전체 카테고리'),('70.00%','quant-ph 주분류'),('47.42%','복수 카테고리'),(f'{HHI:.3f}','주분류 HHI')]; t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(val,label) in enumerate(kpis):
    cell=t.cell(0,i); set_cell_shading(cell,[SKY,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]); set_cell_border(cell,[BLUE,GREEN,ORANGE,PURPLE,GRAY][i],'7'); set_cell_margins(cell,100,70,100,70); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,15,True,NAVY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_run(r,8.6,True,GRAY)
add_callout(doc,'본 장의 핵심 판단',[
    '주 카테고리 기준 양자물리(quant-ph)가 1,680건·70.00%로 강한 학술 코어를 형성함.',
    '전체 카테고리는 117개로 확대되고 복수 카테고리 논문이 1,138건·47.42%로 나타나 코어 집중과 융합 확장이 동시에 존재함.',
    '광학, 기계학습, 암호/보안, 응집물질, 수리물리와의 교차는 양자센싱·컴퓨팅·통신·보안·소자·시뮬레이션의 실무 기술축으로 연결됨.',
],SKY,BLUE)
doc.add_page_break()

add_heading(doc,'4.1 분석 목적 및 제3장과의 연결')
add_bullet(doc,'제3장에서 확인한 월평균 200편의 안정적 연구활동이 어떤 학술분야와 교차 영역에서 형성되는지 분해하여 분석함.',True)
add_bullet(doc,'주 카테고리는 논문의 대표 학술분야를, 전체 카테고리는 주·보조 분류를 모두 포함한 연결 분야를 의미함.')
add_bullet(doc,'카테고리 분석 결과는 제5장에서 국가·기관·저자의 전문분야를 비교하고 제6장에서 융합 연구와 협업구조를 연결하는 기준축으로 활용함.')
add_bullet(doc,'arXiv 카테고리는 학술 분류이므로 산업·정책 기술분류와 1:1 대응하지 않으며, 후반부에서 실무형 기술군으로 N:M 재매핑함.')
add_caption(doc,'표 4-1. 분야·카테고리 분석 항목과 해석 목적')
add_table(doc,['분석 항목','집계 기준','해석 목적','후속 연결'],[
    ['주 카테고리 분포','primary_category별 논문 수·비중','핵심 연구축과 편중 수준을 파악함','제5장 주체별 전문분야'],
    ['전체 카테고리 분포','categories 전체 출현 논문 수','주분류 밖 인접 분야와 응용 확산을 확인함','제7장 기술축 도출'],
    ['복수 카테고리','논문별 category_count','융합 연구 비중과 다학제 강도를 측정함','제6장 협업과 교차검증'],
    ['정확 조합·동시출현','카테고리 조합 및 quant-ph 동시 부여','융합 방향과 브리지 분야를 식별함','유망기술 후보군'],
    ['집중도·성장','TOP N, HHI, 전·후반기 증감','코어 집중과 롱테일·부상 신호를 함께 판단함','상시 모니터링'],
], [3.3,5.0,5.6,5.0],8.2)

add_heading(doc,'4.2 주 카테고리: 양자물리 코어의 절대적 중심성')
doc.add_picture(str(FIGS['core']),width=Cm(13.5)); add_caption(doc,'그림 4-1. 주 카테고리의 양자물리(quant-ph) 코어 비중',True)
add_bullet(doc,'양자물리(quant-ph)는 1,680건·70.00%로 확인되어 검색식이 양자 핵심 논문군을 안정적으로 포착함.',True)
add_bullet(doc,'TOP1 비중이 매우 높아 주분류만 보면 단일 코어 중심 구조이나, 나머지 30.00%에 86개 주 카테고리가 분포하여 응용 확장 저변이 넓음.')
add_bullet(doc,f'HHI는 {HHI:.3f}, 유효 카테고리 수는 약 {EFFECTIVE_CATS:.1f}개로 산출되어 형식적 카테고리 수 87개와 실질적 집중도의 차이가 큼.')

doc.add_picture(str(FIGS['primary']),width=Cm(18.1)); add_caption(doc,'그림 4-2. 주 카테고리 TOP2~15 확대 분포',True)
add_bullet(doc,'메조스코픽/홀 효과(cond-mat.mes-hall) 69건은 양자소자·나노구조·메조스코픽 물성 연구가 주요 보조축임을 보임.',True)
add_bullet(doc,'기계학습(cs.LG) 44건, 암호/보안(cs.CR) 43건, 컴퓨터 비전(cs.CV) 34건이 상위권에 포함되어 계산·AI·보안·이미지 기반 응용이 물리 코어와 병행됨.')
add_bullet(doc,'강상관 전자(cond-mat.str-el), 광학(physics.optics), 화학물리(physics.chem-ph), 재료과학(cond-mat.mtrl-sci)은 하드웨어·소재·측정 플랫폼의 기반 연구축으로 해석함.')
add_caption(doc,'표 4-2. 주 카테고리 TOP15')
add_table(doc,['순위','주 카테고리','논문 수','비중','전략 해석'],[[i+1,n,f'{v:,}',f'{s:.2f}%',('핵심 코어' if i==0 else '상위 보조축' if i<5 else '롱테일 상위')] for i,(n,v,s) in enumerate(PRIMARY_CATS)],[1.5,7.8,2.5,2.4,4.7],8.0)

add_heading(doc,'4.3 분야군 재집계: 계산·소자·응용 축의 병렬 형성')
doc.add_picture(str(FIGS['groups']),width=Cm(18.1)); add_caption(doc,'그림 4-3. 실무 해석용 상위 분야군 분포',True)
add_bullet(doc,'양자물리 코어가 70.00%를 차지하나 컴퓨터과학 9.38%, 응집물질 7.04%, 응용/원자/화학물리 4.79%가 의미 있는 후속 축을 형성함.',True)
add_bullet(doc,'컴퓨터과학은 알고리즘·기계학습·암호·소프트웨어 축, 응집물질은 큐비트 소자·재료·다체계 축, 응용물리는 센싱·광학·원자계 축으로 연결됨.')
add_bullet(doc,'미국 NQI가 양자센싱, 컴퓨팅, 네트워킹, 기초과학, 양자기술을 별도 프로그램 영역으로 정의하는 구조와 본 데이터의 다축 분포가 정합적임 [R1].')
add_bullet(doc,'EU Quantum Europe Strategy가 연구혁신, 인프라, 생태계, 우주·듀얼유스, 인력의 5개 축을 제시한 점은 학술 코어가 산업·인프라·응용 분야로 확장되는 정책적 배경으로 활용함 [R2].')
add_caption(doc,'표 4-3. 분야군별 분포와 실무 기술축')
add_table(doc,['분야군','논문 수','비중','주요 연결 기술'],[[n,f'{v:,}',f'{s:.2f}%',d] for (n,v,s),d in zip(FIELD_GROUPS,['양자정보·기초이론·범용 기반','양자알고리즘·AI·보안·SW','큐비트 소자·물성·나노구조','양자센싱·광학·원자계·화학','기초물리·시뮬레이션·우주/핵','수리모델·통계·최적화','제어·신호처리·시스템 통합','융합 응용·기타'])],[7.0,2.5,2.4,7.0],8.1)

add_heading(doc,'4.4 전체 카테고리 포함 분포: 주분류 밖 연결성 확대')
add_bullet(doc,'전체 카테고리 기준 117개가 확인되어 주 카테고리 87개보다 30개 확대됨.',True)
add_bullet(doc,'양자물리(quant-ph)는 주·보조 포함 1,901건·79.21%에 나타나 비양자 주분류 논문에서도 양자물리가 연결 허브로 기능함.')
doc.add_picture(str(FIGS['all']),width=Cm(18.1)); add_caption(doc,'그림 4-4. 전체 카테고리 포함 기준 TOP2~15',True)
add_bullet(doc,'기계학습(cs.LG)은 주분류 44건에서 전체 포함 143건으로 확대되어 다수 논문에서 보조 방법론·계산 도구로 결합되는 구조를 보임.',True)
add_bullet(doc,'광학(physics.optics)은 주분류 26건에서 전체 포함 103건, 통계역학(cond-mat.stat-mech)은 22건에서 100건으로 확대되어 플랫폼·이론 연결성이 강함.')
add_bullet(doc,'암호/보안(cs.CR) 90건과 인공지능(cs.AI) 83건은 양자기술이 계산·보안 인프라와 결합되는 실질적 응용 신호로 해석함.')
add_caption(doc,'표 4-4. 전체 카테고리 포함 TOP15')
add_table(doc,['순위','전체 카테고리','포함 논문 수','포함 비중'],[[i+1,n,f'{v:,}',f'{s:.2f}%'] for i,(n,v,s) in enumerate(ALL_CATS)],[1.5,10.1,3.4,3.6],8.1)

add_heading(doc,'4.5 복수 카테고리: 단일 분과와 융합 연구의 균형')
doc.add_picture(str(FIGS['counts']),width=Cm(16.8)); add_caption(doc,'그림 4-5. 논문별 부여 카테고리 수 분포',True)
add_bullet(doc,'단일 카테고리 논문은 1,262건·52.58%, 복수 카테고리 논문은 1,138건·47.42%로 거의 균형을 이룸.',True)
add_bullet(doc,f'전체 카테고리 할당 수는 약 {ASSIGNMENTS:,}건, 논문당 평균 카테고리는 {AVG_CATS:.2f}개로 산출됨.')
add_bullet(doc,f'복수 카테고리 논문만 보면 평균 {MULTI_AVG:.2f}개가 부여되어 단순한 2분야 결합을 넘어 3~5개 분야를 포괄하는 고차 융합도 존재함.')
add_bullet(doc,'기술분류체계 구축 시 단일 대분류 귀속보다 주 기술 1개와 보조 기술 복수개를 허용하는 N:M 매핑이 적합함.')
add_caption(doc,'표 4-5. 논문별 카테고리 수와 융합 강도')
add_table(doc,['카테고리 수','논문 수','비중','해석'],[[n,f'{v:,}',f'{s:.2f}%',('단일 분야 중심' if i==0 else '다학제·융합형')] for i,(n,v,s) in enumerate(CAT_COUNTS)],[3.4,3.4,3.4,8.5],8.5)

add_heading(doc,'4.6 정확 복수 조합: 융합기술의 직접 연결 경로')
doc.add_picture(str(FIGS['pairs']),width=Cm(18.1)); add_caption(doc,'그림 4-6. 양자물리 중심 정확 복수 카테고리 조합 TOP10',True)
add_bullet(doc,'정확 조합 1위는 양자물리(quant-ph)+광학(physics.optics) 44건으로 양자광학·포토닉스·정밀계측·양자통신 기반을 대표함.',True)
add_bullet(doc,'양자물리+수리물리 34건은 알고리즘·시뮬레이션·오류정정의 이론 기반, 양자물리+암호/보안 31건은 QKD·PQC·보안 프로토콜의 연결축으로 해석함.')
add_bullet(doc,'양자물리+기계학습 27건은 양자머신러닝뿐 아니라 오류정정 디코더, 회로 최적화, 실험 제어 자동화와 연결될 가능성이 있음.')
add_bullet(doc,f'상위 10개 정확 조합은 총 {TOP10_PAIRS_N}건으로 복수 카테고리 논문의 {TOP10_PAIRS_MULTI_SHARE:.1f}%를 설명하여 상위 융합축과 광범위한 롱테일 조합이 공존함.')
add_caption(doc,'표 4-6. 정확 복수 카테고리 조합 TOP10과 기술적 의미')
meanings=['양자광학·센싱·포토닉스·통신','이론모델·알고리즘·시뮬레이션','양자암호·QKD·PQC·보안','다체계·열역학·최적화','양자머신러닝·디코더·회로최적화','양자물성·큐비트 플랫폼','신흥 양자컴퓨팅·시스템','양자화학·분자·소재 시뮬레이션','메조스코픽 소자·나노구조','원자계 센싱·시계·계측']
add_table(doc,['순위','정확 조합','논문 수','전체 비중','기술적 의미'],[[i+1,n,f'{v}',f'{s:.2f}%',meanings[i]] for i,(n,v,s) in enumerate(PAIR_EXACT)],[1.3,7.3,2.4,2.6,5.8],8.0)

add_heading(doc,'4.7 동시출현 교차 신호: quant-ph 허브와 주변 분야')
doc.add_picture(str(FIGS['cooccur']),width=Cm(18.1)); add_caption(doc,'그림 4-7. 양자물리(quant-ph) 동시 부여 교차 카테고리 TOP10',True)
add_bullet(doc,'통계역학(cond-mat.stat-mech) 90건, 기계학습(cs.LG) 86건, 광학(physics.optics) 80건이 가장 강한 교차 신호로 확인됨.',True)
add_bullet(doc,'통계역학은 양자다체계·열역학·동역학, 기계학습은 데이터 기반 모델링·디코딩·제어, 광학은 광자·센싱·네트워크 플랫폼과 연결함.')
add_bullet(doc,'NIST의 FIPS 203·204·205 승인으로 양자내성암호 전환이 실제 표준 단계로 이동하여 암호/보안(cs.CR) 교차 신호의 정책·산업적 중요성이 커짐 [R3].')
add_bullet(doc,'DARPA QBI가 유틸리티 규모 양자컴퓨터의 비용 대비 계산가치와 검증을 강조함에 따라 수리물리·기계학습·소프트웨어·벤치마킹 카테고리의 결합을 별도 추적할 필요가 있음 [R4].')
add_caption(doc,'표 4-7. 정확 조합과 동시출현 분석의 차이')
add_table(doc,['구분','집계 방식','장점','활용'],[
    ['정확 조합','논문의 전체 카테고리 집합이 동일한 경우','구체적인 융합 유형을 명확히 보여줌','대표 융합기술 후보 선정'],
    ['동시출현','quant-ph와 특정 카테고리가 함께 존재하면 집계','3개 이상 복합 조합까지 포함하여 허브 연결성을 보여줌','교차 네트워크와 브리지 분야 탐색'],
], [3.0,5.7,5.2,5.2],8.4)

add_heading(doc,'4.8 집중도와 롱테일 구조')
add_bullet(doc,'주 카테고리 TOP1 누적비중은 70.00%, TOP3는 74.71%, TOP5는 77.92%, TOP10은 83.75%로 확인됨.',True)
add_bullet(doc,'HHI 0.493은 강한 집중 구조를 의미하지만 TOP20 이후에도 67개 주 카테고리가 존재하여 응용 확장은 롱테일에서 탐지될 가능성이 큼.')
add_bullet(doc,'따라서 전략분석은 양자물리 코어의 세부 토픽을 깊게 분석하는 축과 주변 카테고리의 부상·교차 신호를 빠르게 탐지하는 축을 병행해야 함.')
add_caption(doc,'표 4-8. 분야 집중도와 확장성 핵심 지표')
add_table(doc,['지표','값','해석'],[
    ['주 / 전체 카테고리','87개 / 117개','보조 분류를 포함하면 연결 분야가 30개 확대됨'],['TOP1 / TOP5 / TOP10','70.00% / 77.92% / 83.75%','코어 집중이 강하나 상위 10개 밖 롱테일도 16.25% 존재함'],['HHI / 유효 카테고리 수',f'{HHI:.3f} / {EFFECTIVE_CATS:.1f}개','실질 분포가 quant-ph 중심으로 압축됨'],['복수 카테고리','1,138건·47.42%','집중 구조 내부에서 논문 단위 융합성은 높음'],['논문당 평균 카테고리',f'{AVG_CATS:.2f}개','단일 주분류만으로 기술 의미를 설명하기 어려움'],
], [4.4,4.2,9.3],8.4)

add_heading(doc,'4.9 전·후반기 변화: 코어 확대와 소규모 부상 신호')
doc.add_picture(str(FIGS['growth']),width=Cm(18.1)); add_caption(doc,'그림 4-8. 주요 카테고리 전·후반 6개월 변화',True)
add_bullet(doc,'양자물리(quant-ph)는 전반 824건에서 후반 856건으로 +32건·+3.9% 증가하여 코어 중심성이 유지되면서 절대 규모도 확대됨.',True)
add_bullet(doc,'통계역학(cond-mat.stat-mech)은 6건에서 16건, 정보이론(cs.IT)은 3건에서 10건으로 증가하여 절대 규모는 작지만 부상도 관점에서 우선 검토 가치가 있음.')
add_bullet(doc,'암호/보안(cs.CR)은 20건에서 23건으로 +15.0% 증가하여 PQC·QKD·양자보안 전환 이슈와 정합적임.')
add_bullet(doc,'소규모 카테고리의 높은 증가율은 기저효과가 크므로 절대 증가량, 최근 3개월 모멘텀, 키워드·기관·국가 변화를 함께 검증해야 함.')
add_caption(doc,'표 4-9. 주요 증가 카테고리의 전·후반기 비교')
add_table(doc,['카테고리','전반 6개월','후반 6개월','증감','증감률','해석'],[[n,b,a,f'{d:+d}',f'{r:+.1f}%',('코어 확대' if i==0 else '소규모 부상 신호')] for i,(n,b,a,d,r) in enumerate(GROWTH)],[6.4,2.5,2.5,2.0,2.3,3.8],7.9)

add_heading(doc,'4.10 최신 시장·기술·정책·표준화 환경과의 통합 해석')
add_bullet(doc,'미국 NQI는 양자센싱·계측, 양자컴퓨팅, 양자네트워킹, 기초과학, 양자기술을 별도 프로그램 영역으로 정의하고 지속적인 연방 투자를 유지함. 이는 분야군을 물리 코어와 계산·네트워크·센싱·지원기술로 분리하는 분석 논리와 부합함 [R1].',True)
add_bullet(doc,'EU Quantum Europe Strategy는 연구혁신, 양자인프라, 생태계·공급망, 우주·듀얼유스, 인력의 5개 영역을 제시함. 광학·응집물질·컴퓨터과학·보안 카테고리의 병렬 성장은 연구에서 산업·인프라로 확장되는 정책 방향과 연결됨 [R2].')
add_bullet(doc,'NIST의 PQC 표준은 암호/보안(cs.CR)과 양자물리(quant-ph)의 교차를 양자컴퓨터 연구뿐 아니라 현재 시스템의 양자안전 전환이라는 실무 과제로 확장함 [R3].')
add_bullet(doc,'Google Quantum AI의 below-threshold 표면코드 실증과 IBM의 2026년 양자 이점·2029년 내결함성 로드맵은 수리물리, 통계역학, 기계학습, 소프트웨어, 제어·신호처리의 결합 수요를 강화하는 기술 배경임 [R5][R6].')
add_bullet(doc,'McKinsey는 2026년 양자기술 모니터에서 양자컴퓨팅·통신·센싱을 3대 영역으로 추적하고 300개 이상의 기업 채택·협업을 제시함. 이는 학술 카테고리를 실무 기술군으로 재매핑할 필요성을 뒷받침함 [R7].')
add_callout(doc,'정량–정성 통합 판단',[
    '70.00%의 양자물리 코어는 원천 이론과 공통 과학기반의 집중을 의미함.',
    '47.42%의 복수 카테고리는 산업화에 필요한 소프트웨어·소재·광학·보안·시스템 결합이 이미 학술 분류에 반영되고 있음을 보임.',
    '분야별 시장규모를 직접 추정하기보다 교차 카테고리와 외부 로드맵을 이용해 기술 전환 경로와 후속 분석 우선순위를 판단함.',
],LIGHT_ORANGE,ORANGE)

add_heading(doc,'4.11 학술 카테고리의 실무형 양자기술 재매핑')
doc.add_picture(str(FIGS['mapping']),width=Cm(18.1)); add_caption(doc,'그림 4-9. arXiv 카테고리의 실무형 양자기술 N:M 재매핑 프레임',True)
add_bullet(doc,'양자컴퓨팅·SW/AI는 quant-ph, cs.LG, cs.AI, cs.ET, eess.SP를 결합하여 알고리즘·오류정정·디코더·컴파일러·제어를 포괄함.',True)
add_bullet(doc,'양자통신·보안은 quant-ph, physics.optics, cs.CR을 결합하여 광자통신·QKD·PQC·네트워크 프로토콜을 포괄함.')
add_bullet(doc,'양자센싱·계측은 quant-ph, physics.optics, physics.atom-ph, eess.SP를 결합하여 시계·자기장·관성·광계측·신호처리를 포괄함.')
add_bullet(doc,'양자소자·소재는 cond-mat.mes-hall, cond-mat.str-el, cond-mat.mtrl-sci를 중심으로 큐비트 소자·재료·공정·패키징으로 재구성함.')
add_bullet(doc,'양자이론·시뮬레이션과 양자화학·응용은 math-ph, stat-mech, hep-th, physics.chem-ph를 목적에 따라 중복 매핑함.')
add_caption(doc,'표 4-10. 실무 기술군별 권장 매핑과 후속 분석 변수')
add_table(doc,['실무 기술군','주요 arXiv 카테고리','후속 키워드·지표','외부 연계'],[
    ['양자컴퓨팅·SW/AI','quant-ph, cs.LG, cs.AI, cs.ET, eess.SP','algorithm, circuit, error correction, decoder, compiler, benchmark','기업 로드맵·DARPA QBI'],
    ['양자통신·보안','quant-ph, physics.optics, cs.CR','QKD, repeater, network, PQC, protocol','NIST·통신 표준'],
    ['양자센싱·계측','quant-ph, physics.optics, physics.atom-ph, eess.SP','clock, magnetic field, inertia, metrology, readout','국가 센싱 전략·시험인증'],
    ['양자소자·소재','mes-hall, str-el, mtrl-sci, supr-con','qubit material, junction, defect, fabrication, cryogenic','반도체·소재 공급망'],
    ['양자이론·시뮬레이션','math-ph, stat-mech, hep-th, physics.comp-ph','many-body, simulation, optimization, dynamics','기초과학·HPC'],
    ['양자화학·응용','physics.chem-ph, math-ph, materials','molecule, catalyst, drug, energy, material discovery','제약·소재·에너지 시장'],
], [3.6,6.2,6.4,3.5],7.7)

add_heading(doc,'4.12 전략적 시사점 및 후속 장 연결')
add_bullet(doc,'기술전략은 양자물리 코어를 단일 분야로만 관리하지 않고 컴퓨팅·통신·센싱·소자·소재·보안의 교차축을 별도 포트폴리오로 구성해야 함.',True)
add_bullet(doc,'유망기술 후보는 절대 건수 상위와 증가율 상위를 분리한 뒤 카테고리 조합, 최근 3개월 모멘텀, 키워드 신규성, 주요 기관·국가를 결합하여 선정해야 함.')
add_bullet(doc,'제5장에서는 양자물리 코어, 컴퓨터과학, 응집물질, 광학·응용물리의 상위 기관·국가가 서로 다른지 비교함.')
add_bullet(doc,'제6장에서는 복수 카테고리 수가 증가할수록 공저 규모, 다기관 비율, 국제협력률이 높아지는지 검증함.')
add_bullet(doc,'제7장에서는 본 장의 교차 구조를 양자컴퓨팅, 양자통신·보안, 양자센싱, 양자소자·소재, 양자SW/AI의 전략축으로 통합함.')
add_caption(doc,'표 4-11. 본 장의 핵심 발견과 후속 분석 과제')
add_table(doc,['핵심 발견','해석','후속 분석'],[
    ['quant-ph 70.00%','명확한 과학 코어와 연구 커뮤니티가 존재함','코어 내부 토픽·성능·플랫폼을 세분화함'],['복수 카테고리 47.42%','다학제 결합이 보편적임','N:M 기술분류와 교차 네트워크를 구축함'],['cs.LG·cs.CR·optics 확대','AI·보안·센싱·통신 응용축이 형성됨','최근 성장·기관·국가·특허를 결합함'],['HHI 0.493와 67개 롱테일','코어 집중과 주변 다양성이 공존함','상위 분야와 롱테일 부상도를 이원화함'],['소규모 고성장 카테고리','기저효과가 있으나 조기신호 가치가 있음','3개월 모멘텀·키워드·전문가 검토로 검증함'],
], [4.6,7.0,8.1],8.2)

add_heading(doc,'4.13 분석 한계 및 검증 과제')
add_bullet(doc,'arXiv 카테고리는 저자의 제출 선택과 운영 분류체계에 의존하므로 실제 산업기술의 세부 경계를 완전하게 반영하지 않음.',True)
add_bullet(doc,'검색식에 QUANTIZATION이 포함되어 컴퓨터비전·기계학습 분야의 모델 양자화 논문이 일부 유입될 수 있으므로 초록·전문 기반 정합성 검증이 필요함.')
add_bullet(doc,'카테고리 성장률은 소규모 분야의 기저효과가 크므로 최소 건수 기준, Z-score, 최근 3개월 지속성 조건을 함께 적용해야 함.')
add_bullet(doc,'정확 조합과 동시출현은 집계 의미가 다르므로 동일 순위표로 혼합하지 않고 조합 구조와 허브 연결성을 별도로 해석해야 함.')
add_caption(doc,'표 4-12. 카테고리 분석의 한계와 보완 방향')
add_table(doc,['한계','영향','보완 방향'],[
    ['학술분류와 산업분류 불일치','실무 기술군의 규모가 과대·과소 추정될 수 있음','본문 키워드·임베딩·전문가 정의서를 결합하여 재분류함'],['QUANTIZATION 노이즈','AI 모델 압축 논문이 양자기술로 오인될 수 있음','문맥 규칙과 지도학습 분류로 관련성을 검증함'],['짧은 분석기간','장기 성장과 일시적 변동을 구분하기 어려움','24~36개월 확장과 분기별 모니터링을 수행함'],['소규모 증가율','기저효과로 부상도가 과장될 수 있음','절대 증가량·최근 지속성·외부 이슈를 함께 평가함'],['저자 선택 분류','동일 기술의 분류가 논문별로 달라질 수 있음','복수 카테고리·키워드·인용 네트워크를 통합함'],
], [4.5,6.3,8.9],8.2)

add_heading(doc,'4.14 본 장 소결')
add_bullet(doc,'주 카테고리 87개 중 양자물리(quant-ph)가 1,680건·70.00%를 차지하여 분석 DB의 과학 코어가 명확함.',True)
add_bullet(doc,'전체 카테고리는 117개, 복수 카테고리 논문은 1,138건·47.42%로 확인되어 코어 집중 안에서 다학제 확장이 활발함.')
add_bullet(doc,'광학·수리물리·암호/보안·기계학습·응집물질과의 결합은 양자센싱·통신·보안·컴퓨팅·소자·소재의 산업화 경로를 보여줌.')
add_bullet(doc,'전·후반기 비교에서 quant-ph는 +32건 증가하고 통계역학·정보이론·암호/보안 등은 소규모 부상 신호를 보임.')
add_bullet(doc,'다음 제5장에서는 본 장의 분야축을 기준으로 상위 저자·기관·국가의 전문화와 연구 주도·참여 구조를 분석함.')

add_heading(doc,'참고자료')
refs=[
    ('[R1]','National Quantum Initiative, Annual Report FY2025','https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf'),
    ('[R2]','European Commission, Quantum Europe Strategy, 2025','https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy'),
    ('[R3]','NIST CSRC, Post-Quantum Cryptography FIPS Approved','https://csrc.nist.gov/news/2024/postquantum-cryptography-fips-approved'),
    ('[R4]','DARPA, Quantum Benchmarking Initiative','https://www.darpa.mil/research/programs/quantum-benchmarking-initiative'),
    ('[R5]','Nature, Quantum Error Correction Below the Surface Code Threshold','https://www.nature.com/articles/s41586-024-08449-y'),
    ('[R6]','IBM Quantum, Hardware and Roadmap','https://www.ibm.com/quantum/hardware'),
    ('[R7]','McKinsey, Quantum Technology Monitor 2026','https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/mckinsey-quantum-technology-monitor-2026-a-commercial-tipping-point'),
]
for rid,title,url in refs:
    p=doc.add_paragraph(); pformat(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(.42); p.paragraph_format.first_line_indent=Cm(-.42); r=p.add_run(rid+' '); set_run(r,8.8,True,NAVY); add_hyperlink(p,title,url)

cp=doc.core_properties; cp.title='arXiv 양자 분야 제4장 분야·카테고리 분포 분석'; cp.subject='2025년 6월~2026년 5월, 12개월·2,400건 기준 정량·정성 통합 분석'; cp.author='OpenAI'; cp.keywords='arXiv, 양자, 카테고리, 융합, 시장, 기술, 정책, 표준화'
settings=doc.settings._element; update=OxmlElement('w:updateFields'); update.set(qn('w:val'),'true'); settings.append(update)
doc.save(OUTPUT); print(f'Created {OUTPUT}')

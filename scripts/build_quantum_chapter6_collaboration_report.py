from __future__ import annotations

from pathlib import Path
import shutil
import subprocess

import numpy as np
import matplotlib.pyplot as plt
from matplotlib import font_manager

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path('artifacts')
ASSET_DIR = OUT_DIR / 'chapter6_assets'
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / 'arXiv_양자분야_제6장_협업구조분석_고품질_재작성_20260704.docx'

months = ['2025-06','2025-07','2025-08','2025-09','2025-10','2025-11','2025-12','2026-01','2026-02','2026-03','2026-04','2026-05']
paper_counts = np.array([182,187,183,219,229,171,227,186,165,211,220,220], dtype=float)
avg_authors = np.array([9.42,4.26,4.66,4.76,5.77,5.53,5.47,4.77,8.90,4.56,4.01,3.98])
coauthor_rate = np.array([94.0,87.2,86.9,90.4,90.4,92.4,84.1,88.7,87.9,87.2,86.4,82.7])
multi_inst_rate = np.array([22.0,29.9,25.7,36.1,30.6,31.0,30.0,31.2,33.3,28.9,30.0,25.9])
intl_rate = np.array([7.7,11.2,5.5,11.0,14.8,12.3,11.9,10.8,11.5,9.0,8.2,6.4])

team_labels = ['1명','2명','3명','4명','5명','6명','7명','8명','9명','10명 이상']
team_counts = np.array([287,435,452,397,230,167,123,91,48,170])
team_shares = team_counts / 2400 * 100

country_pairs = [
    ('독일↔미국',14),('인도↔미국',14),('영국↔미국',14),('중국↔홍콩',12),
    ('중국↔미국',10),('독일↔이탈리아',9),('중국↔싱가포르',8),
    ('중국↔영국',8),('스페인↔미국',8),('스위스↔미국',7)
]
hub_countries = [
    ('미국',133,93),('독일',79,54),('영국',73,51),('중국',66,44),
    ('이탈리아',54,28),('인도',44,22),('스페인',35,25),('폴란드',30,12),
    ('프랑스',30,24),('네덜란드',26,13)
]
institution_pairs = [
    ('MIT↔Mit',5),('Ahmedabad Univ.↔Univ. of Oklahoma',3),
    ('Beihang Univ.↔Normal Univ.',3),('Hunan Univ.↔Univ. of Technology',3),
    ('Imperial↔Imperial College London',3),('Institut quantique↔U Sherbrooke',3),
    ('Jilin Univ.↔Normal Univ.',3),('Lebedev Physical Institute↔Moscow',3),
    ('Nanjing Univ.↔Normal Univ.',3),('Normal Univ.↔South Univ.',3),
]

TOTAL=2400; COAUTHOR_N=2113; SINGLE_N=287; MULTI_INST_N=710; INTL_N=241
AVG_AUTHOR=5.41; MEDIAN_AUTHOR=4; AVG_INST=1.44; MEDIAN_INST=1; COUNTRY_VALID_N=1998
P90=8; P95=11; P99=22; MAX_AUTHORS=732
weighted_coauthor_first=np.average(coauthor_rate[:6],weights=paper_counts[:6])
weighted_coauthor_second=np.average(coauthor_rate[6:],weights=paper_counts[6:])
weighted_multi_first=np.average(multi_inst_rate[:6],weights=paper_counts[:6])
weighted_multi_second=np.average(multi_inst_rate[6:],weights=paper_counts[6:])
weighted_intl_first=np.average(intl_rate[:6],weights=paper_counts[:6])
weighted_intl_second=np.average(intl_rate[6:],weights=paper_counts[6:])
corr_volume_coauthor=float(np.corrcoef(paper_counts,coauthor_rate)[0,1])
corr_volume_multi=float(np.corrcoef(paper_counts,multi_inst_rate)[0,1])
corr_volume_intl=float(np.corrcoef(paper_counts,intl_rate)[0,1])

NAVY='17365D'; BLUE='2F75B5'; LIGHT_BLUE='D9EAF7'; GREEN='548235'; LIGHT_GREEN='E2F0D9'; ORANGE='ED7D31'; LIGHT_ORANGE='FCE4D6'; PURPLE='7030A0'; LIGHT_PURPLE='E4DFEC'; GRAY='6B7280'; LIGHT_GRAY='F2F4F7'; DARK='243447'; WHITE='FFFFFF'; BORDER='CAD4DF'; YELLOW='FFD966'

def rgb(h): return RGBColor.from_string(h)

font_candidates=['/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc','/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc','/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc']
font_path=next((p for p in font_candidates if Path(p).exists()),None)
if font_path:
    fp=font_manager.FontProperties(fname=font_path); plt.rcParams['font.family']=fp.get_name()
else:
    plt.rcParams['font.family']='DejaVu Sans'
plt.rcParams['axes.unicode_minus']=False

# --- charts: use matplotlib default colors ---
def finish(fig,path):
    fig.tight_layout(pad=1.0); fig.savefig(path,dpi=220,bbox_inches='tight'); plt.close(fig)

def chart_funnel(path):
    labels=['전체 논문','공동저자','다기관','국제협력']; values=[100,COAUTHOR_N/TOTAL*100,MULTI_INST_N/TOTAL*100,INTL_N/TOTAL*100]; counts=[TOTAL,COAUTHOR_N,MULTI_INST_N,INTL_N]
    fig,ax=plt.subplots(figsize=(10.8,4.8)); y=np.arange(4); bars=ax.barh(y,values); ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold'); ax.invert_yaxis(); ax.set_xlim(0,108); ax.set_xlabel('전체 논문 대비 비율(%)',fontweight='bold'); ax.set_title('협업 구조의 단계별 축소: 연구자→기관→국가',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='x',alpha=.25)
    for b,v,n in zip(bars,values,counts): ax.text(v+1.2,b.get_y()+b.get_height()/2,f'{n:,}건 · {v:.1f}%',va='center',fontsize=10,fontweight='bold',clip_on=False)
    finish(fig,path)

def chart_team(path):
    fig,ax=plt.subplots(figsize=(11.2,5.2)); x=np.arange(len(team_labels)); bars=ax.bar(x,team_counts); ax.set_xticks(x); ax.set_xticklabels(team_labels,fontweight='bold'); ax.set_ylabel('논문 수(건)',fontweight='bold'); ax.set_title('공저 규모 분포',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='y',alpha=.25); ax.set_ylim(0,500)
    for b,n,s in zip(bars,team_counts,team_shares): ax.text(b.get_x()+b.get_width()/2,n+9,f'{n}\n({s:.1f}%)',ha='center',fontsize=9,fontweight='bold',clip_on=False)
    ax.text(2.5,470,'2~4인 연구 1,284건 · 53.5%',ha='center',fontsize=10.5,fontweight='bold'); finish(fig,path)

def chart_monthly(path):
    fig,ax=plt.subplots(figsize=(12,5.4)); x=np.arange(12); ax.plot(x,coauthor_rate,marker='o',linewidth=2.2,label='공동저자 비율'); ax.plot(x,multi_inst_rate,marker='s',linewidth=2.2,label='다기관 비율'); ax.plot(x,intl_rate,marker='^',linewidth=2.2,label='국제협력 비율'); ax.set_xticks(x); ax.set_xticklabels([m.replace('2025-','25.').replace('2026-','26.') for m in months],fontweight='bold'); ax.set_ylabel('비율(%)',fontweight='bold'); ax.set_ylim(0,100); ax.set_title('월별 공동저자·다기관·국제협력 비율',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='y',alpha=.25); ax.legend(frameon=False,ncol=3,loc='lower left')
    for i in [0,4,8,11]:
        ax.text(i,coauthor_rate[i]+2.2,f'{coauthor_rate[i]:.1f}',ha='center',fontsize=8.5,fontweight='bold'); ax.text(i,multi_inst_rate[i]+2.2,f'{multi_inst_rate[i]:.1f}',ha='center',fontsize=8.5,fontweight='bold'); ax.text(i,intl_rate[i]+2.2,f'{intl_rate[i]:.1f}',ha='center',fontsize=8.5,fontweight='bold')
    finish(fig,path)

def chart_avg_authors(path):
    fig,ax=plt.subplots(figsize=(11.8,4.8)); x=np.arange(12); ax.plot(x,avg_authors,marker='o',linewidth=2.2); ax.axhline(AVG_AUTHOR,linestyle='--',linewidth=1.5,label=f'전체 평균 {AVG_AUTHOR:.2f}명'); ax.set_xticks(x); ax.set_xticklabels([m.replace('2025-','25.').replace('2026-','26.') for m in months],fontweight='bold'); ax.set_ylabel('평균 저자 수(명)',fontweight='bold'); ax.set_ylim(3,10.5); ax.set_title('월별 평균 저자 수와 대형 공동연구 영향',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='y',alpha=.25); ax.legend(frameon=False,loc='upper right')
    for i in [0,4,8,11]: ax.text(i,avg_authors[i]+.25,f'{avg_authors[i]:.2f}',ha='center',fontsize=9,fontweight='bold')
    finish(fig,path)

def chart_pairs(path):
    labels=[x[0] for x in country_pairs][::-1]; vals=[x[1] for x in country_pairs][::-1]; fig,ax=plt.subplots(figsize=(11.5,5.8)); y=np.arange(10); bars=ax.barh(y,vals); ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold'); ax.set_xlabel('공동 논문 수(건)',fontweight='bold'); ax.set_xlim(0,16.5); ax.set_title('국가 간 국제협력쌍 TOP10',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='x',alpha=.25)
    for b,v in zip(bars,vals): ax.text(v+.25,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9.5,fontweight='bold')
    finish(fig,path)

def chart_hubs(path):
    labels=[x[0] for x in hub_countries]; weighted=np.array([x[1] for x in hub_countries]); docs=np.array([x[2] for x in hub_countries]); x=np.arange(10); w=.38; fig,ax=plt.subplots(figsize=(12,5.4)); b1=ax.bar(x-w/2,weighted,w,label='가중 연결 수'); b2=ax.bar(x+w/2,docs,w,label='국제협력 논문 수'); ax.set_xticks(x); ax.set_xticklabels(labels,fontweight='bold'); ax.set_ylabel('건수',fontweight='bold'); ax.set_title('국제협력 허브 국가: 연결 강도와 참여 논문',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='y',alpha=.25); ax.legend(frameon=False); ax.set_ylim(0,145)
    for bars in (b1,b2):
        for b in bars: ax.text(b.get_x()+b.get_width()/2,b.get_height()+2,f'{int(b.get_height())}',ha='center',fontsize=8.5,fontweight='bold')
    finish(fig,path)

def chart_inst_pairs(path):
    labels=[x[0] for x in institution_pairs][::-1]; vals=[x[1] for x in institution_pairs][::-1]; fig,ax=plt.subplots(figsize=(12,6.2)); y=np.arange(10); bars=ax.barh(y,vals); ax.set_yticks(y); ax.set_yticklabels(labels,fontweight='bold',fontsize=8.7); ax.set_xlabel('추출된 공동 논문 수(건)',fontweight='bold'); ax.set_xlim(0,6.2); ax.set_title('기관 협업쌍 TOP10: 기관명 표준화 전 탐색값',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='x',alpha=.25)
    for b,v in zip(bars,vals): ax.text(v+.12,b.get_y()+b.get_height()/2,f'{v}건',va='center',fontsize=9.3,fontweight='bold')
    ax.text(3,-1.35,'동일기관 표기 변형이 협업쌍으로 잡힌 사례를 포함하므로 확정 순위가 아님',ha='center',fontsize=9.5,fontweight='bold'); finish(fig,path)

def chart_half(path):
    labels=['공동저자','다기관','국제협력']; first=[weighted_coauthor_first,weighted_multi_first,weighted_intl_first]; second=[weighted_coauthor_second,weighted_multi_second,weighted_intl_second]; x=np.arange(3); w=.35; fig,ax=plt.subplots(figsize=(9.8,5)); b1=ax.bar(x-w/2,first,w,label='전반 6개월'); b2=ax.bar(x+w/2,second,w,label='후반 6개월'); ax.set_xticks(x); ax.set_xticklabels(labels,fontweight='bold'); ax.set_ylabel('논문수 가중 비율(%)',fontweight='bold'); ax.set_title('전·후반 6개월 협업 강도 비교',fontweight='bold',fontsize=15,pad=16); ax.grid(axis='y',alpha=.25); ax.legend(frameon=False); ax.set_ylim(0,100)
    for bars in (b1,b2):
        for b in bars: ax.text(b.get_x()+b.get_width()/2,b.get_height()+1.2,f'{b.get_height():.1f}',ha='center',fontsize=9.3,fontweight='bold')
    finish(fig,path)

FIGS={'funnel':ASSET_DIR/'fig_6_1_funnel.png','team':ASSET_DIR/'fig_6_2_team.png','monthly':ASSET_DIR/'fig_6_3_monthly.png','authors':ASSET_DIR/'fig_6_4_avg_authors.png','pairs':ASSET_DIR/'fig_6_5_country_pairs.png','hubs':ASSET_DIR/'fig_6_6_hubs.png','inst':ASSET_DIR/'fig_6_7_inst_pairs.png','half':ASSET_DIR/'fig_6_8_half.png'}
chart_funnel(FIGS['funnel']); chart_team(FIGS['team']); chart_monthly(FIGS['monthly']); chart_avg_authors(FIGS['authors']); chart_pairs(FIGS['pairs']); chart_hubs(FIGS['hubs']); chart_inst_pairs(FIGS['inst']); chart_half(FIGS['half'])

# --- Word helpers ---
def set_cell_shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_border(cell,color=BORDER,size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ['top','left','bottom','right']:
        el=borders.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)

def set_cell_margins(cell,top=80,start=90,bottom=80,end=90):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_run(run,size=10,bold=False,color=DARK):
    run.font.name='맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=rgb(color)

def pformat(p,before=0,after=4,line=1.25,keep=False):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line; pf.keep_with_next=keep

def add_page_number(p):
    run=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text='PAGE'; e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); run._r.extend([b,i,e])

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); pformat(p,before=12 if level==1 else 8,after=6,line=1.0,keep=True); r=p.add_run(text); set_run(r,15 if level==1 else 12,True,NAVY)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'3'); bottom.set(qn('w:color'),BLUE); pBdr.append(bottom); pPr.append(pBdr)
    return p

def add_bullet(doc,text,major=False):
    p=doc.add_paragraph(); pformat(p,after=3,line=1.28); p.paragraph_format.left_indent=Cm(.48); p.paragraph_format.first_line_indent=Cm(-.34); r=p.add_run('□ ' if major else '· '); set_run(r,10.2,True,BLUE if major else GRAY); r=p.add_run(text); set_run(r,10.2,major,DARK); return p

def add_caption(doc,text,figure=False):
    p=doc.add_paragraph(); pformat(p,before=3,after=6,line=1.0,keep=not figure); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); set_run(r,9.2,True,GRAY if figure else NAVY)

def add_table(doc,headers,rows,widths,font=8.0):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,NAVY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,after=0,line=1.05); r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cell=cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,WHITE if ri%2==0 else LIGHT_GRAY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pformat(p,after=0,line=1.08); r=p.add_run(str(v)); set_run(r,font,i==0,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0); return table

def add_callout(doc,title,lines,fill=LIGHT_BLUE,accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,accent,'8'); set_cell_margins(cell,130,150,130,150); p=cell.paragraphs[0]; pformat(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=cell.add_paragraph(); pformat(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.first_line_indent=Cm(-.25); r=p.add_run('• '); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_hyperlink(p,text,url):
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rPr.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)

# --- document ---
doc=Document(); sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35); sec.header_distance=Cm(.55); sec.footer_distance=Cm(.55)
style=doc.styles['Normal']; style.font.name='맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); style.font.size=Pt(10)
header=sec.header.paragraphs[0]; r=header.add_run('arXiv 양자 분야 논문 분석 | 제6장 협업 구조 분석'); set_run(r,8.4,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('분석 기준: 2025.06~2026.05, 12개월·2,400건  |  '); set_run(r,8.2,False,GRAY); add_page_number(footer)

cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; cell=cover.cell(0,0); set_cell_shading(cell,NAVY); set_cell_margins(cell,720,300,720,300); set_cell_border(cell,NAVY,'0'); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CHAPTER 6'); set_run(r,14,True,'9DC3E6'); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('협업 구조 분석'); set_run(r,27,True,WHITE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('2025년 6월~2026년 5월 | 12개월·2,400건'); set_run(r,12,True,LIGHT_BLUE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,before=16); r=p.add_run('공동저자 보편화 × 선택적 다기관·국제협력 × 허브국가 중심 네트워크'); set_run(r,10.8,True,YELLOW)

doc.add_paragraph(); kpis=[('88.0%','공동저자 연구'),('5.41명','평균 저자 수'),('29.6%','다기관 연구'),('10.0%','국제협력'),('7.1%','10명 이상 대형팀')]; t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(val,label) in enumerate(kpis):
    cell=t.cell(0,i); fill=[LIGHT_BLUE,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]; accent=[BLUE,GREEN,ORANGE,PURPLE,GRAY][i]; set_cell_shading(cell,fill); set_cell_border(cell,accent,'7'); set_cell_margins(cell,100,70,100,70); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,15,True,NAVY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_run(r,8.6,True,GRAY)
add_callout(doc,'본 장의 핵심 판단',['양자 분야의 기본 수행 방식은 공동저자 연구이나, 기관과 국가 수준의 협업은 선택적으로 형성되는 다층 구조로 판단함.','2~4인 연구가 53.5%를 차지하여 소·중규모 팀이 중심이며, 10명 이상 대형팀 7.1%가 실험·장비·플랫폼·컨소시엄형 연구를 보완함.','국제협력은 전체의 10.0%로 제한적이나 미국을 중심으로 독일·영국·인도·중국 등과 연결되는 허브형 네트워크가 확인됨.'],LIGHT_BLUE,BLUE)
doc.add_page_break()

add_heading(doc,'6.1 분석 목적 및 앞선 장과의 연결')
for txt,major in [
('제3장에서 확인한 월평균 200편의 연구 생산이 단독 연구, 소규모 공저, 다기관 연구, 국제 공동연구 중 어떠한 수행 구조를 통해 형성되는지 분석함.',True),
('제4장에서 확인한 양자물리(quant-ph) 코어와 컴퓨터과학·응집물질·광학·암호/보안 확장이 실제로 더 큰 팀과 기관·국가 간 협력을 요구하는지 해석함.',False),
('제5장에서 확인한 분산형 연구자·기관 기반과 미국·중국·독일·영국·일본 중심의 국가축이 협업 네트워크에서는 어떤 허브 구조로 나타나는지 연결함.',False),
('본 장의 결과는 제7장에서 기술군별 국제협력 전략, 공동연구 파트너 선정, 연구인프라·표준화 연계, 국내 협력역량 강화 방안으로 전환함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-1. 협업 구조 분석 단위와 지표')
add_table(doc,['분석 축','판정 기준','핵심 지표','해석 목적'],[
['공저 구조','저자 수 1명 / 2명 이상','단독·공동 비율, 평균·중앙값, 공저 규모','연구 수행의 기본 팀 구조를 파악함'],['기관 협업','정제 기관 수 1개 / 2개 이상','다기관 비율, 평균 기관 수, 기관쌍','인프라·장비·전문성 결합 수준을 진단함'],['국가 협업','기관 국가 1개 / 2개 이상','국제협력률, 국가쌍, 가중 연결 수','국제 개방성·허브국가·파트너를 탐색함'],['월별 변화','월별 논문과 협업 지표 재집계','공동저자·다기관·국제협력 비율','연구량과 협업 강도의 동행 여부를 확인함'],['분야별 차이','주 카테고리별 협업 특성 비교','평균 저자, 다기관·국제협력 상대수준','이론·실험·응용 기술군의 협업 요구를 해석함']],[3.2,5.0,5.6,5.2],8.0)

add_heading(doc,'6.2 협업 구조의 단계별 축소'); doc.add_picture(str(FIGS['funnel']),width=Cm(17.8)); add_caption(doc,'그림 6-1. 협업 구조의 단계별 축소: 전체 논문→공동저자→다기관→국제협력',True)
for txt,major in [
(f'공동저자 연구는 {COAUTHOR_N:,}건·{COAUTHOR_N/TOTAL*100:.1f}%로 확인되어 팀 연구가 양자 분야의 기본 수행 방식으로 정착된 것으로 보임.',True),
(f'다기관 연구는 {MULTI_INST_N:,}건·{MULTI_INST_N/TOTAL*100:.1f}%로 공동저자 연구의 약 {MULTI_INST_N/COAUTHOR_N*100:.1f}%에 해당하여 다수의 공저가 동일 기관 내부에서 이루어짐.',False),
(f'국제협력 연구는 {INTL_N:,}건·{INTL_N/TOTAL*100:.1f}%이며 기관 국가 확인 논문 {COUNTRY_VALID_N:,}건 기준으로는 {INTL_N/COUNTRY_VALID_N*100:.1f}%로 산출됨.',False),
('연구자 수준의 협업은 보편적이나 기관·국가 경계를 넘는 협업은 장비·인프라·실증·표준·정책 조건이 충족되는 경우 선택적으로 형성되는 구조로 판단함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-2. 협업 구조 핵심 지표')
add_table(doc,['지표','값','보조 해석','전략적 의미'],[
['분석 논문',f'{TOTAL:,}건','2026년 6월 부분월 27건 제외','모든 협업 지표의 공통 분모로 사용함'],['공동저자',f'{COAUTHOR_N:,}건·{COAUTHOR_N/TOTAL*100:.1f}%','단독 연구 287건·12.0%','팀 기반 연구가 보편적임'],['평균 / 중앙 저자',f'{AVG_AUTHOR:.2f}명 / {MEDIAN_AUTHOR}명','대형팀이 평균을 상승시킴','전형적 팀은 4인 내외로 해석함'],['다기관',f'{MULTI_INST_N:,}건·{MULTI_INST_N/TOTAL*100:.1f}%','평균 기관 수 1.44개·중앙값 1개','기관 간 협업은 선택적임'],['국제협력',f'{INTL_N:,}건·{INTL_N/TOTAL*100:.1f}%','국가 확인 논문 기준 12.1%','허브국가 중심의 제한적 개방성을 보임']],[3.7,4.8,5.2,5.4],8.2)

add_heading(doc,'6.3 공저 규모: 소·중규모 팀과 대형 컨소시엄의 공존'); doc.add_picture(str(FIGS['team']),width=Cm(18)); add_caption(doc,'그림 6-2. 공저 규모 분포',True)
for txt,major in [
('2인 435건, 3인 452건, 4인 397건으로 2~4인 연구가 1,284건·53.5%를 차지하여 전형적인 협업 단위가 소·중규모 팀으로 나타남.',True),('5~9인 연구는 659건·27.5%로 확인되어 소자 제작, 실험 검증, 알고리즘·제어·데이터 분석이 결합되는 중형팀도 의미 있는 비중을 보임.',False),('10명 이상 대형 공동연구는 170건·7.1%로 나타나 대형 장비, 다중 실험사이트, 검증형 연구, 국제 컨소시엄 가능성을 보여줌.',False),(f'저자 수 90분위 {P90}명, 95분위 {P95}명, 99분위 {P99}명, 최대 {MAX_AUTHORS}명으로 확인되어 평균 저자 수는 대형 논문에 민감하므로 중앙값과 구간별 분포를 우선 해석함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-3. 공저 규모별 논문 분포')
add_table(doc,['공저 규모','논문 수','비중','구조적 해석'],[[lab,f'{n:,}',f'{s:.1f}%','단독 연구' if i==0 else '핵심 소·중규모 팀' if i in [1,2,3] else '중형 협업팀' if i in [4,5,6,7,8] else '대형 실험·컨소시엄형'] for i,(lab,n,s) in enumerate(zip(team_labels,team_counts,team_shares))],[3.4,3.4,3.3,8.8],8.2)

add_heading(doc,'6.4 월별 협업 구조 변화'); doc.add_picture(str(FIGS['monthly']),width=Cm(18.1)); add_caption(doc,'그림 6-3. 월별 공동저자·다기관·국제협력 비율',True)
for txt,major in [
('공동저자 비율은 전 기간 82.7~94.0%로 높게 유지되어 월별 논문 수 증감과 무관하게 팀 연구 중심 구조가 안정적으로 지속됨.',True),('다기관 비율은 2025년 9월 36.1%로 가장 높고 2025년 6월 22.0%로 가장 낮아 특정 월의 실험·응용 주제 구성에 따라 변동하는 것으로 보임.',False),('국제협력 비율은 2025년 10월 14.8%로 최고, 2025년 8월 5.5%로 최저이며 월별 차이가 약 9.3%p로 확인됨.',False),('2026년 3~5월 논문 수는 높은 수준을 유지하였으나 국제협력률은 9.0%→8.2%→6.4%로 하락하여 연구량 증가가 반드시 국제협력 확대로 이어지지 않음을 보임.',False),(f'논문 수와 협업률의 탐색적 상관은 공동저자 {corr_volume_coauthor:+.2f}, 다기관 {corr_volume_multi:+.2f}, 국제협력 {corr_volume_intl:+.2f}로 나타나 12개월 표본에서는 연구량과 협업 강도의 직접 연동성이 제한적으로 보임.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-4. 월별 협업 지표')
add_table(doc,['월','논문 수','평균 저자','공동저자','다기관','국제협력'],[[m,f'{int(paper_counts[i])}',f'{avg_authors[i]:.2f}',f'{coauthor_rate[i]:.1f}%',f'{multi_inst_rate[i]:.1f}%',f'{intl_rate[i]:.1f}%'] for i,m in enumerate(months)],[2.7,2.5,2.8,3,3,3],7.9)

add_heading(doc,'6.5 월별 평균 저자 수와 대형 공동연구 영향'); doc.add_picture(str(FIGS['authors']),width=Cm(18)); add_caption(doc,'그림 6-4. 월별 평균 저자 수와 대형 공동연구 영향',True)
for txt,major in [('2025년 6월 평균 9.42명과 2026년 2월 평균 8.90명은 다른 월보다 크게 높아 일부 대형 공동연구 논문의 영향이 집중된 것으로 판단함.',True),('2026년 4~5월 평균 저자 수는 각각 4.01명, 3.98명으로 낮지만 논문 수는 각 220편으로 높아 최근 연구량 증가는 대형 컨소시엄보다 소·중규모 팀의 다수 생산에 의해 형성된 것으로 보임.',False),('월별 평균 저자 수만으로 협업 성숙도를 판단하면 대형 논문의 영향이 과대 반영되므로 공동저자율, 다기관율, 국제협력률, 중앙값을 함께 적용함.',False)]: add_bullet(doc,txt,major)

add_heading(doc,'6.6 전·후반기 협업 강도 비교'); doc.add_picture(str(FIGS['half']),width=Cm(16.8)); add_caption(doc,'그림 6-5. 전·후반 6개월 협업 강도 비교',True)
for txt,major in [(f'논문 수 가중 공동저자 비율은 전반기 {weighted_coauthor_first:.1f}%에서 후반기 {weighted_coauthor_second:.1f}%로 낮아졌으나 여전히 80%대 후반의 높은 수준을 유지함.',True),(f'다기관 비율은 전반기 {weighted_multi_first:.1f}%와 후반기 {weighted_multi_second:.1f}%로 유사하여 기관 간 협업 강도는 구조적으로 안정적임.',False),(f'국제협력 비율은 전반기 {weighted_intl_first:.1f}%에서 후반기 {weighted_intl_second:.1f}%로 소폭 낮아져 최근 연구량 확대가 국내·단일국 중심 연구 증가와 병행된 가능성이 있음.',False),('반기 비교는 월별 구성효과와 반올림 자료의 영향을 포함한 탐색값이므로 분야·기관·국가별 재집계로 원인을 검증해야 함.',False)]: add_bullet(doc,txt,major)

add_heading(doc,'6.7 기관 간 협업 구조와 데이터 표준화 이슈'); doc.add_picture(str(FIGS['inst']),width=Cm(18.1)); add_caption(doc,'그림 6-6. 기관 협업쌍 TOP10: 기관명 표준화 전 탐색값',True)
for txt,major in [(f'논문당 평균 기관 수는 {AVG_INST:.2f}개, 중앙값은 {MEDIAN_INST}개로 나타나 다수 논문이 하나의 주관기관을 중심으로 수행됨.',True),('상위 기관쌍의 공동 논문 수는 3~5건으로 낮아 특정 기관쌍이 전체 협업을 지배하는 구조는 확인되지 않음.',False),('MIT↔Mit, Imperial↔Imperial College London처럼 동일기관의 표기 변형이 협업쌍으로 추출되어 기관 네트워크는 표준화 전 확정 순위로 사용할 수 없음.',False),('최종 기관 협업망은 ROR·OpenAlex·Crossref·GRID 식별자를 활용해 대학 본부·부속연구소·학과·약어를 통합한 뒤 재산출해야 함.',False),('기관쌍 분석의 전략적 가치는 현재 순위보다 장비·소재·알고리즘·실험 인프라를 서로 보완하는 기관 유형을 탐색하는 데 있음.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-5. 기관 협업쌍 추출값과 검증 필요성')
add_table(doc,['순위','기관 협업쌍','공동 논문','판정','후속 조치'],[[i+1,pair,n,'동일기관 표기 가능' if i in [0,4] else '복수기관 신호','기관명 표준화 후 재검증'] for i,(pair,n) in enumerate(institution_pairs)],[1.2,8.2,2.5,4,4.5],7.8)

add_heading(doc,'6.8 국가 간 협업쌍: 미국 중심 글로벌 축과 권역형 연결'); doc.add_picture(str(FIGS['pairs']),width=Cm(18)); add_caption(doc,'그림 6-7. 국가 간 국제협력쌍 TOP10',True)
for txt,major in [('독일–미국, 인도–미국, 영국–미국이 각각 14건으로 최상위 협업축을 형성하여 미국이 다수 권역을 연결하는 중심국으로 확인됨.',True),('중국–홍콩 12건은 권역 내 연구·인력 네트워크의 강도를 보여주며 중국–미국, 중국–싱가포르, 중국–영국은 중국 연구망의 권역 외 확장을 나타냄.',False),('독일–이탈리아는 유럽 내 협력축, 스페인–미국과 스위스–미국은 대서양 연계축으로 해석할 수 있음.',False),('상위 10개 국가쌍의 합계는 104건이나 동일 국제협력 논문이 다수 국가쌍을 생성할 수 있으므로 241건 국제협력 논문과 단순 합산 비교하지 않음.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-6. 국가 간 협업쌍 TOP10')
add_table(doc,['순위','국가 협업쌍','공동 논문','네트워크 유형','전략적 해석'],[[i+1,pair,n,'미국 중심 글로벌축' if '미국' in pair else '중국 권역·글로벌축' if '중국' in pair else '유럽 권역축','핵심 파트너 후보 및 기술군별 협력 검증 필요'] for i,(pair,n) in enumerate(country_pairs)],[1.2,5.7,2.5,4.5,6],7.8)

add_heading(doc,'6.9 국제협력 허브 국가'); doc.add_picture(str(FIGS['hubs']),width=Cm(18)); add_caption(doc,'그림 6-8. 국제협력 허브 국가: 가중 연결 수와 국제협력 논문 수',True)
for txt,major in [('미국은 가중 연결 수 133과 국제협력 논문 93건으로 가장 강한 허브성을 보이며 독일·영국·중국이 뒤를 이음.',True),('독일과 영국은 제5장에서 확인한 전체 기관 국가 참여 확대와 협업 네트워크 중심성이 동시에 나타나 국제 공동연구 플랫폼 역할이 강한 것으로 보임.',False),('이탈리아·스페인·프랑스·네덜란드 등 유럽국은 절대 연구량보다 다국가 연결에서 의미 있는 위치를 보이며 EU의 분산형 연구 인프라와 정합적임.',False),('가중 연결 수는 국가쌍 동시출현 횟수의 합으로 실제 상대국 수와 다르며, 정밀한 허브성 판단에는 degree, betweenness, eigenvector centrality를 추가해야 함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-7. 국제협력 허브 국가 TOP10')
add_table(doc,['순위','국가','가중 연결 수','국제협력 논문','구조적 해석'],[[i+1,n,w,d,'최상위 글로벌 허브' if i==0 else '주요 국제협력 허브' if i<4 else '권역·브리지 허브 후보'] for i,(n,w,d) in enumerate(hub_countries)],[1.2,3.5,3.1,3.2,8.4],8.0)

add_heading(doc,'6.10 분야별 협업 강도: 이론·실험·응용의 차이')
for txt,major in [('양자물리(quant-ph)는 압도적인 논문 규모와 높은 공동저자 비율을 동시에 유지하여 분야 전체의 기본 협업 플랫폼으로 기능함.',True),('컴퓨터 비전(cs.CV), 기계학습(cs.LG), 신호처리(eess.SP)는 표본 규모는 작지만 다기관 협업이 상대적으로 높아 데이터·알고리즘·시스템 전문성의 결합이 중요한 것으로 보임.',False),('메조스코픽/홀 효과(cond-mat.mes-hall)와 광학(physics.optics)은 장비·측정·소자 제작이 요구되어 평균 저자 수와 다기관 협업이 높은 실험·인프라형 영역으로 해석함.',False),('고에너지 물리 이론(hep-th), 일반상대론/양자중력(gr-qc)은 상대적으로 소규모 팀과 낮은 국제협력 비율이 유지되는 이론형 연구 구조가 일부 나타남.',False),('정확한 분야별 협업률은 주 카테고리 표본 수와 국가정보 커버리지를 함께 제시하고, 복수 카테고리 논문과 기술군 재분류 기준으로 재산출해야 함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-8. 분야 유형별 협업 구조 해석')
add_table(doc,['분야 유형','대표 카테고리','상대적 협업 구조','필요 자원·전문성','후속 분석'],[
['핵심 양자물리','양자물리(quant-ph)','공동저자 보편·다기관 중간','이론·실험·알고리즘 공통 기반','세부 플랫폼별 분해'],['AI·데이터·공학','기계학습(cs.LG), 컴퓨터 비전(cs.CV), 신호처리(eess.SP)','다기관 상대적 높음','데이터·SW·제어·응용기관 결합','기관유형·기업 참여 분석'],['소자·물성·광학','메조스코픽/홀 효과, 강상관 전자, 광학','저자·다기관 규모 높음','장비·공정·광원·극저온·측정','시설·공급망 협업 분석'],['이론·수리','고에너지 이론, 일반상대론/양자중력, 수리물리','소규모 팀·국제협력 상대적 낮음','이론·모델·수학적 전문성','인용·공동저자 지식망 분석'],['통신·보안','암호/보안(cs.CR), 광학, 양자물리','표준·실증 중심 선택적 협업','통신사·장비사·표준기관·보안기관','QKD·PQC·네트워크 실증 분석']],[3,5.8,4.4,5.2,4.4],7.6)

add_heading(doc,'6.11 최신 정책·표준화 환경과 협업 구조의 통합 해석')
for txt,major in [
('미국 NQI는 과학·인력·산업·인프라·안보·국제협력을 정책축으로 제시하고, 독일·한국·영국·일본 등과 양자 협력 공동성명을 추진함. 독일–미국과 영국–미국이 최상위 협업쌍으로 확인된 결과는 이러한 제도적 협력 기반과 구조적으로 부합함 [R1].',True),
('미국 NQI 보고서는 미국 저자가 포함된 광범위한 QIS 출판물의 절반 이상이 국제 공동저자를 포함한다고 제시함. 본 데이터의 국제협력률 10.0%는 제목 검색식, 12개월 범위, 기관국가 커버리지, 기술군 구성 차이로 직접 비교할 수 없으며 협업률 과소추정 가능성을 함께 고려해야 함 [R1].',False),
('EU Quantum Europe Strategy는 유럽의 전략·로드맵 분절을 문제로 지적하고, 조정된 인프라 허브·공급망·인력 이동을 핵심축으로 제시함. 독일·이탈리아 및 유럽권 허브국의 연결은 다국가 인프라형 생태계와 연결해 해석함 [R2].',False),
('영국 National Quantum Strategy는 10년 장기전략과 국가 프로그램을 통해 연구·산업·인프라·인력을 결합함. 영국이 미국·중국과 주요 협업쌍을 형성하고 허브국가 상위에 위치한 점은 국제 네트워크형 전략의 결과로 볼 수 있는 보조 신호임 [R3].',False),
('ITU-T Y.Sup98은 양자 네트워크의 계층모델, 프로토콜, 성능, 전송, 보안, IoT 적용을 여러 Study Group이 협력해 표준화해야 한다고 제시함. 양자통신·QKD 연구는 단일 연구실보다 통신·보안·광학·시스템 기관 간 협력이 필수적인 영역으로 판단함 [R4].',False)]: add_bullet(doc,txt,major)
add_callout(doc,'정량–정성 통합 판단',['공동저자 88.0%는 양자기술의 다학제성을 반영하나, 국제협력 10.0%는 전략기술·인프라·보안·데이터 결측의 영향을 함께 받음.','국제 공동연구는 단순 개방성보다 공동 인프라 접근, 상호보완 전문성, 표준화 참여, 신뢰·연구보안 체계에 의해 선택적으로 형성됨.','협업 우선순위는 전체 논문 수보다 기술군 적합성, 허브 중심성, 장비·실증 인프라, 표준화 역할, 지식재산·수출통제 위험을 함께 평가함.'],LIGHT_ORANGE,ORANGE)

add_heading(doc,'6.12 기술군별 권장 협업 모델'); add_caption(doc,'표 6-9. 양자 기술군별 협업 모델과 파트너 구성')
add_table(doc,['기술군','권장 팀 구조','핵심 파트너','협업 목적','주의요인'],[
['양자컴퓨팅 하드웨어','대형 다기관·국제 컨소시엄','소자·극저온·제어·공정·클라우드 기관','큐비트 확장·오류정정·시스템 통합','IP·장비 공급망·수출통제'],['양자SW·AI','소·중규모 다학제 팀','알고리즘·ML·컴파일러·응용기관','디코더·회로최적화·벤치마크','재현성·성능기준·플랫폼 종속'],['양자통신·보안','통신사·장비사·표준기관 포함 다기관','광학·네트워크·보안·PQC·QKD 기관','상호운용·실증망·마이그레이션','보안등급·표준경쟁·망운영'],['양자센싱·계측','장비·응용기관 중심 공동실증','원자·광학·신호처리·국방·의료기관','성능검증·환경적응·시험인증','측정표준·현장데이터·듀얼유스'],['양자소자·소재','시설공유형 다기관 협력','대학·국가연구소·파운드리·소재기업','공정·결함·패키징·수율 개선','시설 접근·공정기밀·소재 공급망'],['기초이론·시뮬레이션','소규모 국제 네트워크','수리물리·다체계·HPC 연구그룹','모델·알고리즘·검증문제 개발','저자관행·성과귀속·장기성']],[3.4,4.5,5.8,5.2,4.4],7.5)

add_heading(doc,'6.13 한국의 협업 전략 시사점')
for txt,major in [('제5장에서 한국은 제1저자 국가 44건, 전체 기관 국가 48건으로 두 기준 모두 12위이며 참여 확대 폭이 +4건에 그쳐 국제 공동연구 외연을 넓힐 필요가 있음.',True),('미국 NQI의 양자 협력 공동성명 대상에 한국이 포함되어 있으므로 제도적 협력 기반을 실제 공동과제·인력교류·시설공유·표준화 참여로 전환할 필요가 있음.',False),('한국은 전 기술군을 균등 확대하기보다 양자센싱·광학, 암호/보안, 소자·소재, 양자AI 등 국내 강점과 연결되는 분야에서 미국·독일·영국·일본 허브기관과 선택적 협력을 설계하는 것이 적합함.',False),('국제협력 평가지표는 공동논문 건수 외에 공동 제1저자·교신저자, 브리지 연구자, 기관쌍 지속성, 공동 특허·표준, 인력 이동, 공동 실증을 포함해야 함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-10. 한국의 협업 네트워크 강화 실행과제')
add_table(doc,['과제','현재 진단','실행 내용','성과지표'],[
['허브국 연계','미국·독일·영국 등이 글로벌 허브','기술군별 상위 기관·브리지 연구자를 선정해 공동과제를 설계함','신규 국가쌍·기관쌍, 공동 제1저자'],['공동 인프라','다기관 연구 29.6%로 선택적','극저온·광원·파운드리·통신 테스트베드의 공동 접근체계를 구축함','시설 공동활용, 공동 실증'],['표준화 협력','양자 네트워크·PQC 표준화 가속','ITU·ISO/IEC·ETSI·GSMA 활동과 논문·실증 컨소시엄을 연계함','국제표준 기고, 시험·인증'],['인력 이동','국가 간 협력은 인력·신뢰 기반','공동박사·포닥·산업 파견·방문연구를 확대함','인력 교류, 공동 지도, 정착률'],['연구보안·IP','전략기술 특성상 개방성 제약','공동연구 계약·데이터·IP·수출통제 기준을 사전 설계함','협약기간, 분쟁·지연 감소']],[3.2,5.2,7.1,5],7.8)

add_heading(doc,'6.14 분석 한계 및 데이터 보완 과제')
for txt,major in [('기관 국가 커버리지는 83.25%이므로 국제협력률은 실제보다 과소 추정될 수 있으며 국가 미확인 논문의 기술군·월별 편향을 점검해야 함.',True),('기관명 표준화 오류가 기관쌍을 왜곡하므로 현재 기관 네트워크는 탐색 신호로만 활용하고 ROR 등 외부 식별자 기반 재정비가 필수임.',False),('공동논문은 협력의 존재를 보여주지만 역할·기여도·인프라 제공·데이터 제공·교신 책임을 구분하지 못하므로 CRediT 기여문과 교신저자 정보를 연계해야 함.',False),('국가쌍 공동논문 수는 다국가 논문에서 복수쌍이 생성되므로 논문 수, 가중 degree, 상대국 수, betweenness를 구분해야 함.',False),('12개월 데이터는 장기 협력 지속성을 판단하기 어려우므로 최소 3년 누적과 최근 12개월 가속도를 병렬 적용해야 함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-11. 협업 데이터 품질과 보완 방향')
add_table(doc,['한계요인','영향','현재 통제','보완 방향'],[
['기관명 변형','허위 기관쌍·허브 왜곡','동일기관 의심쌍을 주석 처리함','ROR·OpenAlex·GRID 기반 통합'],['기관 국가 결측','국제협력 과소추정','전체·유효분모를 함께 제시함','소속문자열·도메인·외부DB 보완'],['저자 기여도 미확인','실제 리더십·역할 구분 불가','저자수·제1저자와 분리 해석함','교신저자·CRediT·마지막 저자 연계'],['다국가 논문의 복수쌍','쌍 합계가 논문 수 초과','가중 연결 수로 명시함','단순·가중 네트워크 지표 병렬 산출'],['짧은 분석기간','지속 협력과 일회 협력 구분 한계','월별·반기별 변화를 제시함','3~5년 코호트·생존 분석'],['기술군 혼합','협업률의 분야 구성효과','카테고리별 상대적 해석을 수행함','실무 기술분류 재매핑 후 재산출']],[3.4,5.2,5.7,6.2],7.8)

add_heading(doc,'6.15 제7장 연결용 종합 시사점')
for txt,major in [('양자 분야는 공동저자 연구가 보편적이지만 다기관·국제협력은 기술적 필요성과 제도적 조건에 따라 선택적으로 발생하는 다층 협업 생태계로 정의함.',True),('연구량 증가가 협업 강도 증가와 직접 연결되지 않으므로 유망기술 분석에서는 논문 성장률과 다기관·국제협력 성장률을 별도 지표로 적용함.',False),('미국은 국가쌍·허브 지표 모두에서 중심성이 높고 독일·영국·중국은 권역과 글로벌 네트워크를 연결하는 주요 파트너로 확인됨.',False),('기관 협업은 현재 표준화 오류가 크므로 제7장의 전략 파트너 선정 전에 기관명 정비와 기술군별 네트워크 재산출을 선행함.',False),('제7장에서는 연구 규모·분야·주체·협업 결과를 결합하여 고성장–고협업, 고성장–저협업, 저성장–고허브 기술군을 구분하고 R&D·인력·국제협력 전략을 제시함.',False)]: add_bullet(doc,txt,major)
add_caption(doc,'표 6-12. 협업 구조 유형별 전략 대응')
add_table(doc,['협업 유형','정량 특징','전략적 해석','권장 대응'],[
['고공저·저다기관','팀 연구는 활발하나 동일기관 중심','내부 전문성 결합형 연구','기관 간 장비·데이터·실증 연계를 확대함'],['고다기관·저국제','국내 기관 협력은 활발하나 국가 경계 유지','국가 프로젝트·인프라 중심','국제 허브기관과 단계적 공동과제를 설계함'],['고국제·허브집중','소수 허브국가와 반복 연결','전략적 글로벌 네트워크','브리지 연구자·기관과 장기 파트너십을 구축함'],['대형팀·저지속성','특정 월·프로젝트에 대형팀 집중','이벤트·컨소시엄형 연구','3년 지속성·후속 성과·공동특허를 확인함'],['소규모·고성장','소수팀이 빠르게 신규 주제를 생산','신흥 이론·알고리즘 가능성','신규 저자·키워드·토픽의 조기 신호를 추적함']],[3.5,5,6,6],8.0)

add_heading(doc,'6.16 본 장 소결')
for txt,major in [('공동저자 연구는 2,113건·88.0%로 양자 분야의 기본 수행방식이며 전형적인 팀 규모는 중앙값 4명과 2~4인 연구 53.5%로 요약됨.',True),('다기관 연구는 710건·29.6%, 국제협력은 241건·10.0%로 연구자 협업이 기관·국가 경계를 넘을수록 선택성이 커짐.',False),('월별 협업률은 연구량과 직접 동행하지 않으며 최근 고활동 구간에서도 국제협력률이 낮아져 연구 생산 확대와 네트워크 개방성을 별도로 관리해야 함.',False),('미국은 독일·영국·인도·중국 등과 연결되는 최상위 허브이며 유럽과 중국은 권역 내·외 연결축을 동시에 형성함.',False),('기관쌍 결과는 표준화 오류를 포함하므로 최종 전략보고서에서는 기관 식별자 정비, 기술군 재분류, 네트워크 중심성, 지속성 지표를 추가함.',False),('다음 제7장에서는 본 장의 협업 구조를 연구 규모·카테고리·주체 분석과 결합하여 기술군별 국제협력 우선순위와 실행전략으로 전환함.',False)]: add_bullet(doc,txt,major)

add_heading(doc,'참고자료')
refs=[('[R1]','National Quantum Initiative, Annual Report FY2025','https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf'),('[R2]','European Commission, Quantum Europe Strategy, 2025','https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy'),('[R3]','UK Government, National Quantum Strategy','https://www.gov.uk/government/publications/national-quantum-strategy'),('[R4]','ITU-T Y.Sup98, Technical Considerations Towards Quantum Networks, 2025','https://www.itu.int/rec/dologin_pub.asp?id=T-REC-Y.Sup98-202511-I!!PDF-E&lang=s&type=items')]
for rid,title,url in refs:
    p=doc.add_paragraph(); pformat(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(.42); p.paragraph_format.first_line_indent=Cm(-.42); r=p.add_run(rid+' '); set_run(r,8.8,True,NAVY); add_hyperlink(p,title,url)

cp=doc.core_properties; cp.title='arXiv 양자 분야 제6장 협업 구조 분석'; cp.subject='2025년 6월~2026년 5월, 12개월·2,400건 기준 정량·정성 통합 분석'; cp.author='OpenAI'; cp.keywords='arXiv, 양자, 협업, 공동저자, 다기관, 국제협력, 네트워크, 표준화'
settings=doc.settings._element; update=OxmlElement('w:updateFields'); update.set(qn('w:val'),'true'); settings.append(update)
doc.save(OUTPUT); print(f'Created {OUTPUT}')

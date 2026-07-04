from __future__ import annotations

from pathlib import Path
import math

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Polygon
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path('artifacts')
ASSET_DIR = OUT_DIR / 'chapter7_assets'
OUT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / 'arXiv_양자분야_제7장_종합인사이트_및_전략적시사점_고품질_재작성_20260704.docx'

# -----------------------------------------------------------------------------
# Integrated evidence from Chapters 2-6
# -----------------------------------------------------------------------------
TOTAL = 2400
MONTHLY_AVG = 200.0
RECENT3_AVG = 217.0
HALF_GROWTH = 5.0
CV = 11.1
QUANT_PH = 1680
QUANT_PH_SHARE = 70.00
ALL_CATEGORY_N = 117
MULTI_CATEGORY_N = 1138
MULTI_CATEGORY_SHARE = 47.42
CATEGORY_HHI = 0.493
UNIQUE_AUTHORS = 10514
UNIQUE_FIRST_AUTHORS = 2248
VALID_INSTITUTIONS = 1995
COUNTRY_N = 35
COAUTHOR_N = 2113
COAUTHOR_SHARE = 88.0
MULTI_INST_N = 710
MULTI_INST_SHARE = 29.6
INTL_N = 241
INTL_SHARE = 10.0
AVG_AUTHORS = 5.41
MEDIAN_AUTHORS = 4

COUNTRY_DATA = [
    ('미국',453,524),('중국',277,296),('독일',161,199),('영국',113,142),
    ('일본',112,124),('인도',95,111),('이탈리아',91,106),
    ('프랑스',56,77),('캐나다',71,73),('호주',59,70),('한국',44,48)
]
COUNTRY_HUB = {'미국':133,'독일':79,'영국':73,'중국':66,'이탈리아':54,'인도':44,'프랑스':30,'한국':6}

# Expert synthesis scores: 1-5 scale. These are not raw measurements.
TECH_PORTFOLIO = [
    ('양자컴퓨팅·SW/AI',4.8,4.6,5.0,4.2,4.4,'최우선'),
    ('양자통신·보안',3.8,4.2,4.8,4.5,4.0,'최우선'),
    ('양자센싱·계측',4.1,4.0,4.4,4.2,3.5,'최우선'),
    ('양자소자·소재',4.3,3.8,4.1,5.0,4.7,'핵심기반'),
    ('양자이론·시뮬레이션',4.4,3.2,2.8,3.2,2.5,'장기기반'),
    ('양자화학·산업응용',3.3,3.5,4.0,3.8,3.6,'응용발굴'),
]

NAVY='17365D'; BLUE='2F75B5'; LIGHT_BLUE='D9EAF7'; GREEN='548235'; LIGHT_GREEN='E2F0D9'; ORANGE='ED7D31'; LIGHT_ORANGE='FCE4D6'; PURPLE='7030A0'; LIGHT_PURPLE='E4DFEC'; GRAY='6B7280'; LIGHT_GRAY='F2F4F7'; DARK='243447'; WHITE='FFFFFF'; BORDER='CAD4DF'; YELLOW='FFD966'; RED='C00000'

def rgb(h): return RGBColor.from_string(h)

font_candidates = [
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    '/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc',
    '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
]
font_path = next((p for p in font_candidates if Path(p).exists()), None)
if font_path:
    fp = font_manager.FontProperties(fname=font_path)
    plt.rcParams['font.family'] = fp.get_name()
else:
    plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['axes.unicode_minus'] = False
plt.rcParams['font.weight'] = 'bold'

# -----------------------------------------------------------------------------
# Figures - matplotlib default color cycle for consistency and legibility
# -----------------------------------------------------------------------------
def finish(fig, path):
    fig.tight_layout(pad=1.0)
    fig.savefig(path, dpi=220, bbox_inches='tight')
    plt.close(fig)


def fig_integrated_dashboard(path):
    fig, ax = plt.subplots(figsize=(12, 6.3))
    ax.axis('off'); ax.set_xlim(0, 1); ax.set_ylim(0, 1)
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    cards = [
        (0.03,0.55,0.215,0.34,colors[0],'연구 규모·성장',[
            f'{TOTAL:,}건 · 월평균 {MONTHLY_AVG:.1f}편',
            f'최근 3개월 {RECENT3_AVG:.1f}편',
            f'후반기 +{HALF_GROWTH:.1f}% · CV {CV:.1f}%']),
        (0.265,0.55,0.215,0.34,colors[1],'분야·융합',[
            f'양자물리(quant-ph) {QUANT_PH_SHARE:.2f}%',
            f'전체 카테고리 {ALL_CATEGORY_N}개',
            f'복수 카테고리 {MULTI_CATEGORY_SHARE:.2f}%']),
        (0.50,0.55,0.215,0.34,colors[2],'연구 주체',[
            f'저자 {UNIQUE_AUTHORS:,}명',
            f'기관 {VALID_INSTITUTIONS:,}개',
            f'{COUNTRY_N}개국 · 한국 12위']),
        (0.735,0.55,0.235,0.34,colors[3],'협업 구조',[
            f'공동저자 {COAUTHOR_SHARE:.1f}%',
            f'다기관 {MULTI_INST_SHARE:.1f}%',
            f'국제협력 {INTL_SHARE:.1f}%']),
    ]
    for x,y,w,h,color,title,lines in cards:
        ax.add_patch(FancyBboxPatch((x,y),w,h,boxstyle='round,pad=.015,rounding_size=.025',facecolor='white',edgecolor=color,linewidth=2.3))
        ax.text(x+w/2,y+h-.065,title,ha='center',fontsize=12.5,fontweight='bold',color=color)
        ax.text(x+w/2,y+.14,'\n'.join(lines),ha='center',va='center',fontsize=10.5,fontweight='bold',color='black',linespacing=1.45)
    ax.add_patch(FancyBboxPatch((.10,.11),.80,.27,boxstyle='round,pad=.02,rounding_size=.03',facecolor='white',edgecolor='black',linewidth=2.2))
    ax.text(.50,.30,'통합 진단',ha='center',fontsize=11.5,fontweight='bold')
    ax.text(.50,.225,'높은 연구활동과 강한 과학 코어 위에서 다학제 융합과 상용화 전환이 동시 진행됨',ha='center',fontsize=14,fontweight='bold')
    ax.text(.50,.155,'다만 국제협력·기관표준화·시장 검증은 연구량 대비 보완이 필요한 전략적 병목으로 판단함',ha='center',fontsize=10.5,fontweight='bold')
    for x in [.14,.37,.61,.84]:
        ax.add_patch(FancyArrowPatch((x,.54),(.50,.39),arrowstyle='-|>',mutation_scale=14,linewidth=1.2,alpha=.7))
    ax.text(.50,.96,'제3~6장 핵심 결과의 통합 증거 맵',ha='center',fontsize=17,fontweight='bold')
    finish(fig,path)


def fig_technology_portfolio(path):
    fig, ax = plt.subplots(figsize=(11.5, 7.0))
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    for i,(name,research,growth,market,collab,risk,priority) in enumerate(TECH_PORTFOLIO):
        size = 260 + collab * 110
        ax.scatter(research, market, s=size, alpha=.72, edgecolors='black', linewidths=1.0, color=colors[i % len(colors)])
        dx = 0.05 if name != '양자컴퓨팅·SW/AI' else -0.78
        dy = 0.08 if name not in ['양자통신·보안','양자소자·소재'] else -0.22
        ax.text(research+dx, market+dy, f'{name}\n성장 {growth:.1f}·리스크 {risk:.1f}', fontsize=9.4, fontweight='bold')
    ax.axvline(4.0, linestyle='--', linewidth=1.1)
    ax.axhline(4.0, linestyle='--', linewidth=1.1)
    ax.set_xlim(2.6,5.25); ax.set_ylim(2.3,5.25)
    ax.set_xlabel('연구기반·학술역량 종합점수(5점)',fontweight='bold')
    ax.set_ylabel('시장·정책·표준화 준비도(5점)',fontweight='bold')
    ax.set_title('양자 기술축 전략 포트폴리오',fontsize=16,fontweight='bold',pad=18)
    ax.grid(alpha=.25)
    ax.text(4.63,5.08,'선도·상용화 가속',ha='center',fontsize=11,fontweight='bold')
    ax.text(3.25,5.08,'표준·응용 선점',ha='center',fontsize=11,fontweight='bold')
    ax.text(4.63,2.45,'기반기술 장기투자',ha='center',fontsize=11,fontweight='bold')
    ax.text(3.25,2.45,'선택적 탐색·검증',ha='center',fontsize=11,fontweight='bold')
    ax.text(2.68,2.58,'※ 원의 크기: 협업·인프라 의존도\n※ 점수: 정량결과와 외부자료를 결합한 5점 전문가 통합평가',fontsize=8.9,fontweight='bold')
    finish(fig,path)


def fig_country_position(path):
    fig, ax = plt.subplots(figsize=(9.5, 7.0))
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    selected = ['미국','중국','독일','영국','일본','한국']
    data = [x for x in COUNTRY_DATA if x[0] in selected]
    maxv = 560
    ax.plot([0,maxv],[0,maxv],linestyle='--',linewidth=1.2,label='주도=참여 기준선')
    for i,(name,lead,participation) in enumerate(data):
        hub = COUNTRY_HUB.get(name,5)
        ax.scatter(lead,participation,s=180+hub*5,alpha=.72,edgecolors='black',linewidths=1.0,color=colors[i])
        offsets={'미국':(-45,10),'중국':(8,5),'독일':(8,5),'영국':(8,5),'일본':(8,-18),'한국':(8,-18)}
        dx,dy=offsets[name]
        ax.annotate(f'{name}\n{lead}/{participation}',(lead,participation),xytext=(dx,dy),textcoords='offset points',fontsize=9.6,fontweight='bold')
    ax.set_xlim(0,570); ax.set_ylim(0,570)
    ax.set_xlabel('제1저자 국가 논문 수: 연구 주도성',fontweight='bold')
    ax.set_ylabel('전체 기관 국가 연결 수: 참여 외연',fontweight='bold')
    ax.set_title('주요국 연구 주도–참여–허브 포지션',fontsize=16,fontweight='bold',pad=18)
    ax.grid(alpha=.25)
    ax.text(360,500,'대각선 상단: 공동연구·비제1저자 참여 외연이 큼',fontsize=9.5,fontweight='bold',ha='center')
    ax.text(175,70,'한국: 자체 주도 기반은 존재하나\n국제 네트워크 외연 확대 필요',fontsize=9.5,fontweight='bold',ha='center')
    ax.legend(frameon=False,loc='lower right')
    finish(fig,path)


def fig_transition_timeline(path):
    fig, ax = plt.subplots(figsize=(12,5.6))
    ax.axis('off'); ax.set_xlim(2023.5,2035.8); ax.set_ylim(0,1)
    ax.plot([2024,2035],[.48,.48],linewidth=2)
    milestones = [
        (2024,.72,'NIST FIPS\n203·204·205','표준 전환'),
        (2025,.26,'EU Quantum Europe\nStrategy·ITU Y.Sup98','인프라·표준'),
        (2026,.75,'300+ 기업 채택\n미국 EO 14413','상용화·배치'),
        (2028,.25,'시장 매출\n최대 44억 달러 전망','확대 시나리오'),
        (2029,.74,'IBM 대규모\n내결함성 목표','기술 로드맵'),
        (2033,.25,'DARPA QBI\n유틸리티 규모 검증','독립 검증'),
        (2035,.74,'한국 1만 인력·2천 기업\n시장가치 장기 전망','생태계 목표'),
    ]
    for i,(year,y,text,sub) in enumerate(milestones):
        ax.scatter([year],[.48],s=90,zorder=3)
        ax.plot([year,year],[.48,y],linewidth=1.3)
        ax.add_patch(FancyBboxPatch((year-.65,y-.11),1.3,.20,boxstyle='round,pad=.015,rounding_size=.03',facecolor='white',edgecolor=plt.rcParams['axes.prop_cycle'].by_key()['color'][i%10],linewidth=1.8))
        ax.text(year,y+.015,text,ha='center',va='center',fontsize=9.0,fontweight='bold',linespacing=1.2)
        ax.text(year,y-.075,sub,ha='center',fontsize=7.8,fontweight='bold')
        ax.text(year,.43,str(year),ha='center',fontsize=9,fontweight='bold')
    ax.text(2029.7,.96,'연구 중심 → 검증·표준화 → 상용화·배치로 이동하는 외부환경',ha='center',fontsize=16,fontweight='bold')
    ax.text(2029.7,.05,'기업·정부·표준화 로드맵은 목표와 전망이며 실제 실현 시점은 기술검증·비용·공급망에 따라 변동 가능함',ha='center',fontsize=9.2,fontweight='bold')
    finish(fig,path)


def fig_priority_pyramid(path):
    fig, ax = plt.subplots(figsize=(10.5,6.2))
    ax.axis('off'); ax.set_xlim(0,1); ax.set_ylim(0,1)
    colors = plt.rcParams['axes.prop_cycle'].by_key()['color']
    layers = [
        ((.20,.10),(.80,.10),(.70,.34),(.30,.34),colors[2],'3단계 장기기반·응용발굴','양자이론·시뮬레이션 / 양자화학·산업응용'),
        ((.30,.36),(.70,.36),(.62,.59),(.38,.59),colors[1],'2단계 핵심기반·공급망','양자소자·소재 / 공정·극저온·광자·제어'),
        ((.40,.61),(.60,.61),(.53,.88),(.47,.88),colors[0],'1단계 선도·상용화 가속','컴퓨팅·SW/AI / 통신·보안 / 센싱·계측'),
    ]
    for pts,color,title,sub in [(x[:4],x[4],x[5],x[6]) for x in layers]:
        poly=Polygon(pts,closed=True,facecolor=color,alpha=.55,edgecolor='black',linewidth=1.4); ax.add_patch(poly)
        cx=np.mean([p[0] for p in pts]); cy=np.mean([p[1] for p in pts])
        ax.text(cx,cy+.035,title,ha='center',fontsize=11.5,fontweight='bold')
        ax.text(cx,cy-.035,sub,ha='center',fontsize=9.2,fontweight='bold')
    ax.text(.50,.96,'한국형 양자기술 투자 우선순위',ha='center',fontsize=16,fontweight='bold')
    ax.text(.50,.025,'상위층은 시장·정책·표준화 창구가 열려 있는 영역, 하위층은 장기 경쟁력과 산업 파급을 위한 기반영역으로 구성함',ha='center',fontsize=9.2,fontweight='bold')
    finish(fig,path)


def fig_roadmap(path):
    fig, ax = plt.subplots(figsize=(12,6.2))
    ax.axis('off'); ax.set_xlim(0,36); ax.set_ylim(0,6.3)
    phases=[(0,6,'0~6개월\n정비·선별'),(6,18,'6~18개월\n검증·연계'),(18,36,'18~36개월\n확대·상용화')]
    colors=plt.rcParams['axes.prop_cycle'].by_key()['color']
    for i,(s,e,title) in enumerate(phases):
        ax.add_patch(FancyBboxPatch((s+.2,5.35),e-s-.4,.65,boxstyle='round,pad=.02,rounding_size=.08',facecolor=colors[i],alpha=.65,edgecolor='black'))
        ax.text((s+e)/2,5.67,title,ha='center',va='center',fontsize=11,fontweight='bold')
    rows=[
        ('데이터·분류체계',[(0,6,'DB 표준화·기술분류'),(6,18,'논문–특허–표준 연계'),(18,36,'상시 조기경보 플랫폼')]),
        ('핵심기술',[(0,6,'6대 기술축 후보 검증'),(6,18,'실증·벤치마크 과제'),(18,36,'산업별 확산·사업화')]),
        ('국제협력',[(0,6,'허브기관·브리지 연구자'),(6,18,'공동과제·시설공유'),(18,36,'공동표준·공급망 연대')]),
        ('인력·조직',[(0,6,'핵심인력·역량맵'),(6,18,'공동박사·포닥·산업파견'),(18,36,'클러스터·기업 생태계')]),
        ('성과관리',[(0,6,'기준선·KPI 설정'),(6,18,'분기 모니터링·중간평가'),(18,36,'포트폴리오 재배분')]),
    ]
    for r,(label,items) in enumerate(rows):
        y=4.65-r*.87
        ax.text(-.2,y+.22,label,ha='right',va='center',fontsize=9.8,fontweight='bold')
        for j,(s,e,txt) in enumerate(items):
            ax.add_patch(FancyBboxPatch((s+.4,y),e-s-.8,.48,boxstyle='round,pad=.015,rounding_size=.05',facecolor='white',edgecolor=colors[j],linewidth=1.5))
            ax.text((s+e)/2,y+.24,txt,ha='center',va='center',fontsize=8.9,fontweight='bold')
    ax.text(18,6.2,'36개월 전략 실행 로드맵',ha='center',fontsize=16,fontweight='bold')
    ax.text(18,.15,'단계 종료 시 기술성·시장성·협업성·데이터 신뢰도를 재평가하여 투자 우선순위를 조정함',ha='center',fontsize=9.4,fontweight='bold')
    finish(fig,path)

FIGS = {
    'dashboard':ASSET_DIR/'fig_7_1_dashboard.png',
    'portfolio':ASSET_DIR/'fig_7_2_portfolio.png',
    'country':ASSET_DIR/'fig_7_3_country.png',
    'timeline':ASSET_DIR/'fig_7_4_transition_timeline.png',
    'pyramid':ASSET_DIR/'fig_7_5_priority_pyramid.png',
    'roadmap':ASSET_DIR/'fig_7_6_roadmap.png',
}
fig_integrated_dashboard(FIGS['dashboard'])
fig_technology_portfolio(FIGS['portfolio'])
fig_country_position(FIGS['country'])
fig_transition_timeline(FIGS['timeline'])
fig_priority_pyramid(FIGS['pyramid'])
fig_roadmap(FIGS['roadmap'])

# -----------------------------------------------------------------------------
# Word helpers
# -----------------------------------------------------------------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_border(cell, color=BORDER, size='5'):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ['top','left','bottom','right']:
        el=borders.find(qn(f'w:{edge}'))
        if el is None: el=OxmlElement(f'w:{edge}'); borders.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),size); el.set(qn('w:color'),color)

def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=tcMar.find(qn(f'w:{m}'))
        if node is None: node=OxmlElement(f'w:{m}'); tcMar.append(node)
        node.set(qn('w:w'),str(v)); node.set(qn('w:type'),'dxa')

def set_run(run,size=10,bold=False,color=DARK):
    run.font.name='맑은 고딕'; run._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); run.font.size=Pt(size); run.font.bold=bold; run.font.color.rgb=rgb(color)

def pformat(p,before=0,after=4,line=1.25,keep=False):
    pf=p.paragraph_format; pf.space_before=Pt(before); pf.space_after=Pt(after); pf.line_spacing=line; pf.keep_with_next=keep

def add_page_number(p):
    run=p.add_run(); b=OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'),'begin'); i=OxmlElement('w:instrText'); i.set(qn('xml:space'),'preserve'); i.text='PAGE'; e=OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'),'end'); run._r.extend([b,i,e])

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(); pformat(p,before=12 if level==1 else 8,after=6,line=1.0,keep=True); r=p.add_run(text); set_run(r,15 if level==1 else 12,True,NAVY)
    if level==1:
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'3'); bottom.set(qn('w:color'),BLUE); pBdr.append(bottom); pPr.append(pBdr)
    return p

def add_bullet(doc,text,major=False):
    p=doc.add_paragraph(); pformat(p,after=3,line=1.28); p.paragraph_format.left_indent=Cm(.48); p.paragraph_format.first_line_indent=Cm(-.34); r=p.add_run('□ ' if major else '· '); set_run(r,10.2,True,BLUE if major else GRAY); r=p.add_run(text); set_run(r,10.2,major,DARK); return p

def add_caption(doc,text,figure=False):
    p=doc.add_paragraph(); pformat(p,before=3,after=6,line=1.0,keep=not figure); p.alignment=WD_ALIGN_PARAGRAPH.CENTER if figure else WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run(text); set_run(r,9.2,True,GRAY if figure else NAVY)

def add_table(doc,headers,rows,widths,font=8.0):
    table=doc.add_table(rows=1,cols=len(headers)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,NAVY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,after=0,line=1.05); r=p.add_run(str(h)); set_run(r,font,True,WHITE)
    for ri,row in enumerate(rows):
        cells=table.add_row().cells
        for i,v in enumerate(row):
            cell=cells[i]; cell.width=Cm(widths[i]); set_cell_shading(cell,WHITE if ri%2==0 else LIGHT_GRAY); set_cell_border(cell); set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER if i<2 else WD_ALIGN_PARAGRAPH.LEFT; pformat(p,after=0,line=1.08); r=p.add_run(str(v)); set_run(r,font,i==0,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0); return table

def add_callout(doc,title,lines,fill=LIGHT_BLUE,accent=BLUE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False; cell=t.cell(0,0); set_cell_shading(cell,fill); set_cell_border(cell,accent,'8'); set_cell_margins(cell,130,150,130,150); p=cell.paragraphs[0]; pformat(p,after=4,line=1.0); r=p.add_run(title); set_run(r,11,True,NAVY)
    for line in lines:
        p=cell.add_paragraph(); pformat(p,after=2,line=1.2); p.paragraph_format.left_indent=Cm(.35); p.paragraph_format.first_line_indent=Cm(-.25); r=p.add_run('• '); set_run(r,9.6,True,accent); r=p.add_run(line); set_run(r,9.6,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_hyperlink(p,text,url):
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr'); col=OxmlElement('w:color'); col.set(qn('w:val'),'0563C1'); rPr.append(col); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u); r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); p._p.append(h)

# -----------------------------------------------------------------------------
# Document
# -----------------------------------------------------------------------------
doc=Document(); sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(1.35); sec.bottom_margin=Cm(1.35); sec.left_margin=Cm(1.35); sec.right_margin=Cm(1.35); sec.header_distance=Cm(.55); sec.footer_distance=Cm(.55)
style=doc.styles['Normal']; style.font.name='맑은 고딕'; style._element.rPr.rFonts.set(qn('w:eastAsia'),'맑은 고딕'); style.font.size=Pt(10)
header=sec.header.paragraphs[0]; r=header.add_run('arXiv 양자 분야 논문 분석 | 제7장 종합 인사이트 및 전략적 시사점'); set_run(r,8.4,True,NAVY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run('분석 기준: 2025.06~2026.05, 12개월·2,400건  |  '); set_run(r,8.2,False,GRAY); add_page_number(footer)

# Cover
cover=doc.add_table(rows=1,cols=1); cover.alignment=WD_TABLE_ALIGNMENT.CENTER; cell=cover.cell(0,0); set_cell_shading(cell,NAVY); set_cell_margins(cell,720,300,720,300); set_cell_border(cell,NAVY,'0'); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CHAPTER 7'); set_run(r,14,True,'9DC3E6'); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('종합 인사이트 및 전략적 시사점'); set_run(r,24,True,WHITE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('연구 규모·분야·주체·협업 구조의 통합 분석'); set_run(r,12,True,LIGHT_BLUE); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; pformat(p,before=16); r=p.add_run('2025년 6월~2026년 5월 | 12개월·2,400건'); set_run(r,10.8,True,YELLOW)

doc.add_paragraph(); kpis=[('2,400건','분석 논문'),('217.0편','최근 3개월 평균'),('47.42%','복수 카테고리'),('10,514명','고유 저자'),('10.0%','국제협력')]; t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(val,label) in enumerate(kpis):
    cell=t.cell(0,i); fill=[LIGHT_BLUE,LIGHT_GREEN,LIGHT_ORANGE,LIGHT_PURPLE,LIGHT_GRAY][i]; accent=[BLUE,GREEN,ORANGE,PURPLE,GRAY][i]; set_cell_shading(cell,fill); set_cell_border(cell,accent,'7'); set_cell_margins(cell,100,70,100,70); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(val); set_run(r,14.5,True,NAVY); p=cell.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(label); set_run(r,8.5,True,GRAY)

add_callout(doc,'본 장의 종합 판단',[
    '양자 연구는 높은 활동성과 강한 양자물리 코어를 유지하면서 AI·광학·보안·소재·응용으로 확장되는 안정 성장형 융합 생태계로 판단함.',
    '연구자·기관 저변은 넓지만 국가·인프라·국제협력은 주요국 중심으로 집중되어 기술경쟁력은 규모보다 연결망·시설·표준·산업전환 역량에서 차별화될 것으로 보임.',
    '전략 우선순위는 컴퓨팅·SW/AI, 통신·보안, 센싱·계측을 선도·상용화 영역으로, 소자·소재를 핵심 공급망으로, 이론·응용을 장기기반으로 구성함.',
],LIGHT_BLUE,BLUE)
doc.add_page_break()

# 7.1
add_heading(doc,'7.1 통합 분석 프레임과 활용 원칙')
for txt,major in [
('제2장의 분석 기준, 제3장의 규모·성장, 제4장의 분야·융합, 제5장의 연구 주체, 제6장의 협업 구조를 하나의 전략 프레임으로 통합함.',True),
('정량 결과는 논문 수 자체보다 성장 지속성, 분야 집중·융합성, 연구 주체 분산성, 국가·기관 네트워크, 외부 시장·정책·표준화 신호를 함께 해석함.',False),
('본 장의 기술축 우선순위는 원자료에서 직접 산출된 단일 지표가 아니라 정량 지표와 최신 외부자료를 결합한 전문가 통합 판단으로 제시함.',False),
('2026년 6월 이후 발표된 정책은 분석기간 내 논문 변화의 직접 원인이 아니라 향후 연구·산업 방향을 검증하는 후속 환경신호로 구분함.',False),
('본 장은 향후 Executive Summary의 핵심 메시지와 실행과제의 근거를 제공하도록 장별 인사이트를 중복 없이 연결함.',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-1. 제2~6장 결과의 전략적 통합 구조')
add_table(doc,['근거 장','핵심 정량 결과','종합 해석','제7장 활용'],[
['제2장 분석 대상·방법론','12개월·2,400건, 2026년 6월 제외','비교 가능한 완전월 기준 확보','전략지표의 공통 분모로 적용함'],
['제3장 연구 규모·성장',f'월평균 {MONTHLY_AVG:.1f}편, 최근 3개월 {RECENT3_AVG:.1f}편, 후반기 +{HALF_GROWTH:.1f}%','급팽창보다 조정 후 재상승하는 안정 성장','투자 지속성과 모니터링 필요성을 판단함'],
['제4장 분야·카테고리',f'quant-ph {QUANT_PH_SHARE:.2f}%, 복수 카테고리 {MULTI_CATEGORY_SHARE:.2f}%','강한 과학 코어와 다학제 확장이 공존','6대 실무 기술축으로 재매핑함'],
['제5장 연구 주체',f'저자 {UNIQUE_AUTHORS:,}명, 기관 {VALID_INSTITUTIONS:,}개, {COUNTRY_N}개국','개인·기관은 분산, 국가는 주요국 중심','인력·기관·국가 전략을 차별화함'],
['제6장 협업 구조',f'공동저자 {COAUTHOR_SHARE:.1f}%, 다기관 {MULTI_INST_SHARE:.1f}%, 국제협력 {INTL_SHARE:.1f}%','팀 연구는 보편적이나 국제개방성은 선택적','허브기관·표준·실증 중심 협력전략을 제시함'],
],[3.1,5.8,5.8,5.4],7.8)

# 7.2
add_heading(doc,'7.2 종합 증거 맵: 성장·융합·주체·협업의 결합')
doc.add_picture(str(FIGS['dashboard']),width=Cm(18.1)); add_caption(doc,'그림 7-1. 제3~6장 핵심 결과의 통합 증거 맵',True)
for txt,major in [
('연구활동은 월평균 200편, 최근 3개월 217편으로 분석기간 말에 재상승하여 양자 분야가 일시적 이슈가 아닌 지속적 고활동 영역임을 보임.',True),
('양자물리(quant-ph) 70.00%의 코어 집중과 복수 카테고리 47.42%의 융합성이 동시에 나타나 원천과학과 응용확장이 병행됨.',False),
('고유 저자 10,514명과 기관 1,995개는 넓은 생태계를 의미하지만 국가 상위권과 협업 허브는 미국·중국·독일·영국 등 주요국에 집중됨.',False),
('공동저자 88.0%에 비해 국제협력 10.0%는 낮아 다학제성은 높지만 국경을 넘는 인프라·데이터·보안·표준 협력에는 추가 장벽이 존재함.',False),
('종합적으로 연구 규모 확대보다 분야 간 연결의 질, 연구 인프라 접근성, 국제 네트워크 중심성, 상용화 검증역량이 다음 경쟁단계를 결정할 것으로 판단함.',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-2. 종합 핵심 인사이트')
add_table(doc,['번호','종합 인사이트','정량 근거','전략적 의미'],[
['I-1','안정적 고활동과 최근 모멘텀 강화',f'최근 3개월 {RECENT3_AVG:.1f}편, 전체 평균 대비 +8.5%','중단 없는 기초·응용 R&D 파이프라인이 필요함'],
['I-2','과학 코어와 융합 응용의 동시 성장',f'quant-ph {QUANT_PH_SHARE:.2f}%, 복수 분류 {MULTI_CATEGORY_SHARE:.2f}%','기술분류와 사업은 N:M 구조로 설계함'],
['I-3','연구주체는 분산되나 국가역량은 집중',f'{UNIQUE_AUTHORS:,}명·{VALID_INSTITUTIONS:,}개, 주요 5개국 주도','시설·정책·기업·인력이 국가경쟁력을 좌우함'],
['I-4','협업은 보편적이나 국제협력은 선택적',f'공동저자 {COAUTHOR_SHARE:.1f}%, 국제협력 {INTL_SHARE:.1f}%','허브기관·공동인프라·신뢰체계를 전략적으로 확보함'],
['I-5','연구 중심에서 검증·표준·시장으로 전환','NIST·ITU·DARPA·기업 로드맵 강화','성능·비용·상호운용·보안을 평가축에 추가함'],
],[1.4,6.0,5.2,6.9],7.7)

# 7.3
add_heading(doc,'7.3 기술성숙도 진단: 연구 확대에서 상용화 검증 단계로 전환')
for txt,major in [
('분석결과는 양자기술이 단순 논문 증가 단계에서 오류정정·벤치마킹·표준화·산업적용을 검증하는 단계로 이동하고 있음을 시사함.',True),
('McKinsey는 2026년 300개 이상의 기업이 양자컴퓨팅을 채택·협업하고, 2025년 관련 기업 매출이 10억 달러를 넘었다고 제시하여 기업수요가 실험적 관심을 넘어 내부역량 구축으로 이동하고 있음을 보임 [R5].',False),
('DARPA QBI는 2033년까지 비용을 초과하는 계산가치를 가진 산업적으로 유용한 컴퓨터의 가능성을 독립 검증하는 체계를 운영하여 성능주장보다 검증가능성과 경제성을 중시함 [R6].',False),
('IBM은 2026년 말 근접 양자이점과 2029년 대규모 내결함성 양자컴퓨터를 목표로 제시하고 있어 컴퓨팅·오류정정·제어·HPC 연계의 시간표가 구체화됨 [R7].',False),
('미국의 2026년 행정명령은 컴퓨팅·센싱·네트워킹의 배치와 상용화, 제조·산업 파트너십을 강조하여 연구성과의 산업전환을 국가정책으로 명확히 함 [R3].',False),
]: add_bullet(doc,txt,major)

doc.add_picture(str(FIGS['timeline']),width=Cm(18.1)); add_caption(doc,'그림 7-2. 연구·표준·검증·상용화 전환의 주요 외부환경',True)
add_caption(doc,'표 7-3. 기술성숙 단계별 핵심 판단')
add_table(doc,['단계','현재 신호','핵심 병목','정책·산업 대응'],[
['원천연구','높은 논문량과 quant-ph 코어 유지','성능·재현성·플랫폼 간 비교','기초연구·공용 데이터·개방형 벤치마크를 유지함'],
['시스템 통합','복수 카테고리·다기관 연구 확대','소자·제어·SW·극저온·광학 결합','공동 인프라와 풀스택 과제를 설계함'],
['검증·표준','DARPA QBI, NIST·ITU 표준화','비용·유용성·상호운용·보안','독립 검증기관·시험인증·표준 연계를 강화함'],
['상용화·배치','기업 채택·정부 배치 정책 확대','사용사례 ROI·인력·레거시 통합','산업별 실증과 양자–HPC–AI 하이브리드를 확대함'],
['확산·생태계','기업·클러스터·공급망 목표 구체화','핵심부품·인력·IP·수출통제','공급망·인력·국제협력 거버넌스를 구축함'],
],[3.2,5.8,5.8,6.0],7.8)

# 7.4
add_heading(doc,'7.4 기술축별 전략 포트폴리오')
add_bullet(doc,'제4장의 카테고리 구조를 양자컴퓨팅·SW/AI, 양자통신·보안, 양자센싱·계측, 양자소자·소재, 양자이론·시뮬레이션, 양자화학·산업응용의 6대 기술축으로 재구성함.',True)
add_bullet(doc,'연구기반, 최근 성장, 시장·정책·표준화 준비도, 협업·인프라 의존도, 기술·사업 리스크를 5점 척도로 평가함.')
add_bullet(doc,'점수는 기술축 간 전략적 비교를 위한 전문가 통합평가이며 원자료의 절대 성능이나 시장규모를 직접 의미하지 않음.')

doc.add_picture(str(FIGS['portfolio']),width=Cm(17.6)); add_caption(doc,'그림 7-3. 양자 기술축 전략 포트폴리오',True)
add_caption(doc,'표 7-4. 기술축 통합평가 기준')
add_table(doc,['평가 축','정의','주요 근거','해석 기준'],[
['연구기반','논문 규모·분야 코어·연구주체 저변','제3~5장 규모·카테고리·주체 결과','5점일수록 연구기반이 광범위함'],
['성장 모멘텀','전·후반기·최근 증가와 교차분야 확장','제3~4장 성장·카테고리 변화','5점일수록 최근 확대가 뚜렷함'],
['시장·정책 준비도','기업채택·정부전략·표준·로드맵','NQI·EU·NIST·McKinsey·기업 로드맵','5점일수록 상용화 창구가 구체적임'],
['협업 의존도','다기관·국제협력·시설·공급망 필요성','제6장 협업 및 기술 특성','5점일수록 컨소시엄·공동인프라가 필수임'],
['리스크','기술불확실성·비용·표준·IP·공급망 위험','외부 로드맵과 전문가 해석','5점일수록 단계적 검증이 필요함'],
],[3.4,5.5,6.0,5.8],7.7)

add_caption(doc,'표 7-5. 6대 기술축 전략 포트폴리오')
tech_rows=[]
for name,research,growth,market,collab,risk,priority in TECH_PORTFOLIO:
    action={
        '양자컴퓨팅·SW/AI':'오류정정·컴파일러·하이브리드·산업문제 실증을 묶어 추진함',
        '양자통신·보안':'PQC 전환과 QKD·양자네트워크 표준·실증을 병렬 추진함',
        '양자센싱·계측':'의료·국방·반도체·항법의 현장성능과 시험인증을 확보함',
        '양자소자·소재':'큐비트·광자·극저온·패키징의 공급망과 공동시설을 확보함',
        '양자이론·시뮬레이션':'장기 원천성과와 알고리즘·검증문제의 개방형 기반을 유지함',
        '양자화학·산업응용':'제약·소재·에너지의 고가치 문제 중심으로 단계적 PoC를 수행함',
    }[name]
    tech_rows.append([name,f'{research:.1f}',f'{growth:.1f}',f'{market:.1f}',f'{collab:.1f}',f'{risk:.1f}',priority,action])
add_table(doc,['기술축','연구','성장','준비도','협업','리스크','우선군','핵심 실행방향'],tech_rows,[4.3,1.5,1.5,1.6,1.5,1.6,2.3,6.4],7.2)

# 7.5
add_heading(doc,'7.5 글로벌 경쟁·협력 구조의 전략적 해석')
doc.add_picture(str(FIGS['country']),width=Cm(15.7)); add_caption(doc,'그림 7-4. 주요국 연구 주도–참여–허브 포지션',True)
for txt,major in [
('미국은 제1저자 453건, 기관 국가 524건과 최상위 협업 허브성을 동시에 보여 연구·인프라·산업·국제네트워크를 결합한 종합 선도국으로 판단함.',True),
('중국은 제1저자 277건과 기관 국가 296건으로 연구 규모와 자체 주도성이 높으며 미국·홍콩·싱가포르·영국과의 선택적 연결을 형성함.',False),
('독일과 영국은 연구량 대비 기관 참여 확대와 국제협업 허브성이 크게 나타나 유럽의 분산형 인프라·컨소시엄을 연결하는 네트워크형 선도국으로 보임.',False),
('일본은 제1저자 112건과 기관 국가 124건으로 자체 주도형 기반이 강하고 하드웨어·소재·산업화 정책과 결합할 가능성이 큼.',False),
('한국은 제1저자 44건과 기관 국가 48건으로 기본 연구기반은 존재하나 상위 허브국 대비 네트워크 외연과 절대 규모가 작아 기술 특화와 국제 연계의 동시 강화가 필요함.',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-6. 주요국 유형과 협력전략')
add_table(doc,['국가 유형','대표 국가','정량 신호','전략적 특징','한국의 대응'],[
['종합 선도·글로벌 허브','미국','주도 453·참여 524·허브 133','연구·산업·인프라·동맹을 통합함','공동인프라·벤치마크·산업실증 파트너로 활용함'],
['규모·자체주도형','중국','주도 277·참여 296','대규모 국가투자와 자체 연구망이 강함','공개영역 동향·공급망·표준경쟁을 면밀히 관찰함'],
['네트워크·인프라 허브','독일·영국','주도 대비 참여·허브성이 큼','다국가 인프라·컨소시엄·산업연계에 강점','EU 공동과제·시험인증·인력교류를 확대함'],
['하드웨어·산업기반형','일본','주도 112·참여 124','소재·정밀제조·산학연 연계 잠재력이 큼','소자·소재·센싱 공동개발을 추진함'],
['성장·집중형','한국','주도 44·참여 48','기본역량은 있으나 규모·허브성 보완 필요','선택적 기술특화와 허브국 연계를 병행함'],
],[3.5,3.6,4.4,5.6,6.0],7.5)

# 7.6
add_heading(doc,'7.6 시장·정책·표준화 신호와 전략창구')
for txt,major in [
('미국 NQI는 FY2025 QIS R&D 예산으로 9.98억 달러를 요청하고 과학·인력·산업·인프라·안보·국제협력을 정책축으로 운영하여 장기 공공투자의 기반을 유지함 [R1].',True),
('EU Quantum Europe Strategy는 연구혁신, 양자인프라, 생태계·공급망·산업화, 우주·듀얼유스, 인력의 5개 영역을 제시하여 과학성과의 시장전환을 통합적으로 추진함 [R2].',False),
('NIST가 승인한 FIPS 203·204·205는 양자내성암호가 연구·후보 알고리즘 단계에서 실제 전환표준 단계로 이동했음을 의미함 [R4].',False),
('ITU-T Y.Sup98은 양자네트워크의 아키텍처·프로토콜·성능·전송·보안·IoT 적용을 여러 Study Group이 분담해야 한다고 제시하여 양자통신의 표준화가 다기관 협업형 과제임을 보여줌 [R8].',False),
('한국의 제1차 양자종합계획은 2035년 전문인력 1만 명, 양자기업 2천 개, 글로벌 표준 선도 3위권과 컴퓨팅·통신·센서·소재부품·알고리즘 클러스터를 제시함 [R9].',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-7. 최신 외부환경과 분석결과의 연결')
add_table(doc,['외부 신호','핵심 내용','본 분석과의 연결','전략 창구'],[
['지속 공공투자','미국·EU·영국·한국의 장기전략과 인프라 투자','안정적 논문량과 주요국 연구 집중을 뒷받침함','공동센터·공용시설·인력프로그램'],
['상용화 전환','기업 300+ 채택·미국의 배치·상용화 정책','최근 연구 모멘텀과 응용 카테고리 확장을 설명함','산업별 PoC·하이브리드 워크플로'],
['검증 강화','DARPA QBI와 기업 로드맵 구체화','성과평가가 논문 수에서 유용성·비용으로 이동함','독립 벤치마크·시험평가·경제성'],
['표준화 창구','NIST PQC·ITU 양자네트워크','암호/보안·광학·네트워크 교차분야 가치가 상승함','표준기고·상호운용·마이그레이션'],
['생태계·공급망','소자·광자·극저온·제어·클러스터 정책','다기관 협업과 소자·소재 분야의 중요성을 강화함','공급망 맵·공동파운드리·시험인증'],
],[3.4,5.7,6.4,5.5],7.7)

# 7.7
add_heading(doc,'7.7 한국형 투자 우선순위와 포트폴리오 원칙')
doc.add_picture(str(FIGS['pyramid']),width=Cm(16.2)); add_caption(doc,'그림 7-5. 한국형 양자기술 투자 우선순위',True)
for txt,major in [
('1단계는 시장·표준·정책 창구가 구체화된 컴퓨팅·SW/AI, 통신·보안, 센싱·계측을 대상으로 실증·표준·산업수요 연계를 가속함.',True),
('2단계는 전 기술축의 성능과 자립성을 좌우하는 소자·소재·극저온·광자·제어·패키징을 국가 공용인프라와 공급망 관점에서 확보함.',False),
('3단계는 이론·시뮬레이션과 화학·산업응용을 장기 원천성과 및 고가치 산업문제 발굴의 이중 트랙으로 유지함.',False),
('우선순위는 고정하지 않고 최근 12개월 성장, 국제협력, 표준화, 기업수요, 기술검증 결과에 따라 매년 재배분함.',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-8. 전략목표와 핵심 실행과제')
add_table(doc,['전략목표','핵심 문제','우선 실행과제','기대효과'],[
['기술 포트폴리오 정교화','학술 카테고리와 산업기술 간 불일치','6대 기술축·세부기술 정의서·검색식·전문가 검증을 구축함','투자 중복을 줄이고 유망기술을 조기 식별함'],
['검증 중심 R&D 전환','논문 수와 실제 유용성 간 격차','성능·비용·재현성·상호운용·안전 기준을 과제평가에 반영함','상용화 가능성과 실패 조기탐지가 향상됨'],
['공동 인프라·공급망','장비·공정·부품의 높은 진입장벽','극저온·광원·파운드리·제어·테스트베드를 공동 구축함','중소 연구팀과 기업의 진입비용이 낮아짐'],
['국제협력 고도화','논문 국제협력률과 허브성이 제한적','허브기관·브리지 연구자·공동표준·공동실증을 연계함','네트워크 중심성과 기술흡수력이 강화됨'],
['산업수요·인력 연결','연구성과와 기업 워크플로 간 단절','양자–HPC–AI 실증, 공동박사·포닥·산업파견을 확대함','기업 채택과 전문인력 순환이 촉진됨'],
['데이터 기반 거버넌스','DB·기관명·기술분류의 품질 편차','논문–특허–과제–표준–시장 통합 모니터링을 운영함','근거기반 투자·정책 조정이 가능해짐'],
],[3.6,5.3,7.0,5.4],7.6)

# 7.8
add_heading(doc,'7.8 기술축별 실행전략')
add_caption(doc,'표 7-9. 6대 기술축별 중점 전략')
add_table(doc,['기술축','단기 중점','중기 중점','핵심 KPI'],[
['양자컴퓨팅·SW/AI','오류정정·컴파일러·디코더·벤치마크 DB를 정비함','HPC·AI 연계 산업문제와 독립 검증체계를 구축함','논리오류율, 회로성능, 재현성, 산업 PoC'],
['양자통신·보안','PQC 전환대상·QKD·양자망 표준을 매핑함','통신망 상호운용·금융·국방 실증을 확대함','표준기고, 전환율, 상호운용, 실증망'],
['양자센싱·계측','의료·항법·반도체·국방 사용사례를 선별함','현장환경 성능·시험인증·공공조달을 연계함','감도·정확도·안정성·현장실증'],
['양자소자·소재','핵심부품·소재·장비 공급망과 기술격차를 분석함','공동파운드리·패키징·극저온·제어 생태계를 구축함','수율·결함·국산화·시설활용률'],
['양자이론·시뮬레이션','장기 기초과제와 개방형 검증문제를 유지함','알고리즘·물성·에너지·우주 분야로 확장함','고영향 논문·오픈코드·후속응용'],
['양자화학·산업응용','산업별 고가치 난제를 정의하고 데이터 준비도를 평가함','제약·소재·에너지 PoC와 경제성 검증을 수행함','문제축소율·예측개선·ROI·특허'],
],[4.3,6.1,6.2,5.0],7.4)

# 7.9
add_heading(doc,'7.9 36개월 실행 로드맵')
doc.add_picture(str(FIGS['roadmap']),width=Cm(18.1)); add_caption(doc,'그림 7-6. 데이터·기술·협력·인력·성과관리 통합 로드맵',True)
for txt,major in [
('0~6개월은 기관·저자·국가 표준화, 실무 기술분류, 핵심기술 후보와 기준선 KPI를 확정하는 정비 단계로 운영함.',True),
('6~18개월은 논문–특허–표준–시장 연계, 공동과제·시설공유, 독립 벤치마크와 산업 PoC를 수행하는 검증 단계로 운영함.',False),
('18~36개월은 기술군별 성과를 기반으로 사업화·표준·클러스터·국제 공급망을 확대하고 포트폴리오를 재배분함.',False),
('각 단계 종료 시 기술성·시장성·협업성·데이터 신뢰도를 재평가하고 실패 가능성이 높은 과제는 축소·전환함.',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-10. 단계별 주요 산출물')
add_table(doc,['단계','핵심 산출물','의사결정 포인트','책임 주체'],[
['0~6개월','통합 DB, 6대 기술축 정의서, 핵심주체·허브맵, KPI 기준선','기술후보의 포함·제외와 우선순위를 확정함','전담 PMO·데이터팀·분야전문가'],
['6~18개월','벤치마크, 실증과제, 공동인프라 운영안, 국제협력·표준화 계획','기술성·경제성·파트너 적합성을 중간검증함','R&D기관·기업·시험기관·표준전문가'],
['18~36개월','산업확산 모델, 공동표준·특허, 공급망·클러스터, 투자재배분안','상용화·확대·중단 여부를 결정함','정부·기업·투자기관·국제파트너'],
],[3.0,7.3,6.3,5.0],7.8)

# 7.10
add_heading(doc,'7.10 상시 모니터링 KPI와 의사결정 체계')
add_bullet(doc,'연구성과는 논문 수 단일지표에서 성장·융합·영향력·협업·표준·실증·시장성과를 결합한 균형지표로 전환함.',True)
add_bullet(doc,'선행지표는 최근 3개월 성장률, 신규 키워드·카테고리, 신규 기관·저자, 국제협력 증가, 표준 신규과제를 활용함.')
add_bullet(doc,'중간지표는 독립 벤치마크, 재현성, 시설 활용, 공동특허·표준, 인력 이동, 산업 PoC를 활용함.')
add_bullet(doc,'후행지표는 매출·투자·조달·제품배치·공급망 자립도·산업생산성 개선을 활용함.')

add_caption(doc,'표 7-11. 통합 모니터링 KPI')
add_table(doc,['영역','핵심 지표','정의·계산','활용 기준'],[
['연구활동','최근 3개월 모멘텀','최근 3개월 평균 / 전체 월평균 - 1','증가 지속성과 조정구간을 판단함'],
['융합성','복수 카테고리·교차강도','복수분류 논문 비율, 카테고리 동시출현','기술융합과 응용확장을 탐지함'],
['기술품질','영향력·재현성·벤치마크','인용·코드·데이터·독립검증 통과율','양적 성장과 실질 성능을 분리함'],
['주체역량','반복 연구자·기관 전문화','최근 참여·분야특화·인력 이동','핵심인력과 인프라 거점을 선정함'],
['협업성','다기관·국제협력·중심성','협업률·degree·betweenness·지속성','허브·브리지·협력공백을 식별함'],
['표준·IP','표준기고·SEP·공동특허','기고·채택·특허군·라이선스','시장 규칙 선점 여부를 판단함'],
['상용화','PoC·배치·투자·매출','실증 성공·사용사례·조달·기업투자','확대·전환·중단을 결정함'],
['데이터품질','DOI·기관·국가·분류 정확도','커버리지·오류율·표준화율','분석 신뢰도와 재작업 우선순위를 관리함'],
],[3.1,5.1,6.6,5.8],7.6)

# 7.11
add_heading(doc,'7.11 주요 위험요인과 대응방안')
add_caption(doc,'표 7-12. 전략 실행 리스크 레지스터')
add_table(doc,['리스크','발생 가능성','영향도','주요 징후','대응방안'],[
['기술 일정 지연','높음','매우 높음','오류정정·수율·비용 목표 반복 연기','단계형 투자·독립 검증·대체 플랫폼 포트폴리오를 운영함'],
['시장 과대기대','중간','높음','PoC 증가 대비 실제 워크플로·ROI 부족','사용사례 가치·통합비용·기회비용을 사전평가함'],
['표준 분절','중간','높음','지역·기관별 상이한 프로토콜과 인증','ITU·ISO/IEC·ETSI·NIST 연계와 상호운용 시험을 강화함'],
['핵심부품 공급망','높음','높음','극저온·광원·공정·제어 장비 편중','다중공급·공동시설·핵심부품 R&D를 추진함'],
['인력 미스매치','높음','높음','물리·공학·SW·산업인력 간 단절','융합교육·산업파견·공동박사·재교육을 확대함'],
['연구보안·수출통제','중간','매우 높음','공동연구 승인·데이터·장비 이동 지연','협력 등급·IP·보안·수출통제 절차를 사전 설계함'],
['데이터 편향','중간','중간','arXiv·검색식·기관표준화에 따른 순위 변동','다중 DB·민감도 분석·전문가 검증을 병행함'],
],[3.8,2.8,2.8,6.2,6.3],7.4)

# 7.12
add_heading(doc,'7.12 데이터 해석의 한계와 보완 원칙')
for txt,major in [
('arXiv 데이터는 연구활동의 빠른 선행신호이나 동료평가 완료, 기술성능, 시장성과를 직접 의미하지 않음.',True),
('제목 검색식에 QUANTIZATION이 포함되어 AI 모델 양자화 등 비핵심 논문이 일부 포함될 수 있으므로 초록·전문 기반 정합성 검증이 필요함.',False),
('기관·국가 정보와 DOI의 커버리지 한계가 연구주체·국제협력 지표를 과소 또는 왜곡할 수 있음.',False),
('12개월 데이터는 장기 계절성·기술주기·협업 지속성을 확정하기 어려우므로 3~5년 장기분석과 최근 12개월 모멘텀을 병렬 적용함.',False),
('시장 전망과 기업 로드맵은 목표·시나리오이므로 독립 성능검증과 실제 배치·매출·사용사례 데이터로 지속 점검함.',False),
]: add_bullet(doc,txt,major)

add_caption(doc,'표 7-13. 종합분석 신뢰도 향상을 위한 보완과제')
add_table(doc,['보완 축','현재 한계','우선 보완','기대효과'],[
['데이터 범위','arXiv 중심·12개월','Crossref·OpenAlex·WoS/Scopus·특허·과제 연계','장기성과와 품질을 교차검증함'],
['기술분류','학술 카테고리와 산업기술 불일치','초록·전문 임베딩·N:M 분류·전문가 정의서','실무 기술군 규모와 융합관계를 정교화함'],
['주체표준화','저자·기관 이명·동명이인','ORCID·ROR·소속·공동저자 기반 식별','핵심인력·기관·네트워크 정확도가 향상됨'],
['성과품질','논문 수 중심','인용·재현성·벤치마크·오픈소스·실증 연계','실제 기술진전과 양적 증가를 구분함'],
['시장·정책','외부자료의 시점·가정 차이','정기 업데이트·다중전망·시나리오 분석','전망 편향과 정책 변화에 대응함'],
],[3.4,5.8,6.5,6.1],7.6)

# 7.13
add_heading(doc,'7.13 Executive Summary 연결용 핵심 메시지')
add_bullet(doc,'양자 분야는 12개월 2,400건·월평균 200편의 높은 연구기반과 최근 3개월 217편의 재상승 모멘텀을 보임.',True)
add_bullet(doc,'양자물리 코어 70.00%와 복수 카테고리 47.42%가 공존하여 원천과학 기반 위에서 AI·광학·보안·소재·응용이 확장됨.')
add_bullet(doc,'저자 10,514명·기관 1,995개의 넓은 생태계와 미국·중국·독일·영국 중심의 국가집중이 동시에 나타남.')
add_bullet(doc,'공동저자 88.0%에 비해 국제협력 10.0%로 나타나 국내·기관 내부 협업은 강하지만 글로벌 허브 연계는 전략적으로 확대할 필요가 있음.')
add_bullet(doc,'시장·정책 환경은 연구 중심에서 검증·표준·상용화·배치로 이동하고 있으며 컴퓨팅·SW/AI, 통신·보안, 센싱·계측이 우선 대응영역으로 판단됨.')
add_bullet(doc,'한국은 기술 전면추격보다 선택적 특화, 공동인프라, 국제 허브 연계, 표준·시험인증, 논문–특허–시장 통합모니터링을 병행하는 전략이 적합함.')

add_caption(doc,'표 7-14. 최종 전략메시지와 실행 우선순위')
add_table(doc,['우선순위','전략 메시지','즉시 실행','중장기 목표'],[
['1','연구량을 기술검증·시장성과로 전환함','6대 기술축·벤치마크·산업 PoC를 확정함','상용화·배치 가능한 성과 포트폴리오를 구축함'],
['2','표준과 인프라를 기술경쟁력의 핵심축으로 관리함','PQC·양자망·시험인증·공동시설 계획을 수립함','국제표준·공급망·공용인프라의 선도지위를 확보함'],
['3','국제허브와 국내 네트워크를 동시에 강화함','허브기관·브리지 연구자·공동과제를 발굴함','지속적 국가·기관 협업망과 인력순환을 구축함'],
['4','데이터 기반 포트폴리오 거버넌스를 운영함','DB 정비·KPI·분기 모니터링을 시작함','투자 확대·전환·중단을 정량근거로 결정함'],
['5','장기기반과 산업응용을 균형 있게 유지함','이론·소재 기반과 고가치 응용문제를 병렬 선정함','원천경쟁력과 산업파급을 동시에 확보함'],
],[2.1,5.8,6.8,7.1],7.7)

# References
add_heading(doc,'참고자료')
refs=[
('[R1]','National Quantum Initiative, Annual Report FY2025','https://www.quantum.gov/wp-content/uploads/2024/12/NQI-Annual-Report-FY2025.pdf'),
('[R2]','European Commission, Quantum Europe Strategy, 2025','https://digital-strategy.ec.europa.eu/en/library/quantum-europe-strategy'),
('[R3]','The White House, Executive Order 14413: Ushering in the Next Frontier of Quantum Innovation, 2026','https://www.whitehouse.gov/presidential-actions/2026/06/ushering-in-the-next-frontier-of-quantum-innovation/'),
('[R4]','NIST CSRC, Post-Quantum Cryptography FIPS Approved','https://csrc.nist.gov/news/2024/postquantum-cryptography-fips-approved'),
('[R5]','McKinsey, Quantum Technology Monitor 2026: A Commercial Tipping Point','https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/mckinsey-quantum-technology-monitor-2026-a-commercial-tipping-point'),
('[R6]','DARPA, Quantum Benchmarking Initiative','https://www.darpa.mil/research/programs/quantum-benchmarking-initiative'),
('[R7]','IBM Quantum, Hardware and Roadmap','https://www.ibm.com/quantum/hardware'),
('[R8]','ITU-T Y.Sup98, Technical Considerations Towards Quantum Networks, 2025','https://www.itu.int/rec/dologin_pub.asp?id=T-REC-Y.Sup98-202511-I!!PDF-E&lang=s&type=items'),
('[R9]','과학기술정보통신부, 제1차 양자종합계획, 2026','https://www.msit.go.kr/eng/bbs/view.do?bbsSeqNo=42&mId=4&mPid=2&nttSeqNo=1222&sCode=eng'),
]
for rid,title,url in refs:
    p=doc.add_paragraph(); pformat(p,after=3,line=1.15); p.paragraph_format.left_indent=Cm(.42); p.paragraph_format.first_line_indent=Cm(-.42); r=p.add_run(rid+' '); set_run(r,8.8,True,NAVY); add_hyperlink(p,title,url)

cp=doc.core_properties; cp.title='arXiv 양자 분야 제7장 종합 인사이트 및 전략적 시사점'; cp.subject='제2~6장 정량결과와 최신 시장·기술·정책·표준화 자료의 통합 분석'; cp.author='OpenAI'; cp.keywords='arXiv, 양자, 종합 인사이트, 전략, 기술 포트폴리오, 정책, 시장, 표준화'
settings=doc.settings._element; update=OxmlElement('w:updateFields'); update.set(qn('w:val'),'true'); settings.append(update)
doc.save(OUTPUT)
print(f'Created {OUTPUT}')
